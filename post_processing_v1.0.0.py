#%% Import

import pandas as pd
import time
from io import BytesIO # 파일처럼 취급되는 문자열 객체 생성(메모리 낭비 down)
import multiprocessing as mp

import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docx.oxml.ns import qn

import PBD_p3d as pbd

#%% 시간 측정(START)
time_start = time.time()

print('\n########## Result ##########\n')

###############################################################################
###############################################################################
#%% User Input

# Building Name
bldg_name = '105동'

# Analysis Result
result_path = r'K:\2105-이형우\from 박재성\test\Results_E.Column'
result_xlsx = 'Analysis Result' # 해석결과에 공통으로 포함되는 이름 (확장자X)

# Data Conversion Sheet, Column Sheet, Beam Sheet
input_path = r'K:\2105-이형우\from 박재성\test\Results_E.Column'
input_xlsx = 'SW-105D_Data Conversion(E.Column)_Ver.1.1_230117.xlsx'
column_xlsx = 'Results_E.Column(105D)_Ver.1.3.xlsx'
beam_xlsx = 'Results_E.Beam(103_16).xlsx'

# Post-processed Result
output_path = result_path # result_path와 동일하게 설정. 바꿔도 됨
output_docx = '105 해석결과(E.Column).docx'
column_pdf_name = '105D E.Column 결과' # 출력될 pdf 파일의 이름 (확장자X)
appendix_docx = 'Appendix. Wall SF(elementwise).docx'

c_beam_group = 'C.Beam'
g_col_group = 'G.Column'

# Base Shear
ylim = 30000 #kN, 그래프의 y limit

# Story Shear
xlim = 30000 # kN, 그래프의 x limit
story_gap = 2 # 층간격

# BR
yticks = 2

#%% Post Processing

# 밑면 전단력
# base_SF = pbd.base_SF(result_path, ylim=ylim)

# 층 전단력
# story_SF = pbd.story_SF(input_path, input_xlsx, result_path\
                            # , yticks=story_gap, xlim=xlim)

# 층간변위비
# IDR = pbd.IDR(input_path, input_xlsx, result_path, yticks=story_gap)

# 벽체 압축/인장 변형률
# AS = pbd.AS(input_path, input_xlsx, result_path, yticks=story_gap)

# 벽체 전단강도
# wall_SF = pbd.wall_SF(input_path, input_xlsx, result_path, graph=True\
                        # , yticks=story_gap, xlim=3)

# 벽체 전단력(only graph)
# wall_SF_graph = pbd.wall_SF_graph(input_path, input_xlsx, yticks=story_gap)

# 벽체 소성회전각
# SWR = pbd.SWR(input_path, input_xlsx, result_path, yticks=story_gap\
                # , DE_criteria=0.002, MCE_criteria=0.004/1.2)

# 벽체 소성회전각(DCR)
# SWR_DCR = pbd.SWR_DCR(input_path, input_xlsx, result_path\
                        # , yticks=story_gap, xlim=3)
# 
# 연결보 소성회전각(Gage 설치 X)    
# BR_no_gage = pbd.BR_no_gage(result_path, result_xlsx, input_path\
                                # , input_xlsx, cri_DE=0.01, cri_MCE=0.025/1.2\
                                # , yticks=2, xlim=0.03)

# 연결보 소성회전각(DCR)
# BR_DCR = pbd.BR_DCR(result_path, result_xlsx, input_path, input_xlsx
                    # , yticks=3, xlim=3)

# 전이보 전단력
# trans_beam_SF = pbd.trans_beam_SF_2(result_path, result_xlsx, input_path\
                                        # , input_xlsx, beam_xlsx, contour=True)

# 전이기둥 전단강도
trans_column_SF = pbd.trans_column_SF(result_path, result_xlsx, input_path\
                                , input_xlsx, column_xlsx, export_to_pdf=True\
                                    , pdf_name=column_pdf_name)

# 전이기둥 전단강도(only pdf)
# trans_column_SF_pdf = pbd.trans_column_SF_pdf(input_path, column_xlsx\
                                                   # , pdf_name=column_pdf_name)

# 일반기둥 전단강도                                                  
# gen_column_SF = pbd.general_column_SF(result_path, result_xlsx, input_path, input_xlsx)

# 일반기둥 소성회전각(DCR)
# CR = pbd.CR(result_path, result_xlsx, input_path, input_xlsx
            # , yticks=2)  

# 일반기둥 소성회전각(DCR)
# CR_DCR = pbd.CR_DCR(result_path, result_xlsx, input_path, input_xlsx
                    # , yticks=2, xlim=3)                             
                                                  
# Pushover
# pushover = system.pushover(result_path, x_result_txt, y_result_txt, base_SF_design, pp_x=, pp_y=)



###############################################################################
###############################################################################
#########################     will be depriecated     #########################
# =============================================================================
# 연결보 소성회전각
# BR = pbd.BR(result_path, result_xlsx, input_path, input_xlsx\
#               , m_hinge_group_name, s_hinge_group_name=s_hinge_group_name\
#               , s_cri_DE=0.01, s_cri_MCE=0.025/1.2, yticks=2, xlim=0.03)
# 
# 전이보 전단력(구버전)
# trans_beam_SF_old = pbd.trans_beam_SF(result_path, result_xlsx\
#                                         , input_path, input_xlsx, beam_xlsx)
# =============================================================================

#%% 그래프 & df 리스트 만들기

function_list = ['base_SF', 'story_SF', 'IDR', 'AS', 'wall_SF'\
                 , 'wall_SF_graph', 'SWR', 'SWR_DCR', 'BR', 'BR_no_gage'\
                 , 'BR_DCR', 'CR', 'CR_DCR', 'trans_beam_SF', 'trans_column_SF', 'pushover']

# plot & dataframe 합쳐진 list 만들기
plot_df_list = []
for i in function_list:
    if i in globals():
        plot_df_list.append(globals()['{}'.format(i)])
        
plot_df_list = list(filter(None, plot_df_list))

# generator -> tuple
plot_df_tuple_list = []
for i in plot_df_list:
    if not isinstance(i, tuple):
        plot_df_tuple_list.append(tuple(i))        
    else:
        plot_df_tuple_list.append(i)
        
# =============================================================================
# generator -> tuple
# def create_plot_df(plot_df_list):
#     plot_df_tuple_list = []    
#     for i in plot_df_list:
#         if not isinstance(i, tuple):
#             plot_df_tuple_list.append(tuple(i))        
#         else:
#             plot_df_tuple_list.append(i)            
#     return plot_df_tuple_list
# 
# Multiprocessing으로 figure append
# num_core = mp.cpu_count()
# 
# pool = mp.Pool(processes = num_core)
# plot_df_mp_result = pool.map(create_plot_df, plot_df_list)
# pool.close()
# pool.join()
# =============================================================================


# tuple을 list로 펼치기
plot_df_list_flat = list(sum(plot_df_tuple_list, ()))

# 전체 plot
plot_list = [x for x in plot_df_list_flat if not isinstance(x, pd.DataFrame)]

# 전체 dataframe
df_list = [x for x in plot_df_list_flat if isinstance(x, pd.DataFrame)]

#%% Word로 결과 정리

# Document 생성
output_word = docx.Document()

# Changing the page margins
output_word_sections = output_word.sections
for section in output_word_sections:
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(0.44)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

# 제목
output_word_title_para = output_word.add_paragraph()
output_word_title_run = output_word_title_para.add_run(bldg_name)
output_word_title_run.font.size = Pt(12)
output_word_title_run.bold = True

# 표 삽입  # int(-(-x//1)) = math.ceil()
output_word_table = output_word.add_table(int(-(-len(plot_list)//2)), 2)
output_word_table_faster = output_word_table._cells

# 그래프 사이즈(inch)
figsize_x, figsize_y = 8.7, 2.3 # cm

#%% 전체 결과 그래프 그리기

count = 0
for i in plot_list:
    
    memfile = BytesIO()
    i.savefig(memfile)
    
    output_word_table_faster_para = output_word_table_faster[count]\
                                    .paragraphs[0]
    output_word_table_faster_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    output_word_table_faster_run = output_word_table_faster_para.add_run()
    # output_word_table_faster_run.add_picture(memfile, height = figsize_y\
    #                                          , width = figsize_x)
    # output_word_table_faster_run.add_picture(memfile, width=Cm(figsize_x)
    output_word_table_faster_run.add_picture(memfile, width=Cm(figsize_x))
    output_word_table_faster_run.add_break(WD_BREAK.PAGE)
    
    memfile.close()    
    count += 1

# Table 스타일  
output_word_table.style = 'Table Grid'
output_word_table.autofit = False
output_word_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 스타일 지정(global)
output_word_style = output_word.styles['Normal']
output_word_style.font.name = '맑은 고딕'
output_word_style._element.rPr.rFonts\
    .set(qn('w:eastAsia'), '맑은 고딕') # 한글 폰트를 따로 설정해 준다
output_word_style.font.size = Pt(8) 
        
# 저장~
output_word.save(output_path + '\\' + output_docx)
'''
#%% Appendix

# Document 생성
appendix_word = docx.Document()

# Changing the page margins
appendix_word_sections = appendix_word.sections
for section in appendix_word_sections:
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(0.44)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

# 제목
appendix_word_title_para = appendix_word.add_paragraph()
appendix_word_title_run = appendix_word_title_para.add_run(bldg_name+'(Appendix)')
appendix_word_title_run.font.size = Pt(12)
appendix_word_title_run.bold = True

# 표 삽입  # int(-(-x//1)) = math.ceil()
appendix_word_table = appendix_word.add_table(int(-(-len(plot_list)//2)), 2)
appendix_word_table_faster = appendix_word_table._cells


#%% 전체 결과 그래프 그리기

# 지진파별 그래프
count = 1

for i in input_wall_name:  
    
    shear_force_major_max_temp = shear_force_major_max[(shear_force_major_max.index.str.contains(i + '_'))\
                                                       & (shear_force_major_max.index.str[0] == i[0])]

    story_temp = shear_force_major_max_temp.index.tolist()
    story_temp = pd.Series(list(map(lambda x: x.rsplit('_', 1)[1], story_temp))).tolist()
    
    ### DE
    memfile = BytesIO()
    plt.figure(count, dpi=150, figsize=(5, 4.8))
    # plt.xlim(0, story_shear_xlim)
    
    # 지진파별 plot
    for j in range(14):
        plt.plot(shear_force_major_max_temp.iloc[:,j], range(shear_force_major_max_temp.shape[0]), label=seismic_load_name_list[j], linewidth=0.7)
        
    # 평균 plot
    plt.plot(shear_force_major_max_temp.iloc[:,0:14].mean(axis=1), range(shear_force_major_max_temp.shape[0]), color='k', label='Average', linewidth=2)
    
    plt.yticks(range(shear_force_major_max_temp.shape[0])[::story_shear_yticks], story_temp[::story_shear_yticks])
    # plt.xticks(range(14), range(1,15))
    
    # 기타
    plt.grid(linestyle='-.')
    plt.xlabel('Story Shear(kN)')
    plt.ylabel('Story')
    plt.legend(loc=1, fontsize=8)
    plt.title('{}'.format(i.split('_')[0]) + ' (DE)')
    
    plt.tight_layout()
    plt.savefig(memfile)
    plt.close()
    
    SF_word_table_faster_para = SF_word_table_faster[count-1].paragraphs[0]
    SF_word_table_faster_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    SF_word_table_faster_run = SF_word_table_faster_para.add_run()
    # SF_word_table_faster_run.add_picture(memfile, height = figsize_y, width = figsize_x)
    SF_word_table_faster_run.add_picture(memfile, width=Cm(figsize_x))

    memfile.close()
    count += 1
    
    ### MCE
    memfile2 = BytesIO()
    plt.figure(count, dpi=150, figsize=(5, 4.8))
    # plt.xlim(0, story_shear_xlim)
    
    # 지진파별 plot
    for j in range(14):
        plt.plot(shear_force_major_max_temp.iloc[:,j+14], range(shear_force_major_max_temp.shape[0]), label=seismic_load_name_list[j+14], linewidth=0.7)
        
    # 평균 plot
    plt.plot(shear_force_major_max_temp.iloc[:,14:28].mean(axis=1), range(shear_force_major_max_temp.shape[0]), color='k', label='Average', linewidth=2)
    
    plt.yticks(range(shear_force_major_max_temp.shape[0])[::story_shear_yticks], story_temp[::story_shear_yticks])
    # plt.xticks(range(14), range(1,15))
    
    # 기타
    plt.grid(linestyle='-.')
    plt.xlabel('Story Shear(kN)')
    plt.ylabel('Story')
    plt.legend(loc=1, fontsize=8)
    plt.title('{}'.format(i.split('_')[0]) + ' (MCE)')
    
    plt.tight_layout()
    plt.savefig(memfile2)
    plt.close()
    
    SF_word_table_faster_para2 = SF_word_table_faster[count-1].paragraphs[0]
    SF_word_table_faster_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    SF_word_table_faster_run2 = SF_word_table_faster_para2.add_run()
    # SF_word_table_faster_run2.add_picture(memfile2, height = figsize_y, width = figsize_x)
    SF_word_table_faster_run2.add_picture(memfile2, width=Cm(figsize_x))
        
    memfile2.close()
    count += 1

# Table 스타일  
appendix_word_table.style = 'Table Grid'
appendix_word_table.autofit = False
appendix_word_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 스타일 지정(global)
appendix_word_style = appendix_word.styles['Normal']
appendix_word_style.font.name = '맑은 고딕'
appendix_word_style._element.rPr.rFonts\
    .set(qn('w:eastAsia'), '맑은 고딕') # 한글 폰트를 따로 설정해 준다
appendix_word_style.font.size = Pt(8) 
        
# 저장~
appendix_word.save(output_path + '\\' + appendix_docx)

'''
#%% 시간 측정(END)
time_end = time.time()
time_run = (time_end-time_start)/60
print('\n', 'total time = %0.7f min' %(time_run))