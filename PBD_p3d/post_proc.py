import pandas as pd
import numpy as np
import os
import pickle
from io import BytesIO
import multiprocess as mp
import win32com.client
import pythoncom
from PyPDF2 import PdfMerger, PdfFileReader
import shutil

import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docx.oxml.ns import qn

# 자료형 비교 위해 불러올 것
import matplotlib.pyplot as plt
from decimal import Decimal, ROUND_UP

#%% Post-Processing Class

class PostProc():
    
    def __init__(self, input_xlsx_path, result_xlsx_path, get_base_SF=False
                 , get_story_SF=False, get_IDR=False, get_BR=False, get_BSF=False
                 , get_E_BSF=False, get_CR=False, get_CSF=False, get_E_CSF=False
                 , get_WAS=False, get_WR=False, get_WSF=False, BR_scale_factor=1.0) -> None:
        
        ##### Load Excel Files (Analysis Result Sheets)
        to_load_list = result_xlsx_path    

        ##### Excel 파일 읽는 Function (w/ Xlsx2csv & joblib)
        def read_excel(path:str, sheet_name:str, skip_rows:list=[0,2]) -> pd.DataFrame:
            import pandas as pd
            from io import StringIO # if not import, error occurs when using multiprocessing
            from xlsx2csv import Xlsx2csv
            data_buffer = StringIO()
            Xlsx2csv(path, outputencoding="utf-8", ignore_formats='float').convert(data_buffer, sheetname=sheet_name)
            data_buffer.seek(0)
            data_df = pd.read_csv(data_buffer, low_memory=False, skiprows=skip_rows)
            return data_df
        
        ##### Read Excel Files (Data Conversion Sheets & Analysis Result Sheets)
        # Story Info
        self.story_info = read_excel(input_xlsx_path, sheet_name='Story Data', skip_rows=[0,1,2])
        self.story_info = self.story_info.iloc[:,[0,1,2]]
        self.story_info.dropna(how='all', inplace=True)
        self.story_info.columns = ['Index', 'Story Name', 'Height(mm)']
        # Rebar Info
        self.rebar_info = read_excel(input_xlsx_path, sheet_name='ETC', skip_rows=[0,1,2])
        self.rebar_info = self.rebar_info.iloc[:,np.r_[0,3:12]]
        self.rebar_info.dropna(how='all', inplace=True)
        self.rebar_info.columns = ['Type', '일반용', '내진용', 'Story(from)-Vertical'
                                   , 'Story(to)-Vertical', 'Concrete-Vertical', 'Story(from)-Horizontal'
                                   , 'Story(to)-Horizontal', 'Concrete-Horizontal', 'Boundary']
        # Output_Wall Properties
        if (get_WAS == True) | (get_WR == True) | (get_WSF == True):
            self.wall_info = read_excel(input_xlsx_path, sheet_name='Input_S.Wall', skip_rows=[0,1,2])
            self.wall_info = self.wall_info.iloc[:,0:11]
            self.wall_info.dropna(how='all', inplace=True)
            self.wall_info.columns = ['Name', 'Length(mm)', 'Thickness(mm)', 'Concrete Grade', 'V.Rebar Type', 'V.Rebar(DXX)'
                                      , 'V.Rebar Spacing(mm)', 'V.Rebar EA', 'H.Rebar Type', 'H.Rebar(DXX)', 'H.Rebar Spacing(mm)']
        # C.Beam Info
        if (get_BR == True) | (get_BSF == True):
            self.beam_info = read_excel(input_xlsx_path, sheet_name='Input_C.Beam', skip_rows=[0,1,2])
            self.beam_info = self.beam_info.iloc[:,0:32]
            self.beam_info.dropna(how='all', inplace=True)
            self.beam_info.columns = ['Name', 'Length(mm)', 'b(mm)', 'h(mm)', 'd(mm)', 'Concrete Grade'
                                      , 'Arrangement', 'Seismic Detail', 'Main Rebar Type', 'Main Rebar(DXX)'
                                      , 'Stirrup Type', 'Stirrup(DXX)', 'X-Bracing Type', 'X-Bracing(DXX)'
                                      , 'Top(1)', 'Top(2)', 'Top(3)', 'Stirrup EA', 'Stirrup Space(mm)'
                                      , 'X-Bracing EA', 'X-Bracing deg', 'Main Rebar(DXX)_after', 'Stirrup(DXX)_after'
                                      , 'X-Bracing(DXX)_after', 'Top(1)_after', 'Top(2)_after', 'Top(3)_after'
                                      , 'Stirrup EA_after', 'Stirrup Space(mm)_after', 'X-Bracing EA_after', 'X-Bracing deg_after'
                                      , 'Boundary']
            # Divided C.Beam
            self.dbeam_info = read_excel(input_xlsx_path, sheet_name='Input_D.Beam', skip_rows=[0,1,2])
            self.dbeam_info = self.dbeam_info.iloc[:,0:33]
            self.dbeam_info.dropna(how='all', inplace=True)
            self.dbeam_info.columns = ['Name', 'Length(mm)', 'Element Length(mm)', 'b(mm)', 'h(mm)', 'd(mm)'
                                      , 'Concrete Grade', 'Arrangement', 'Seismic Detail', 'Main Rebar Type', 'Main Rebar(DXX)'
                                      , 'Stirrup Type', 'Stirrup(DXX)', 'X-Bracing Type', 'X-Bracing(DXX)'
                                      , 'Top(1)', 'Top(2)', 'Top(3)', 'Stirrup EA', 'Stirrup Space(mm)'
                                      , 'X-Bracing EA', 'X-Bracing deg', 'Main Rebar(DXX)_after', 'Stirrup(DXX)_after'
                                      , 'X-Bracing(DXX)_after', 'Top(1)_after', 'Top(2)_after', 'Top(3)_after'
                                      , 'Stirrup EA_after', 'Stirrup Space(mm)_after', 'X-Bracing EA_after', 'X-Bracing deg_after'
                                      , 'Boundary']
            
        # E.Column Info
        if get_E_CSF == True:
            self.ecol_info = read_excel(input_xlsx_path, sheet_name='Input_E.Column', skip_rows=[0,1,2])
            self.ecol_info = self.ecol_info.iloc[:,np.r_[0:15,18]]
            self.ecol_info.dropna(how='all', inplace=True)
            self.ecol_info.columns = ['Name', 'b(mm)', 'h(mm)', 'Concrete Grade'
                                      , 'Main Rebar Type', 'Main Rebar(DXX)', 'Hoop Type'
                                      , 'Hoop(DXX)', 'Layer1 EA', 'Layer1 Row', 'Layer2 EA'
                                      , 'Layer2 Row', 'Hoop X', 'Hoop Y', 'Hoop Space(mm)', 'Direction']

        # Nodes
        if (get_BR == True) | (get_BSF == True) | (get_WAS == True) | (get_WR == True) | (get_WSF == True)\
            | (get_E_CSF == True):
            self.node_data = read_excel(to_load_list[0], 'Node Coordinate Data')
            column_name_to_slice = ['Node ID', 'H1', 'H2', 'V']
            self.node_data = self.node_data.loc[:, column_name_to_slice]
        # Elements(Wall)
        if (get_WAS == True) | (get_WR == True) | (get_WSF == True):
            self.wall_data = read_excel(to_load_list[0], 'Element Data - Shear Wall')
            column_name_to_slice = ['Element Name', 'Property Name', 'I-Node ID', 'J-Node ID', 'K-Node ID', 'L-Node ID']
            self.wall_data = self.wall_data.loc[:, column_name_to_slice]
        # Elements(Frame)
        if (get_BR == True) | (get_BSF == True) | (get_E_CSF == True):
            self.frame_data = read_excel(to_load_list[0], 'Element Data - Frame Types')
            column_name_to_slice = ['Element Name', 'Property Name', 'I-Node ID', 'J-Node ID']
            self.frame_data = self.frame_data.loc[:, column_name_to_slice]
        # Wall Axial Strain Gage
        if get_WAS == True:
            self.wall_as_gage_data = read_excel(to_load_list[0], 'Gage Data - Bar Type')
            column_name_to_slice = ['Group Name', 'Element Name', 'I-Node ID', 'J-Node ID']
            self.wall_as_gage_data = self.wall_as_gage_data.loc[:, column_name_to_slice]
        # Wall Rotation Gage
        if (get_WR == True):
            self.wall_rot_gage_data = read_excel(to_load_list[0], 'Gage Data - Wall Type')
            column_name_to_slice = ['Element Name', 'I-Node ID', 'J-Node ID', 'K-Node ID', 'L-Node ID']
            self.wall_rot_gage_data = self.wall_rot_gage_data.loc[:, column_name_to_slice]

        # Using multiprocess (library which overcomes the issue made ny using 'pickle' in 'multiprocessing' library)
        cpu_num = mp.cpu_count() # Count the # of CPU
        pool = mp.Pool(processes=cpu_num) # Create a pool equal to the # of CPU
        # Inter-Story Drift
        self.drift_data = pool.starmap(read_excel, [[file_path, 'Drift Output'] for file_path in to_load_list])
        self.drift_data = pd.concat(self.drift_data, ignore_index=True)
        column_name_to_slice = ['Drift Name', 'Drift ID', 'Load Case', 'Step Type', 'Drift']
        self.drift_data = self.drift_data.loc[:, column_name_to_slice]
        # Wall Shear Force
        if (get_base_SF == True) | (get_story_SF == True) | (get_WR == True) | (get_WSF == True):
            self.shear_force_data = pool.starmap(read_excel, [[file_path, 'Structure Section Forces'] for file_path in to_load_list])
            self.shear_force_data = pd.concat(self.shear_force_data, ignore_index=True)
            column_name_to_slice = ['StrucSec Name', 'Load Case', 'Step Type', 'FH1', 'FH2', 'FV']
            self.shear_force_data = self.shear_force_data.loc[:, column_name_to_slice]
            self.shear_force_data.columns = ['Name', 'Load Case', 'Step Type', 'H1(kN)', 'H2(kN)', 'V(kN)']
        # Wall Axial Strain Result
        if get_WAS == True:
            self.wall_as_result_data = pool.starmap(read_excel, [[file_path, 'Gage Results - Bar Type'] for file_path in to_load_list])
            self.wall_as_result_data = pd.concat(self.wall_as_result_data, ignore_index=True)
            column_name_to_slice = ['Group Name', 'Element Name', 'Load Case', 'Step Type', 'Axial Strain', 'Performance Level']
            self.wall_as_result_data = self.wall_as_result_data.loc[:, column_name_to_slice]
        # Wall Rotation Result
        if (get_WR == True):
            self.wall_rot_result_data= pool.starmap(read_excel, [[file_path, 'Gage Results - Wall Type'] for file_path in to_load_list])
            self.wall_rot_result_data = pd.concat(self.wall_rot_result_data, ignore_index=True)
            column_name_to_slice = ['Group Name', 'Element Name', 'Load Case', 'Step Type', 'Rotation', 'Performance Level']
            self.wall_rot_result_data = self.wall_rot_result_data.loc[:, column_name_to_slice]
        # Beam Rotation
        if get_BR == True:
            self.beam_rot_data = pool.starmap(read_excel, [[file_path, 'Frame Results - Bending Deform'] for file_path in to_load_list])
            self.beam_rot_data = pd.concat(self.beam_rot_data, ignore_index=True)
            column_name_to_slice = ['Group Name', 'Element Name', 'Load Case', 'Step Type', 'Point ID', 'Relative Location', 'R2', 'R3']
            self.beam_rot_data = self.beam_rot_data.loc[:, column_name_to_slice]
            self.beam_rot_data.columns = ['Group Name', 'Element Name', 'Load Case', 'Step Type', 'Point ID', 'Relative Location', 'H2 Rotation(rad)', 'H3 Rotation(rad)']
        # Beam Shear Force
        if (get_BSF == True) | (get_E_CSF == True):
            self.beam_shear_force_data = pool.starmap(read_excel, [[file_path, 'Frame Results - End Forces'] for file_path in to_load_list])
            self.beam_shear_force_data = pd.concat(self.beam_shear_force_data, ignore_index=True)
            column_name_to_slice = ['Group Name', 'Element Name', 'Load Case', 'Step Type'
                                    , 'P I-End', 'V2 I-End', 'V2 J-End', 'V3 I-End', 'M2 I-End', 'M3 I-End']
            self.beam_shear_force_data = self.beam_shear_force_data.loc[:, column_name_to_slice]

        ##### Create Seismic Loads List
        self.load_name_list = []
        for i in self.drift_data['Load Case'].drop_duplicates():
            new_i = i.split('+')[1]
            new_i = new_i.strip()
            self.load_name_list.append(new_i)
        self.gravity_load_name = [x for x in self.load_name_list if ('DE' not in x) and ('MCE' not in x)]
        self.seismic_load_name_list = [x for x in self.load_name_list if ('DE' in x) or ('MCE' in x)]
        self.seismic_load_name_list.sort()        
        self.DE_load_name_list = [x for x in self.load_name_list if 'DE' in x] # base shear로 사용할 지진파 개수 산정을 위함
        self.MCE_load_name_list = [x for x in self.load_name_list if 'MCE' in x]
        self.DE_load_name_list.sort()
        self.MCE_load_name_list.sort() 

        # pkl 폴더 생성
        def create_dir(directory):
            try:
                if not os.path.exists(directory):
                    os.makedirs(directory)
            except OSError:
                print("Error: Failed to create the directory.")                
        create_dir('pkl')        

    ##### Import Post-Processing Methods
    from PBD_p3d.system import base_SF, story_SF, IDR
    from PBD_p3d.wall import WAS, WR, WSF, WAS_plot, WR_plot, WSF_plot
    from PBD_p3d.beam import BR, BSF, BR_plot, BSF_plot
    from PBD_p3d.column import CR, CSF, E_CSF

#%% Function to Print the Result into PDF

def print_pdf(beam_design_xlsx_path, dbeam_design_xlsx_path, col_design_xlsx_path
              , wall_design_xlsx_path, get_cbeam=False, get_dbeam=False, get_ecol=False
              , get_wall=False, project_name='성능기반 내진설계', bldg_name='1동'):

    # Call CoInitialize function before using any COM object
    excel = win32com.client.gencache.EnsureDispatch('Excel.Application', pythoncom.CoInitialize()) # 엑셀 실행
    excel.Visible = False # 엑셀창 안보이게
    
    xlTypePDF = 0
    xlQualityStandard = 0

    element_num = 0
    if get_cbeam == True:
        wb_cbeam = excel.Workbooks.Open(beam_design_xlsx_path)
        ws_DE = wb_cbeam.Sheets('Plot_C.Beam_DE')
        ws_MCE = wb_cbeam.Sheets('Plot_C.Beam_MCE')
        ws_row_num = wb_cbeam.Sheets('Table_C.Beam_DE')
        startrow, startcol = 5, 1

        ### 프로젝트 & 건물명 입력
        ws_DE.Range('BS5:BS5').Value = project_name
        ws_DE.Range('BS6:BS6').Value = bldg_name

        ### 부재별 excel 시트 생성 & pdf 생성
        # 부재 개수(for iterration) 구하기
        element_num = ws_row_num.Range('A4:A4').Value
        element_num = int(element_num)
        
        # Path 지정
        result_path = os.path.splitext(beam_design_xlsx_path)[0] # 확장자명(extension) 제거
        
        # pdf Merge를 위한 PdfMerger 클래스 생성
        merger = PdfMerger()
        
        for i in range(element_num):
            
            ws_DE.Range('A8:A8').Value = i + 1
            
            # 왜인지 모르겠지만 result_path에 suffix 붙이면 \가 /로 바뀜... 그래서 다시 바꿔주기
            pdf_file_path = result_path + '_DE({}).pdf'.format(i + 1)
            pdf_file_path = pdf_file_path.replace('/', '\\')
            
            ws_DE.Select()
            
            wb_cbeam.ActiveSheet.ExportAsFixedFormat(xlTypePDF, pdf_file_path\
                                                , xlQualityStandard, True, False)    

            merger.append(pdf_file_path)
            
        for i in range(element_num):
            
            ws_MCE.Range('A8:A8').Value = i + 1
            
            # 왜인지 모르겠지만 result_path에 suffix 붙이면 \가 /로 바뀜... 그래서 다시 바꿔주기
            pdf_file_path = result_path + '_MCE({}).pdf'.format(i + 1)
            pdf_file_path = pdf_file_path.replace('/', '\\')
            
            ws_MCE.Select()
            
            wb_cbeam.ActiveSheet.ExportAsFixedFormat(xlTypePDF, pdf_file_path\
                                                , xlQualityStandard, True, False)    

            merger.append(pdf_file_path)
            
        merger.write(result_path + '.pdf')
        merger.close()
        
        # Merge한 후 개별 파일들 지우기    
        for i in range(element_num):
            DE_pdf_file_path = result_path + '_DE({}).pdf'.format(i + 1)
            MCE_pdf_file_path = result_path + '_MCE({}).pdf'.format(i + 1)
            DE_pdf_file_path = DE_pdf_file_path.replace('/', '\\')
            MCE_pdf_file_path = MCE_pdf_file_path.replace('/', '\\')
            os.remove(DE_pdf_file_path)
            os.remove(MCE_pdf_file_path)
            
        wb_cbeam.Close(SaveChanges=False)
        
    if get_dbeam == True:
        wb_dbeam = excel.Workbooks.Open(dbeam_design_xlsx_path)
        ws_DE = wb_dbeam.Sheets('Plot_C.Beam_DE')
        ws_MCE = wb_dbeam.Sheets('Plot_C.Beam_MCE')
        ws_row_num = wb_dbeam.Sheets('Table_C.Beam_DE')
        startrow, startcol = 5, 1

        ### 프로젝트 & 건물명 입력
        ws_DE.Range('BS5:BS5').Value = project_name
        ws_DE.Range('BS6:BS6').Value = bldg_name

        ### 부재별 excel 시트 생성 & pdf 생성
        # C.beam의 index에 연결해서 d.beam의 indexing을 할 때 사용하기 위해
        # previous element nuimber를 따로 저장함
        prev_element_num = element_num
        
        # 부재 개수(for iterration) 구하기
        element_num = ws_row_num.Range('A4:A4').Value
        element_num = int(element_num)
        
        # Path 지정
        result_path = os.path.splitext(dbeam_design_xlsx_path)[0] # 확장자명(extension) 제거
        
        # pdf Merge를 위한 PdfMerger 클래스 생성
        merger = PdfMerger()
        
        for i in range(element_num):
            
            ws_DE.Range('A8:A8').Value = i + 1 + prev_element_num
            
            # 왜인지 모르겠지만 result_path에 suffix 붙이면 \가 /로 바뀜... 그래서 다시 바꿔주기
            pdf_file_path = result_path + '_DE({}).pdf'.format(i + 1)
            pdf_file_path = pdf_file_path.replace('/', '\\')
            
            ws_DE.Select()
            
            wb_dbeam.ActiveSheet.ExportAsFixedFormat(xlTypePDF, pdf_file_path\
                                                , xlQualityStandard, True, False)    

            merger.append(pdf_file_path)
            
        for i in range(element_num):
            
            ws_MCE.Range('A8:A8').Value = i + 1 + prev_element_num
            
            # 왜인지 모르겠지만 result_path에 suffix 붙이면 \가 /로 바뀜... 그래서 다시 바꿔주기
            pdf_file_path = result_path + '_MCE({}).pdf'.format(i + 1)
            pdf_file_path = pdf_file_path.replace('/', '\\')
            
            ws_MCE.Select()
            
            wb_dbeam.ActiveSheet.ExportAsFixedFormat(xlTypePDF, pdf_file_path\
                                                , xlQualityStandard, True, False)    

            merger.append(pdf_file_path)
            
        merger.write(result_path + '.pdf')
        merger.close()
        
        # Merge한 후 개별 파일들 지우기    
        for i in range(element_num):
            DE_pdf_file_path = result_path + '_DE({}).pdf'.format(i + 1)
            MCE_pdf_file_path = result_path + '_MCE({}).pdf'.format(i + 1)
            DE_pdf_file_path = DE_pdf_file_path.replace('/', '\\')
            MCE_pdf_file_path = MCE_pdf_file_path.replace('/', '\\')
            os.remove(DE_pdf_file_path)
            os.remove(MCE_pdf_file_path)
            
        wb_dbeam.Close(SaveChanges=False)
            
    if get_ecol == True:
        wb_ecol = excel.Workbooks.Open(col_design_xlsx_path)
        ws = wb_ecol.Sheets('Design_E.Column')        
        startrow, startcol = 5, 1
        
        ### 프로젝트 & 건물명 입력
        ws.Range('B1:B1').Value = project_name
        ws.Range('T1:T1').Value = bldg_name
        
        ### 부재별 excel 시트 생성 & pdf 생성
        # 부재 개수(for iterration) 구하기
        ws_row_num = ws.UsedRange.Rows.Count        
        element_name = ws.Range('B%s:B%s' %(startrow, ws_row_num)).Value
        element_name_df = pd.DataFrame(element_name)

        # Drop NoneType object & Rebar Diameter in the end of Data
        element_name_df.iloc[:,0] = element_name_df[element_name_df.iloc[:,0].str.count('_') == 2]
        element_name_df = element_name_df.dropna()
        
        # 부재개수
        element_num = element_name_df.shape[0]
        
        # Path 지정
        result_path = os.path.splitext(col_design_xlsx_path)[0] # 확장자명(extension) 제거
        
        # pdf Merge를 위한 PdfMerger 클래스 생성
        merger = PdfMerger()

        for i in range(element_num):

            wb_ecol.Worksheets(3).Select()            
            wb_ecol.Worksheets(3).Name = '({})'.format(i+1)

            # wb_ecol.SaveAs(pdf_file_path, FileFormat=57)
            
            # 왜인지 모르겠지만 result_path에 suffix 붙이면 \가 /로 바뀜... 그래서 다시 바꿔주기
            pdf_file_path = result_path + '({}).pdf'.format(i+1)
            pdf_file_path = pdf_file_path.replace('/', '\\')            
            wb_ecol.ActiveSheet.ExportAsFixedFormat(xlTypePDF, pdf_file_path\
                                                , xlQualityStandard, True, False)    

            merger.append(pdf_file_path)
            
        merger.write(result_path + '.pdf')
        merger.close()
        
        # Merge한 후 개별 파일들 지우기    
        for i in range(element_num):
            pdf_file_path = result_path + '({}).pdf'.format(i+1)
            pdf_file_path = pdf_file_path.replace('/', '\\')
            os.remove(pdf_file_path)

        wb_ecol.Close(SaveChanges=False)
        
    if get_wall == True:
        wb_wall = excel.Workbooks.Open(wall_design_xlsx_path)
        ws_DE = wb_wall.Sheets('Plot_S.Wall_DE')
        ws_MCE = wb_wall.Sheets('Plot_S.Wall_MCE')
        ws_row_num = wb_wall.Sheets('Table_S.Wall_DE')
        startrow, startcol = 5, 1

        ### 프로젝트 & 건물명 입력
        ws_DE.Range('BS5:BS5').Value = project_name
        ws_DE.Range('BS6:BS6').Value = bldg_name

        ### 부재별 excel 시트 생성 & pdf 생성
        # 부재 개수(for iterration) 구하기
        element_num = ws_row_num.Range('A4:A4').Value
        element_num = int(element_num)
        
        # Path 지정
        result_path = os.path.splitext(wall_design_xlsx_path)[0] # 확장자명(extension) 제거
        
        # pdf Merge를 위한 PdfMerger 클래스 생성
        merger = PdfMerger()

        for i in range(element_num):
            
            ws_DE.Range('A8:A8').Value = i + 1
            
            # 왜인지 모르겠지만 result_path에 suffix 붙이면 \가 /로 바뀜... 그래서 다시 바꿔주기
            pdf_file_path = result_path + '_DE({}).pdf'.format(i + 1)
            pdf_file_path = pdf_file_path.replace('/', '\\')
            
            ws_DE.Select()
            
            wb_wall.ActiveSheet.ExportAsFixedFormat(xlTypePDF, pdf_file_path\
                                                , xlQualityStandard, True, False)    

            merger.append(pdf_file_path)
            
        for i in range(element_num):
            
            ws_MCE.Range('A8:A8').Value = i + 1
            
            # 왜인지 모르겠지만 result_path에 suffix 붙이면 \가 /로 바뀜... 그래서 다시 바꿔주기
            pdf_file_path = result_path + '_MCE({}).pdf'.format(i + 1)
            pdf_file_path = pdf_file_path.replace('/', '\\')
            
            ws_MCE.Select()
            
            wb_wall.ActiveSheet.ExportAsFixedFormat(xlTypePDF, pdf_file_path\
                                                , xlQualityStandard, True, False)    

            merger.append(pdf_file_path)
            
        merger.write(result_path + '.pdf')
        merger.close()
        
        # Merge한 후 개별 파일들 지우기    
        for i in range(element_num):
            DE_pdf_file_path = result_path + '_DE({}).pdf'.format(i + 1)
            MCE_pdf_file_path = result_path + '_MCE({}).pdf'.format(i + 1)
            DE_pdf_file_path = DE_pdf_file_path.replace('/', '\\')
            MCE_pdf_file_path = MCE_pdf_file_path.replace('/', '\\')
            os.remove(DE_pdf_file_path)
            os.remove(MCE_pdf_file_path)
            
        wb_wall.Close(SaveChanges=False)

#%% Function to Print the Result into DOCX

def print_docx(result_xlsx_path, get_base_SF=False, get_story_SF=False
               , get_IDR=False, get_BR=False, get_BSF=False, get_E_BSF=False
               , get_CR=False, get_CSF=False, get_E_CSF=False, get_WAS=False
               , get_WR=False, get_WSF=False, project_name='성능기반 내진설계'
               , bldg_name='1동', story_gap=2, max_shear=60000):  

    # Other Parameters (향 후, UI에서 조작할 수 있게끔)
    cri_DE=0.015 # IDR
    cri_MCE=0.02 # IDR
    max_criteria=0.04 # WAS
    min_criteria=-0.002 # WAS
    DCR_criteria=1
    xlim = 2 # BR
    WAS_gage_group='AS' # WAS
    
    fig_scale = 3/4 # 그래프 크기, 축, 글씨 등 scale up/down. (scale과 반비례함)
    
    # Create docx Class (with Template)
    document = docx.Document("template/report_template.docx")
    
    ### 동이름 입력
    # 동 이름 replace(paragraph level)
    for paragraph in document.paragraphs:
        for run in paragraph.runs:        
            if 'building_name' in run.text:
                run.text = bldg_name
                
    # 동 이름 replace(table level)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if 'building_name' in run.text:
                            run.text = bldg_name    
    
    ### System
    #%% Base Shear
    if get_base_SF == True:
        with open('pkl/base_SF.pkl', 'rb') as f:
            base_SF_result = pickle.load(f)
        
        # 결과값 classify & assign
        base_shear_H1 = base_SF_result[0]
        base_shear_H2 = base_SF_result[1]
        DE_load_name_list = base_SF_result[2]
        MCE_load_name_list = base_SF_result[3]
        
        # Base Shear 표 작성
        # template의 0,1번 표 불러오기
        base_SF_values_table = document.tables[0]
        base_SF_plots_table = document.tables[1]
        
        # Plot
        # DE Plot
        if len(DE_load_name_list) != 0:
            # H1_DE
            fig1, ax1 = plt.subplots(1,1, figsize=(7*fig_scale, 6.5*fig_scale), dpi=200)
            fig1.tight_layout() # 이거 안하면 크기 맘대로 바뀜 ㅠ
            ax1.set_ylim(0, max_shear)
        
            ax1.bar(range(len(DE_load_name_list)), base_shear_H1.iloc[0, 0:len(DE_load_name_list)]\
                    , color='darkblue', edgecolor='k', label = 'Max. Base Shear')
            ax1.axhline(y= base_shear_H1.iloc[0, 0:len(DE_load_name_list)].mean(), color='r', linestyle='-', label='Average')
            ax1.set_xticks(range(14), range(1,15))
            
            ax1.set_xlabel('Ground Motion No.')
            ax1.set_ylabel('Base Shear(kN)')
            ax1.legend(loc = 2)
            # ax1.set_title('X 1.2$\star$DBE')
            
            base_SF_avg_DE_x = Decimal(str(base_shear_H1.iloc[0, 0:len(DE_load_name_list)].mean()))\
                .quantize(Decimal('.01'), rounding=ROUND_UP)        
            
            memfile = BytesIO()
            plt.savefig(memfile, bbox_inches="tight") # bbox_inches=tight : 이미지 크기에 그래프가 맞도록
            plt.close()
            
            # H2_DE
            fig2, ax2 = plt.subplots(1,1, figsize=(7*fig_scale, 6.5*fig_scale), dpi=200)
            fig2.tight_layout()
            ax2.set_ylim(0, max_shear)
            
            ax2.bar(range(len(DE_load_name_list)), base_shear_H2.iloc[0, 0:len(DE_load_name_list)], color='darkblue', edgecolor='k', label = 'Max. Base Shear')
            ax2.axhline(y= base_shear_H2.iloc[0, 0:len(DE_load_name_list)].mean(), color='r', linestyle='-', label='Average')
            ax2.set_xticks(range(14), range(1,15))
            
            ax2.set_xlabel('Ground Motion No.')
            ax2.set_ylabel('Base Shear(kN)')
            ax2.legend(loc = 2)
            # ax2.set_title('Y 1.2$\star$DBE')
            
            base_SF_avg_DE_y = Decimal(str(base_shear_H2.iloc[0, 0:len(DE_load_name_list)].mean()))\
                .quantize(Decimal('.01'), rounding=ROUND_UP)
            
            memfile2 = BytesIO()
            plt.savefig(memfile2, bbox_inches="tight")
            plt.close()

            # 첫번째 표에 avg 값 넣기
            values_row = base_SF_values_table.rows[4]
            values_cell_x = values_row.cells[2]
            values_cell_y = values_row.cells[3]
            values_para_x = values_cell_x.paragraphs[0]
            values_para_y = values_cell_y.paragraphs[0]
            values_para_x.text = f'{base_SF_avg_DE_x:,} kN' # 1000 자리마다 , 찍기
            values_para_y.text = f'{base_SF_avg_DE_y:,} kN'
            values_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
            values_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 두번째 표에 그래프 넣기            
            plots_row = base_SF_plots_table.rows[0]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(7), height=Cm(6.5))
            plots_run_y.add_picture(memfile2, width=Cm(7), height=Cm(6.5))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # MCE Plot
        if len(MCE_load_name_list) != 0:
            # H1_MCE
            fig3, ax3 = plt.subplots(1,1, figsize=(7*fig_scale, 6.5*fig_scale), dpi=200)
            fig3.tight_layout()
            ax3.set_ylim(0, max_shear)
            
            ax3.bar(range(len(MCE_load_name_list)), base_shear_H1\
                    .iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)]\
                    , color='darkblue', edgecolor='k', label = 'Max. Base Shear')
            ax3.axhline(y= base_shear_H1.iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)]\
                        .mean(), color='r', linestyle='-', label='Average')
            ax3.set_xticks(range(14), range(1,15))
            
            ax3.set_xlabel('Ground Motion No.')
            ax3.set_ylabel('Base Shear(kN)')
            ax3.legend(loc = 2)
            # ax3.set_title('X MCE')
            
            base_SF_avg_MCE_x = Decimal(str(base_shear_H1.iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)].mean()))\
                .quantize(Decimal('.01'), rounding=ROUND_UP)
            
            memfile = BytesIO()
            plt.savefig(memfile, bbox_inches="tight")           
            plt.close()            

            # H2_MCE
            fig4, ax4 = plt.subplots(1,1, figsize=(7*fig_scale, 6.5*fig_scale), dpi=200)
            fig4.tight_layout()
            ax4.set_ylim(0, max_shear)
            
            plt.bar(range(len(MCE_load_name_list)), base_shear_H2\
                    .iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)]\
                    , color='darkblue', edgecolor='k', label = 'Max. Base Shear')
            plt.axhline(y= base_shear_H2.iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)]\
                        .mean(), color='r', linestyle='-', label='Average')
            ax4.set_xticks(range(14), range(1,15))
            
            ax4.set_xlabel('Ground Motion No.')
            ax4.set_ylabel('Base Shear(kN)')
            ax4.legend(loc = 2)
            # ax4.set_title('Y MCE')
            
            base_SF_avg_MCE_y = Decimal(str(base_shear_H2.iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)].mean()))\
                .quantize(Decimal('.01'), rounding=ROUND_UP)

            memfile2 = BytesIO()
            plt.savefig(memfile2, bbox_inches="tight")
            plt.close()            

            # 첫번째 표에 avg 값 넣기
            values_row = base_SF_values_table.rows[5]
            values_cell_x = values_row.cells[2]
            values_cell_y = values_row.cells[3]
            values_para_x = values_cell_x.paragraphs[0]
            values_para_y = values_cell_y.paragraphs[0]
            values_para_x.text = f'{base_SF_avg_MCE_x:,} kN' # 1000 자리마다 , 찍기
            values_para_y.text = f'{base_SF_avg_MCE_y:,} kN'
            values_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
            values_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 두번째 표에 그래프 넣기            
            plots_row = base_SF_plots_table.rows[3]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(7), height=Cm(6.5))
            plots_run_y.add_picture(memfile2, width=Cm(7), height=Cm(6.5))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER               
            
    #%% Story Drift
    if get_story_SF == True:
        with open('pkl/story_SF.pkl', 'rb') as f:
            story_SF_result = pickle.load(f)
        
        # 결과값 classify & assign
        shear_force_H1_max = story_SF_result[0]
        shear_force_H2_max = story_SF_result[1]
        DE_load_name_list = story_SF_result[2]
        MCE_load_name_list = story_SF_result[3]

        # Story Shear Force 표 작성
        # template의 3번 표 불러오기
        story_SF_plots_table = document.tables[3]

        # Plot
        # DE Plot
        if len(DE_load_name_list) != 0:
            # H1_DE
            fig5, ax5 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig5.tight_layout()
            
            # 지진파별 plot
            for i in range(len(DE_load_name_list)):
                ax5.plot(shear_force_H1_max.iloc[:,i], range(shear_force_H1_max.shape[0]), label=DE_load_name_list[i], linewidth=0.7)
                
            # 평균 plot
            ax5.plot(shear_force_H1_max.iloc[:,0:len(DE_load_name_list)]\
                    .mean(axis=1), range(shear_force_H1_max.shape[0]), color='k', label='Average', linewidth=2)
            
            ax5.set_xlim(0, max_shear)
            ax5.set_yticks(range(shear_force_H1_max.shape[0])[::story_gap], shear_force_H1_max.index[::story_gap], fontsize=8.5)
            
            # 기타
            ax5.grid(linestyle='-.')
            ax5.set_xlabel('Story Shear(kN)')
            ax5.set_ylabel('Story')
            ax5.legend(loc=1, fontsize=8)
            # ax5.set_title('X 1.2$\star$DBE')
            
            memfile = BytesIO()
            plt.savefig(memfile, bbox_inches="tight")           
            plt.close()
            
            # H2_DE
            fig6, ax6 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig6.tight_layout()
            
            for i in range(len(DE_load_name_list)):
                ax6.plot(shear_force_H2_max.iloc[:,i], range(shear_force_H2_max.shape[0]), label=DE_load_name_list[i], linewidth=0.7)
            
            ax6.plot(shear_force_H2_max.iloc[:,0:len(DE_load_name_list)]\
                    .mean(axis=1), range(shear_force_H2_max.shape[0]), color='k', label='Average', linewidth=2)
            
            ax6.set_xlim(0, max_shear)
            ax6.set_yticks(range(shear_force_H2_max.shape[0])[::story_gap], shear_force_H2_max.index[::story_gap], fontsize=8.5)
        
            ax6.grid(linestyle='-.')
            ax6.set_xlabel('Story Shear(kN)')
            ax6.set_ylabel('Story')
            ax6.legend(loc=1, fontsize=8)
            # ax6.set_title('Y 1.2$\star$DBE')
            
            memfile2 = BytesIO()
            plt.savefig(memfile2, bbox_inches="tight")
            plt.close()
            
            # 표에 그래프 넣기            
            plots_row = story_SF_plots_table.rows[0]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(7), height=Cm(9.5))
            plots_run_y.add_picture(memfile2, width=Cm(7), height=Cm(9.5))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # MCE Plot
        if len(MCE_load_name_list) != 0:
            # H1_MCE
            fig7, ax7 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig7.tight_layout()
            
            for i in range(len(MCE_load_name_list)):
                ax7.plot(shear_force_H1_max.iloc[:,i+len(DE_load_name_list)], range(shear_force_H1_max.shape[0]), label=MCE_load_name_list[i], linewidth=0.7)
            ax7.plot(shear_force_H1_max.iloc[:,len(DE_load_name_list)\
                                                    :len(DE_load_name_list)+len(MCE_load_name_list)]\
                            .mean(axis=1), range(shear_force_H1_max.shape[0]), color='k', label='Average', linewidth=2)
            
            ax7.set_xlim(0, max_shear)
            ax7.set_yticks(range(shear_force_H1_max.shape[0])[::story_gap], shear_force_H1_max.index[::story_gap], fontsize=8.5)
        
            ax7.grid(linestyle='-.')
            ax7.set_xlabel('Story Shear(kN)')
            ax7.set_ylabel('Story')
            ax7.legend(loc=1, fontsize=8)
            # ax7.set_title('X MCE')
        
            memfile = BytesIO()
            plt.savefig(memfile, bbox_inches="tight")           
            plt.close()
            
            # H1_MCE
            fig8, ax8 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig8.tight_layout()
            
            for i in range(len(MCE_load_name_list)):
                ax8.plot(shear_force_H2_max.iloc[:,i+len(DE_load_name_list)], range(shear_force_H2_max.shape[0]), label=MCE_load_name_list[i], linewidth=0.7)
            ax8.plot(shear_force_H2_max.iloc[:,len(DE_load_name_list)\
                                                    :len(DE_load_name_list)+len(MCE_load_name_list)]\
                            .mean(axis=1), range(shear_force_H2_max.shape[0]), color='k', label='Average', linewidth=2)
            
            ax8.set_xlim(0, max_shear)
            ax8.set_yticks(range(shear_force_H2_max.shape[0])[::story_gap], shear_force_H2_max.index[::story_gap], fontsize=8.5)
        
            ax8.grid(linestyle='-.')
            ax8.set_xlabel('Story Shear(kN)')
            ax8.set_ylabel('Story')
            ax8.legend(loc=1, fontsize=8)
            # ax8.set_title('Y MCE')
            
            memfile2 = BytesIO()
            plt.savefig(memfile2, bbox_inches="tight")
            plt.close()
            
            # 표에 그래프 넣기            
            plots_row = story_SF_plots_table.rows[3]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(7), height=Cm(9.5))
            plots_run_y.add_picture(memfile2, width=Cm(7), height=Cm(9.5))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    #%% Inter-Story Drift
    if get_IDR == True:
        with open('pkl/IDR.pkl', 'rb') as f:
            IDR_result = pickle.load(f)
        
        # 결과값 classify & assign
        IDR_result_each = IDR_result[0]
        IDR_result_avg = IDR_result[1]
        DE_load_name_list = IDR_result[2]
        MCE_load_name_list = IDR_result[3]
        story_name_window_reordered = IDR_result[4]
        
        # Inter-Story Drift 표 작성
        # template의 2번 표 불러오기
        IDR_plots_table = document.tables[2]
        
        # Plot
        count_x = 0
        count_y = 2
        count_avg = 0
        
        # DE Plot
        if len(DE_load_name_list) != 0:
            # H1_DE   
            fig9, ax9 = plt.subplots(1,1, figsize=(7*fig_scale, 10*fig_scale), dpi=200)
            fig9.tight_layout()
            
            # 지진파별 plot
            for load_name in DE_load_name_list:
                ax9.plot(IDR_result_each[count_x].iloc[:,-1], IDR_result_each[count_x].iloc[:,0]
                         , label='{}'.format(load_name), linewidth=0.7)
                ax9.plot(IDR_result_each[count_x+1].iloc[:,-1], IDR_result_each[count_x].iloc[:,0]
                         , linewidth=0.7)
                count_x += 4
                
            # 평균 plot
            ax9.plot(IDR_result_avg[count_avg].iloc[:,0], story_name_window_reordered, color='k', label='Average', linewidth=2)
            ax9.plot(IDR_result_avg[count_avg].iloc[:,1], story_name_window_reordered, color='k', linewidth=2)
            
            # reference line 그려서 허용치 나타내기
            ax9.axvline(x=-cri_DE, color='r', linestyle='--', label='LS')
            ax9.axvline(x=cri_DE, color='r', linestyle='--')
            
            ax9.set_xlim(-0.025, 0.025)
            ax9.set_yticks(story_name_window_reordered[::story_gap], story_name_window_reordered[::story_gap])
            
            # 기타
            ax9.grid(linestyle='-.')
            ax9.set_xlabel('Interstory Drift Ratios(m/m)')
            ax9.set_ylabel('Story')
            ax9.legend(loc=4, fontsize=8)
            # ax9.set_title('X 1.2$\star$DBE')
            
            memfile = BytesIO()
            plt.savefig(memfile, bbox_inches="tight")           
            plt.close()
            
            # H2_DE
            fig10, ax10 = plt.subplots(1,1, figsize=(7*fig_scale, 10*fig_scale), dpi=200)
            fig10.tight_layout()

            # 지진파별 plot
            for load_name in DE_load_name_list:
                ax10.plot(IDR_result_each[count_y].iloc[:,-1], IDR_result_each[count_y].iloc[:,0]
                         , label='{}'.format(load_name), linewidth=0.7)
                ax10.plot(IDR_result_each[count_y+1].iloc[:,-1], IDR_result_each[count_y].iloc[:,0]
                         , linewidth=0.7)
                count_y += 4
                
            # 평균 plot
            ax10.plot(IDR_result_avg[count_avg].iloc[:,0], story_name_window_reordered, color='k', label='Average', linewidth=2)
            ax10.plot(IDR_result_avg[count_avg].iloc[:,1], story_name_window_reordered, color='k', linewidth=2)
            count_avg += 1
            
            # reference line 그려서 허용치 나타내기
            ax10.axvline(x=-cri_DE, color='r', linestyle='--', label='LS')
            ax10.axvline(x=cri_DE, color='r', linestyle='--')
            
            ax10.set_xlim(-0.025, 0.025)
            ax10.set_yticks(story_name_window_reordered[::story_gap], story_name_window_reordered[::story_gap])
            
            # 기타
            ax10.grid(linestyle='-.')
            ax10.set_xlabel('Interstory Drift Ratios(m/m)')
            ax10.set_ylabel('Story')
            ax10.legend(loc=4, fontsize=8)
            # ax10.set_title('Y 1.2$\star$DBE')
            
            memfile2 = BytesIO()
            plt.savefig(memfile2, bbox_inches="tight")
            plt.close()

            # 표에 그래프 넣기            
            plots_row = IDR_plots_table.rows[0]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(7), height=Cm(10))
            plots_run_y.add_picture(memfile2, width=Cm(7), height=Cm(10))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # MCE Plot
        if len(MCE_load_name_list) != 0:
            # H1_MCE
            fig11, ax11 = plt.subplots(1,1, figsize=(7*fig_scale, 10*fig_scale), dpi=200)
            fig11.tight_layout()
            
            # 지진파별 plot
            for load_name in MCE_load_name_list:
                ax11.plot(IDR_result_each[count_x].iloc[:,-1], IDR_result_each[count_x].iloc[:,0]
                         , label='{}'.format(load_name), linewidth=0.7)
                ax11.plot(IDR_result_each[count_x+1].iloc[:,-1], IDR_result_each[count_x].iloc[:,0]
                         , linewidth=0.7)
                count_x += 4
                
            # 평균 plot
            ax11.plot(IDR_result_avg[count_avg].iloc[:,0], story_name_window_reordered, color='k', label='Average', linewidth=2)
            ax11.plot(IDR_result_avg[count_avg].iloc[:,1], story_name_window_reordered, color='k', linewidth=2)
            
            # reference line 그려서 허용치 나타내기
            ax11.axvline(x=-cri_MCE, color='r', linestyle='--', label='CP')
            ax11.axvline(x=cri_MCE, color='r', linestyle='--')
            
            ax11.set_xlim(-0.025, 0.025)
            ax11.set_yticks(story_name_window_reordered[::story_gap], story_name_window_reordered[::story_gap])
            
            # 기타
            ax11.grid(linestyle='-.')
            ax11.set_xlabel('Interstory Drift Ratios(m/m)')
            ax11.set_ylabel('Story')
            ax11.legend(loc=4, fontsize=8)
            # ax11.set_title('X MCE')
            
            memfile = BytesIO()
            plt.savefig(memfile, bbox_inches="tight")           
            plt.close()

            # H2_MCE
            fig12, ax12 = plt.subplots(1,1, figsize=(7*fig_scale, 10*fig_scale), dpi=200)
            fig12.tight_layout()

            # 지진파별 plot
            for load_name in MCE_load_name_list:
                ax12.plot(IDR_result_each[count_y].iloc[:,-1], IDR_result_each[count_y].iloc[:,0]
                         , label='{}'.format(load_name), linewidth=0.7)
                ax12.plot(IDR_result_each[count_y+1].iloc[:,-1], IDR_result_each[count_y].iloc[:,0]
                         , linewidth=0.7)
                count_y += 4
                
            # 평균 plot
            ax12.plot(IDR_result_avg[count_avg].iloc[:,0], story_name_window_reordered, color='k', label='Average', linewidth=2)
            ax12.plot(IDR_result_avg[count_avg].iloc[:,1], story_name_window_reordered, color='k', linewidth=2)
            count_avg += 1
            
            # reference line 그려서 허용치 나타내기
            ax12.axvline(x=-cri_MCE, color='r', linestyle='--', label='CP')
            ax12.axvline(x=cri_MCE, color='r', linestyle='--')
            
            ax12.set_xlim(-0.025, 0.025)
            ax12.set_yticks(story_name_window_reordered[::story_gap], story_name_window_reordered[::story_gap])
            
            # 기타
            ax12.grid(linestyle='-.')
            ax12.set_xlabel('Interstory Drift Ratios(m/m)')
            ax12.set_ylabel('Story')
            ax12.legend(loc=4, fontsize=8)
            # ax12.set_title('Y MCE')

            memfile2 = BytesIO()
            plt.savefig(memfile2, bbox_inches="tight")
            plt.close()            
            
            # 표에 그래프 넣기            
            plots_row = IDR_plots_table.rows[3]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(7), height=Cm(10))
            plots_run_y.add_picture(memfile2, width=Cm(7), height=Cm(10))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER       

    ### Beam
    #%% Beam Rotation
    if get_BR == True:
        # Load Pickle Files
        with open('pkl/BR.pkl', 'rb') as f:
            BR_result = pickle.load(f)
        
        # 결과값 classify & assign
        BR_plot = BR_result[0]
        story_info = BR_result[1]
        DE_load_name_list = BR_result[2]
        MCE_load_name_list = BR_result[3]
            
        # Beam Rotation 표 작성
        # template의 7번 표 불러오기
        BR_plots_table = document.tables[7]
        
        # Plot
        # DE Plot
        if len(DE_load_name_list) != 0:

            fig13, ax13 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig13.tight_layout()
            
            # DCR plot                
            ax13.scatter(BR_plot['DCR(DE_pos)'], BR_plot['Height(mm)'], color='k', s=1)
            ax13.scatter(BR_plot['DCR(DE_neg)'], BR_plot['Height(mm)'], color='k', s=1)

            # 허용치(DCR) 기준선
            ax13.axvline(x = DCR_criteria, color='r', linestyle='--')
            ax13.axvline(x = -DCR_criteria, color='r', linestyle='--')

            ax13.set_xlim(-xlim, xlim)
            ax13.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax13.grid(linestyle='-.')
            ax13.set_xlabel('DCR')
            ax13.set_ylabel('Story')
            # ax13.set_title('Beam Rotation (1.2$\star$DBE)')
            
            # plt.tight_layout()
            memfile = BytesIO()
            plt.savefig(memfile, bbox_inches="tight")
            plt.close()
            
            # 첫번째 표에 그래프 넣기            
            plots_row = BR_plots_table.rows[0]
            plots_cell = plots_row.cells[0]
            plots_para = plots_cell.paragraphs[0]
            plots_run = plots_para.add_run()
            plots_run.add_picture(memfile, width=Cm(7), height=Cm(9.5))
            plots_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        if len(MCE_load_name_list) != 0:
            
            fig14, ax14 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig14.tight_layout()

            # DCR plot                
            ax14.scatter(BR_plot['DCR(MCE_pos)'], BR_plot['Height(mm)'], color='k', s=1)
            ax14.scatter(BR_plot['DCR(MCE_neg)'], BR_plot['Height(mm)'], color='k', s=1)

            # 허용치(DCR) 기준선
            ax14.axvline(x = DCR_criteria, color='r', linestyle='--')
            ax14.axvline(x = -DCR_criteria, color='r', linestyle='--')

            ax14.set_xlim(-xlim, xlim)
            ax14.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax14.grid(linestyle='-.')
            ax14.set_xlabel('DCR')
            ax14.set_ylabel('Story')
            # ax14.set_title('Beam Rotation (MCE)')
            
            # plt.tight_layout()   
            memfile = BytesIO()
            plt.savefig(memfile, bbox_inches="tight")
            plt.close()

            # 첫번째 표에 그래프 넣기            
            plots_row = BR_plots_table.rows[0]
            plots_cell = plots_row.cells[1]
            plots_para = plots_cell.paragraphs[0]
            plots_run = plots_para.add_run()
            plots_run.add_picture(memfile, width=Cm(7), height=Cm(9.5))
            plots_para.alignment = WD_ALIGN_PARAGRAPH.CENTER   
            
    #%% Beam Shear Force
    if get_BSF == True:
        # Load Pickle Files
        with open('pkl/BSF.pkl', 'rb') as f:
            BSF_result = pickle.load(f)
        
        # 결과값 classify & assign
        BSF_plot = BSF_result[0]
        story_info = BSF_result[1]
        DE_load_name_list = BSF_result[2]
        MCE_load_name_list = BSF_result[3]
            
        # Beam Shear Force 표 작성
        # template의 8번 표 불러오기
        BSF_plots_table = document.tables[8]
        
        # Plot
        # DE Plot
        if len(DE_load_name_list) != 0:
            fig15, ax15 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig15.tight_layout()
            
            # DCR plot                
            ax15.axes.scatter(BSF_plot['DE'], BSF_plot['Height(mm)'], color='k', s=1)
    
            # 허용치(DCR) 기준선
            ax15.axes.axvline(x = DCR_criteria, color='r', linestyle='--')
    
            ax15.axes.set_xlim(0, xlim)
            ax15.axes.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])
    
            # 기타
            ax15.axes.grid(linestyle='-.')
            ax15.axes.set_xlabel('DCR')
            ax15.axes.set_ylabel('Story')
            # ax15.axes.set_title('Shear Strength (1.2$\star$DBE)')
            
            memfile = BytesIO()
            plt.savefig(memfile, bbox_inches="tight")
            plt.close()
            
            # 첫번째 표에 그래프 넣기            
            plots_row = BSF_plots_table.rows[0]
            plots_cell = plots_row.cells[0]
            plots_para = plots_cell.paragraphs[0]
            plots_run = plots_para.add_run()
            plots_run.add_picture(memfile, width=Cm(7), height=Cm(9.5))
            plots_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Plot
        # MCE Plot
        if len(MCE_load_name_list) != 0:
            fig16, ax16 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig16.tight_layout()
            
            # DCR plot                
            ax16.axes.scatter(BSF_plot['MCE'], BSF_plot['Height(mm)'], color='k', s=1)
    
            # 허용치(DCR) 기준선
            ax16.axes.axvline(x = DCR_criteria, color='r', linestyle='--')
    
            ax16.axes.set_xlim(0, xlim)
            ax16.axes.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])
    
            # 기타
            ax16.axes.grid(linestyle='-.')
            ax16.axes.set_xlabel('DCR')
            ax16.axes.set_ylabel('Story')
            # ax16.axes.set_title('Shear Strength (MCE)')
            
            memfile = BytesIO()
            plt.savefig(memfile, bbox_inches="tight")
            plt.close()
            
            # 첫번째 표에 그래프 넣기            
            plots_row = BSF_plots_table.rows[0]
            plots_cell = plots_row.cells[1]
            plots_para = plots_cell.paragraphs[0]
            plots_run = plots_para.add_run()
            plots_run.add_picture(memfile, width=Cm(7), height=Cm(9.5))
            plots_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    #%% Wall Axial Strain
    if get_WAS == True:
        with open('pkl/WAS.pkl', 'rb') as f:
            WAS_result = pickle.load(f)
        
        # 결과값 classify & assign
        WAS_plot = WAS_result[0]
        story_info = WAS_result[1]
        DE_load_name_list = WAS_result[2]
        MCE_load_name_list = WAS_result[3]
        
        # Wall Axial Strain 표 작성
        # template의 2번 표 불러오기
        WAS_plots_table = document.tables[4]            
        
        # DE Plot
        if len(DE_load_name_list) != 0:
            # DE_Neg
            fig17, ax17 = plt.subplots(1,1, figsize=(7*fig_scale, 10*fig_scale), dpi=200)
            fig17.tight_layout()
            
            # WAS plot
            ax17.scatter(WAS_plot['DE(Compressive)'], WAS_plot['Height(mm)'], color='k', s=5)
            ax17.scatter(WAS_plot['DE(Tensile)'], WAS_plot['Height(mm)'], color='k', s=5)

            # 허용치 기준선
            ax17.axvline(x=min_criteria, color='r', linestyle='--')
            ax17.axvline(x=max_criteria, color='r', linestyle='--')

            ax17.set_xlim(-0.003, 0)
            ax17.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax17.grid(linestyle='-.')
            ax17.set_xlabel('Axial Strain (m/m)')
            ax17.set_ylabel('Story')
            # ax17.set_title('1.2$\star$DBE (Compressive)')
            
            memfile = BytesIO()
            plt.savefig(memfile, bbox_inches="tight")           
            plt.close()
            
            # DE_Pos
            fig18, ax18 = plt.subplots(1,1, figsize=(7*fig_scale, 10*fig_scale), dpi=200)
            fig18.tight_layout()
            
            # WAS plot
            ax18.scatter(WAS_plot['DE(Compressive)'], WAS_plot['Height(mm)'], color='k', s=5)
            ax18.scatter(WAS_plot['DE(Tensile)'], WAS_plot['Height(mm)'], color='k', s=5)

            # 허용치 기준선
            ax18.axvline(x=min_criteria, color='r', linestyle='--')
            ax18.axvline(x=max_criteria, color='r', linestyle='--')

            ax18.set_xlim(0, 0.013)
            ax18.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax18.grid(linestyle='-.')
            ax18.set_xlabel('Axial Strain (m/m)')
            ax18.set_ylabel('Story')
            # ax18.set_title('1.2$\star$DBE (Tensile)')
            
            memfile2 = BytesIO()
            plt.savefig(memfile2, bbox_inches="tight")
            plt.close()          
            
            # 표에 그래프 넣기            
            plots_row = WAS_plots_table.rows[0]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(7), height=Cm(10))
            plots_run_y.add_picture(memfile2, width=Cm(7), height=Cm(10))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER   
            
        # MCE Plot
        if len(MCE_load_name_list) != 0:
            # MCE_Neg
            fig19, ax19 = plt.subplots(1,1, figsize=(7*fig_scale, 10*fig_scale), dpi=200)
            fig19.tight_layout()
            
            # WAS plot
            ax19.scatter(WAS_plot['MCE(Compressive)'], WAS_plot['Height(mm)'], color='k', s=5)
            ax19.scatter(WAS_plot['MCE(Tensile)'], WAS_plot['Height(mm)'], color='k', s=5)

            # 허용치 기준선
            ax19.axvline(x=min_criteria, color='r', linestyle='--')
            ax19.axvline(x=max_criteria, color='r', linestyle='--')

            ax19.set_xlim(-0.003, 0)
            ax19.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax19.grid(linestyle='-.')
            ax19.set_xlabel('Axial Strain (m/m)')
            ax19.set_ylabel('Story')
            # ax19.set_title('MCE (Compressive)')
            
            memfile = BytesIO()
            plt.savefig(memfile, bbox_inches="tight")           
            plt.close()
            
            # MCE_Pos
            fig20, ax20 = plt.subplots(1,1, figsize=(7*fig_scale, 10*fig_scale), dpi=200)
            fig20.tight_layout()
            
            # WAS plot
            ax20.scatter(WAS_plot['MCE(Compressive)'], WAS_plot['Height(mm)'], color='k', s=5)
            ax20.scatter(WAS_plot['MCE(Tensile)'], WAS_plot['Height(mm)'], color='k', s=5)

            # 허용치 기준선
            ax20.axvline(x=min_criteria, color='r', linestyle='--')
            ax20.axvline(x=max_criteria, color='r', linestyle='--')

            ax20.set_xlim(0, 0.013)
            ax20.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax20.grid(linestyle='-.')
            ax20.set_xlabel('Axial Strain (m/m)')
            ax20.set_ylabel('Story')
            # ax20.set_title('MCE (Tensile)')
            
            memfile2 = BytesIO()
            plt.savefig(memfile2, bbox_inches="tight")
            plt.close()      
            
            # 표에 그래프 넣기            
            plots_row = WAS_plots_table.rows[3]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(7), height=Cm(10))
            plots_run_y.add_picture(memfile2, width=Cm(7), height=Cm(10))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER       
            
    #%% Wall Rotation
    if get_WR == True:
        with open('pkl/WR.pkl', 'rb') as f:
            WR_result = pickle.load(f)
        
        # 결과값 classify & assign
        WR_plot = WR_result[0]
        story_info = WR_result[1]
        DE_load_name_list = WR_result[2]
        MCE_load_name_list = WR_result[3]
        
        # Wall Rotation 표 작성
        # template의 2번 표 불러오기
        WR_plots_table = document.tables[5]            
        
        # DE Plot
        if len(DE_load_name_list) != 0:
            
            fig21, ax21 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig21.tight_layout()
            
            # WR plot
            ax21.scatter(WR_plot['DCR(DE_pos)'], WR_plot['Height(mm)'], color='k', s=1)
            ax21.scatter(WR_plot['DCR(DE_neg)'], WR_plot['Height(mm)'], color='k', s=1)

            # 허용치 기준선
            ax21.axvline(x = DCR_criteria, color='r', linestyle='--')
            ax21.axvline(x = -DCR_criteria, color='r', linestyle='--')

            ax21.set_xlim(-xlim, xlim)
            ax21.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])
            
            # 기타
            ax21.grid(linestyle='-.')
            ax21.set_xlabel('DCR')
            ax21.set_ylabel('Story')
            # ax21.set_title('Wall Rotation (1.2$\star$DBE)')
            
            memfile = BytesIO()
            plt.savefig(memfile, bbox_inches="tight")           
            plt.close()

        # MCE Plot
        if len(MCE_load_name_list) != 0:
            
            fig22, ax22 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig22.tight_layout()
            
            # WR plot
            ax22.scatter(WR_plot['DCR(MCE_pos)'], WR_plot['Height(mm)'], color='k', s=1)
            ax22.scatter(WR_plot['DCR(MCE_neg)'], WR_plot['Height(mm)'], color='k', s=1)

            # 허용치 기준선
            ax22.axvline(x = DCR_criteria, color='r', linestyle='--')
            ax22.axvline(x = -DCR_criteria, color='r', linestyle='--')

            ax22.set_xlim(-xlim, xlim)
            ax22.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax22.grid(linestyle='-.')
            ax22.set_xlabel('DCR')
            ax22.set_ylabel('Story')
            # ax22.set_title('Wall Rotation (MCE)')
            
            memfile2 = BytesIO()
            plt.savefig(memfile2, bbox_inches="tight")
            plt.close()      
        
        # 표에 그래프 넣기            
        plots_row = WR_plots_table.rows[0]
        plots_cell_x = plots_row.cells[0]
        plots_cell_y = plots_row.cells[1]
        plots_para_x = plots_cell_x.paragraphs[0]
        plots_para_y = plots_cell_y.paragraphs[0]
        plots_run_x = plots_para_x.add_run()
        plots_run_y = plots_para_y.add_run()
        plots_run_x.add_picture(memfile, width=Cm(7), height=Cm(9.5))
        plots_run_y.add_picture(memfile2, width=Cm(7), height=Cm(9.5))
        plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
        plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER       
            
    #%% Wall Shear Force
    if get_WSF == True:
        with open('pkl/WSF.pkl', 'rb') as f:
            WSF_result = pickle.load(f)
        
        # 결과값 classify & assign
        wall_result = WSF_result[0]
        story_info = WSF_result[1]
        DE_load_name_list = WSF_result[2]
        MCE_load_name_list = WSF_result[3]
        
        # Wall Shear Force 표 작성
        # template의 2번 표 불러오기
        WSF_plots_table = document.tables[6]            
        
        # DE Plot
        if len(DE_load_name_list) != 0:
            
            fig23, ax23 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig23.tight_layout()
            
            # WSF plot
            ax23.scatter(wall_result['DE'], wall_result['Height(mm)'], color = 'k', s=1)

            # 허용치 기준선
            ax23.axvline(x = DCR_criteria, color='r', linestyle='--')

            ax23.set_xlim(0, xlim)
            ax23.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax23.grid(linestyle='-.')
            ax23.set_xlabel('DCR')
            ax23.set_ylabel('Story')
            # ax23.set_title('Shear Strength (1.2$\star$DBE)')
            
            memfile = BytesIO()
            plt.savefig(memfile, bbox_inches="tight")           
            plt.close()
            
        # MCE Plot
        if len(MCE_load_name_list) != 0:
            
            fig24, ax24 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig24.tight_layout()
            
            # WSF plot
            ax24.scatter(wall_result['MCE'], wall_result['Height(mm)'], color = 'k', s=1)

            # 허용치 기준선
            ax24.axvline(x = DCR_criteria, color='r', linestyle='--')

            ax24.set_xlim(0, xlim)
            ax24.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax24.grid(linestyle='-.')
            ax24.set_xlabel('DCR')
            ax24.set_ylabel('Story')
            # ax24.set_title('Shear Strength (MCE)')
            
            memfile2 = BytesIO()
            plt.savefig(memfile2, bbox_inches="tight")
            plt.close()      
        
        # 표에 그래프 넣기            
        plots_row = WSF_plots_table.rows[0]
        plots_cell_x = plots_row.cells[0]
        plots_cell_y = plots_row.cells[1]
        plots_para_x = plots_cell_x.paragraphs[0]
        plots_para_y = plots_cell_y.paragraphs[0]
        plots_run_x = plots_para_x.add_run()
        plots_run_y = plots_para_y.add_run()
        plots_run_x.add_picture(memfile, width=Cm(7), height=Cm(9.5))
        plots_run_y.add_picture(memfile2, width=Cm(7), height=Cm(9.5))
        plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
        plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER  
    
    #%%        
    # 결과 저장할 경로
    # Path 지정
    result_path = os.path.dirname(result_xlsx_path[0])
    docx_file_path = os.path.join(result_path,'Result Plots.docx')
    
    count = 1
    while True:
        if os.path.exists(docx_file_path):            
            docx_file_path = os.path.join(result_path,'Result Plots(%s).docx' %count)
            count += 1            
        else:
            # 결과 저장
            document.save(docx_file_path)
            break


#%% Function to Print the Result into HWP

def print_hwp(result_xlsx_path, get_base_SF=False, get_story_SF=False
              , get_IDR=False, get_BR=False, get_BSF=False, get_E_BSF=False
              , get_CR=False, get_CSF=False, get_E_CSF=False, get_WAS=False
              , get_WR=False, get_WSF=False, project_name='성능기반 내진설계'
              , bldg_name='1동', story_gap=2, max_shear=60000):
    
    # Other Parameters (향 후, UI에서 조작할 수 있게끔)
    cri_DE=0.015 # IDR
    cri_MCE=0.02 # IDR
    max_criteria=0.04 # WAS
    min_criteria=-0.002 # WAS
    DCR_criteria=1
    xlim = 2 # BR
    WAS_gage_group='AS' # WAS
    
    fig_scale = 3/4 # 그래프 크기, 축, 글씨 등 scale up/down. (scale과 반비례함)    
    
    # Call CoInitialize Function before using COM object
    hwp = win32com.client.gencache.EnsureDispatch('HWPFrame.HwpObject', pythoncom.CoInitialize())
    hwp.XHwpWindows.Item(0).Visible = False
    hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule") # 작동안함. FilePathCheck dll이 설치되어야 함
    hwp.Open(os.path.join(os.getcwd(), 'template/report_template.hwp')) # 한글 api는 절대경로만 가능하도록 제한되어있음
    
    # global 필드에 변수 입력
    hwp.PutFieldText('프로젝트명', project_name) # int도 입력 가능
    hwp.PutFieldText('건물명', bldg_name) # int도 입력 가능
    
    #%% Base Shear
    if get_base_SF == True:
        with open('pkl/base_SF.pkl', 'rb') as f:
            base_SF_result = pickle.load(f)
        
        # 결과값 classify & assign
        base_shear_H1 = base_SF_result[0]
        base_shear_H2 = base_SF_result[1]
        DE_load_name_list = base_SF_result[2]
        MCE_load_name_list = base_SF_result[3]
        
        # Plot
        # DE Plot
        if len(DE_load_name_list) != 0:
            # H1_DE
            fig1, ax1 = plt.subplots(1,1, figsize=(7*fig_scale, 6.5*fig_scale), dpi=200)
            fig1.tight_layout() # 이거 안하면 크기 맘대로 바뀜 ㅠ
            ax1.set_ylim(0, max_shear)
        
            ax1.bar(range(len(DE_load_name_list)), base_shear_H1.iloc[0, 0:len(DE_load_name_list)]\
                    , color='darkblue', edgecolor='k', label = 'Max. Base Shear')
            ax1.axhline(y= base_shear_H1.iloc[0, 0:len(DE_load_name_list)].mean(), color='r', linestyle='-', label='Average')
            ax1.set_xticks(range(14), range(1,15))
            
            ax1.set_xlabel('Ground Motion No.')
            ax1.set_ylabel('Base Shear(kN)')
            ax1.legend(loc = 2)
            # ax1.set_title('X 1.2$\star$DBE')
            
            base_SF_avg_DE_x = Decimal(str(base_shear_H1.iloc[0, 0:len(DE_load_name_list)].mean()))\
                .quantize(Decimal('1'), rounding=ROUND_UP)        
            
            plt.savefig('images/base_SF_DE_X_fig.png', bbox_inches="tight")
            plt.close()
            
            # H2_DE
            fig2, ax2 = plt.subplots(1,1, figsize=(7*fig_scale, 6.5*fig_scale), dpi=200)
            fig2.tight_layout()
            ax2.set_ylim(0, max_shear)
            
            ax2.bar(range(len(DE_load_name_list)), base_shear_H2.iloc[0, 0:len(DE_load_name_list)], color='darkblue', edgecolor='k', label = 'Max. Base Shear')
            ax2.axhline(y= base_shear_H2.iloc[0, 0:len(DE_load_name_list)].mean(), color='r', linestyle='-', label='Average')
            ax2.set_xticks(range(14), range(1,15))
            
            ax2.set_xlabel('Ground Motion No.')
            ax2.set_ylabel('Base Shear(kN)')
            ax2.legend(loc = 2)
            # ax2.set_title('Y 1.2$\star$DBE')
            
            base_SF_avg_DE_y = Decimal(str(base_shear_H2.iloc[0, 0:len(DE_load_name_list)].mean()))\
                .quantize(Decimal('1'), rounding=ROUND_UP)
            
            plt.savefig('images/base_SF_DE_Y_fig.png', bbox_inches="tight")
            plt.close()

            
            # 그림 개체속성 - 번호 종류:없음 - 으로 바꾸는 함수
            def change_num_type():
                act = hwp.CreateAction('ShapeObjDialog')
                param = act.CreateSet()
                act.GetDefault(param)
                param.SetItem('NumberingType', None)
                act.Execute(param)
            
            # hwp 파일의 각각의 필드에 text or image 넣기
            hwp.PutFieldText('base_SF_DE_X_avg', f'{base_SF_avg_DE_x:,}') # 1000 자리마다 , 찍기
            hwp.PutFieldText('base_SF_DE_Y_avg', f'{base_SF_avg_DE_y:,}') # 1000 자리마다 , 찍기
            hwp.MoveToField('base_SF_DE_fig') # 기준 필드로 이동
            hwp.HAction.Run('TableLeftCell') # 왼쪽 셀로 이동
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/base_SF_DE_X_fig.png'), sizeoption=1, Width=70, Height=65)
            # 번호 종류:없음 - 으로 변경
            hwp.FindCtrl()
            change_num_type()
            hwp.HAction.Run('TableRightCell')
            hwp.HAction.Run('TableRightCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/base_SF_DE_Y_fig.png'), sizeoption=1, Width=70, Height=65)
            hwp.FindCtrl()
            change_num_type()
            
        # MCE Plot
        if len(MCE_load_name_list) != 0:
            # H1_MCE
            fig3, ax3 = plt.subplots(1,1, figsize=(7*fig_scale, 6.5*fig_scale), dpi=200)
            fig3.tight_layout()
            ax3.set_ylim(0, max_shear)
            
            ax3.bar(range(len(MCE_load_name_list)), base_shear_H1\
                    .iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)]\
                    , color='darkblue', edgecolor='k', label = 'Max. Base Shear')
            ax3.axhline(y= base_shear_H1.iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)]\
                        .mean(), color='r', linestyle='-', label='Average')
            ax3.set_xticks(range(14), range(1,15))
            
            ax3.set_xlabel('Ground Motion No.')
            ax3.set_ylabel('Base Shear(kN)')
            ax3.legend(loc = 2)
            # ax3.set_title('X MCE')
            
            base_SF_avg_MCE_x = Decimal(str(base_shear_H1.iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)].mean()))\
                .quantize(Decimal('1'), rounding=ROUND_UP)
            
            plt.savefig('images/base_SF_MCE_X_fig.png', bbox_inches="tight")         
            plt.close()            

            # H2_MCE
            fig4, ax4 = plt.subplots(1,1, figsize=(7*fig_scale, 6.5*fig_scale), dpi=200)
            fig4.tight_layout()
            ax4.set_ylim(0, max_shear)
            
            plt.bar(range(len(MCE_load_name_list)), base_shear_H2\
                    .iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)]\
                    , color='darkblue', edgecolor='k', label = 'Max. Base Shear')
            plt.axhline(y= base_shear_H2.iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)]\
                        .mean(), color='r', linestyle='-', label='Average')
            ax4.set_xticks(range(14), range(1,15))
            
            ax4.set_xlabel('Ground Motion No.')
            ax4.set_ylabel('Base Shear(kN)')
            ax4.legend(loc = 2)
            # ax4.set_title('Y MCE')
            
            base_SF_avg_MCE_y = Decimal(str(base_shear_H2.iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)].mean()))\
                .quantize(Decimal('1'), rounding=ROUND_UP)

            plt.savefig('images/base_SF_MCE_Y_fig.png', bbox_inches="tight")
            plt.close()            
            
            # hwp 파일의 각각의 필드에 text or image 넣기
            hwp.PutFieldText('base_SF_MCE_X_avg', f'{base_SF_avg_MCE_x:,}') # 1000 자리마다 , 찍기
            hwp.PutFieldText('base_SF_MCE_Y_avg', f'{base_SF_avg_MCE_y:,}') # 1000 자리마다 , 찍기
            hwp.MoveToField('base_SF_MCE_fig') # 기준 필드로 이동
            hwp.HAction.Run('TableLeftCell') # 왼쪽 셀로 이동
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/base_SF_MCE_X_fig.png'), sizeoption=1, Width=70, Height=65)
            hwp.FindCtrl()
            change_num_type()
            hwp.HAction.Run('TableRightCell')
            hwp.HAction.Run('TableRightCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/base_SF_MCE_Y_fig.png'), sizeoption=1, Width=70, Height=65)
            hwp.FindCtrl()
            change_num_type()

    #%% Story Drift
    if get_story_SF == True:
        with open('pkl/story_SF.pkl', 'rb') as f:
            story_SF_result = pickle.load(f)
        
        # 결과값 classify & assign
        shear_force_H1_max = story_SF_result[0]
        shear_force_H2_max = story_SF_result[1]
        DE_load_name_list = story_SF_result[2]
        MCE_load_name_list = story_SF_result[3]

        # Plot
        # DE Plot
        if len(DE_load_name_list) != 0:
            # H1_DE
            fig5, ax5 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig5.tight_layout()
            
            # 지진파별 plot
            for i in range(len(DE_load_name_list)):
                ax5.plot(shear_force_H1_max.iloc[:,i], range(shear_force_H1_max.shape[0]), label=DE_load_name_list[i], linewidth=0.7)
                
            # 평균 plot
            ax5.plot(shear_force_H1_max.iloc[:,0:len(DE_load_name_list)]\
                    .mean(axis=1), range(shear_force_H1_max.shape[0]), color='k', label='Average', linewidth=2)
            
            ax5.set_xlim(0, max_shear)
            ax5.set_yticks(range(shear_force_H1_max.shape[0])[::story_gap], shear_force_H1_max.index[::story_gap], fontsize=8.5)
            
            # 기타
            ax5.grid(linestyle='-.')
            ax5.set_xlabel('Story Shear(kN)')
            ax5.set_ylabel('Story')
            ax5.legend(loc=1, fontsize=8)
            # ax5.set_title('X 1.2$\star$DBE')
            
            plt.savefig('images/story_SF_DE_X_fig.png', bbox_inches="tight")          
            plt.close()
            
            # H2_DE
            fig6, ax6 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig6.tight_layout()
            
            for i in range(len(DE_load_name_list)):
                ax6.plot(shear_force_H2_max.iloc[:,i], range(shear_force_H2_max.shape[0]), label=DE_load_name_list[i], linewidth=0.7)
            
            ax6.plot(shear_force_H2_max.iloc[:,0:len(DE_load_name_list)]\
                    .mean(axis=1), range(shear_force_H2_max.shape[0]), color='k', label='Average', linewidth=2)
            
            ax6.set_xlim(0, max_shear)
            ax6.set_yticks(range(shear_force_H2_max.shape[0])[::story_gap], shear_force_H2_max.index[::story_gap], fontsize=8.5)
        
            ax6.grid(linestyle='-.')
            ax6.set_xlabel('Story Shear(kN)')
            ax6.set_ylabel('Story')
            ax6.legend(loc=1, fontsize=8)
            # ax6.set_title('Y 1.2$\star$DBE')
            
            plt.savefig('images/story_SF_DE_Y_fig.png', bbox_inches="tight")
            plt.close()
            
            # hwp 파일의 각각의 필드에 text or image 넣기
            hwp.MoveToField('story_SF_DE_fig')
            hwp.HAction.Run('TableLeftCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/story_SF_DE_X_fig.png'), sizeoption=1, Width=70, Height=95)
            hwp.FindCtrl()
            change_num_type()
            hwp.HAction.Run('TableRightCell')
            hwp.HAction.Run('TableRightCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/story_SF_DE_Y_fig.png'), sizeoption=1, Width=70, Height=95)
            hwp.FindCtrl()
            change_num_type()
            
        # MCE Plot
        if len(MCE_load_name_list) != 0:
            # H1_MCE
            fig7, ax7 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig7.tight_layout()
            
            for i in range(len(MCE_load_name_list)):
                ax7.plot(shear_force_H1_max.iloc[:,i+len(DE_load_name_list)], range(shear_force_H1_max.shape[0]), label=MCE_load_name_list[i], linewidth=0.7)
            ax7.plot(shear_force_H1_max.iloc[:,len(DE_load_name_list)\
                                                    :len(DE_load_name_list)+len(MCE_load_name_list)]\
                            .mean(axis=1), range(shear_force_H1_max.shape[0]), color='k', label='Average', linewidth=2)
            
            ax7.set_xlim(0, max_shear)
            ax7.set_yticks(range(shear_force_H1_max.shape[0])[::story_gap], shear_force_H1_max.index[::story_gap], fontsize=8.5)
        
            ax7.grid(linestyle='-.')
            ax7.set_xlabel('Story Shear(kN)')
            ax7.set_ylabel('Story')
            ax7.legend(loc=1, fontsize=8)
            # ax7.set_title('X MCE')
        
            plt.savefig('images/story_SF_MCE_X_fig.png', bbox_inches="tight")           
            plt.close()
            
            # H1_MCE
            fig8, ax8 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig8.tight_layout()
            
            for i in range(len(MCE_load_name_list)):
                ax8.plot(shear_force_H2_max.iloc[:,i+len(DE_load_name_list)], range(shear_force_H2_max.shape[0]), label=MCE_load_name_list[i], linewidth=0.7)
            ax8.plot(shear_force_H2_max.iloc[:,len(DE_load_name_list)\
                                                    :len(DE_load_name_list)+len(MCE_load_name_list)]\
                            .mean(axis=1), range(shear_force_H2_max.shape[0]), color='k', label='Average', linewidth=2)
            
            ax8.set_xlim(0, max_shear)
            ax8.set_yticks(range(shear_force_H2_max.shape[0])[::story_gap], shear_force_H2_max.index[::story_gap], fontsize=8.5)
        
            ax8.grid(linestyle='-.')
            ax8.set_xlabel('Story Shear(kN)')
            ax8.set_ylabel('Story')
            ax8.legend(loc=1, fontsize=8)
            # ax8.set_title('Y MCE')
            
            plt.savefig('images/story_SF_MCE_Y_fig.png', bbox_inches="tight")
            plt.close()
            
            # hwp 파일의 각각의 필드에 text or image 넣기
            hwp.MoveToField('story_SF_MCE_fig')
            hwp.HAction.Run('TableLeftCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/story_SF_MCE_X_fig.png'), sizeoption=1, Width=70, Height=95)
            hwp.FindCtrl()
            change_num_type()
            hwp.HAction.Run('TableRightCell')
            hwp.HAction.Run('TableRightCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/story_SF_MCE_Y_fig.png'), sizeoption=1, Width=70, Height=95)
            hwp.FindCtrl()
            change_num_type()
            
    #%% Inter-Story Drift
    if get_IDR == True:
        with open('pkl/IDR.pkl', 'rb') as f:
            IDR_result = pickle.load(f)
        
        # 결과값 classify & assign
        IDR_result_each = IDR_result[0]
        IDR_result_avg = IDR_result[1]
        DE_load_name_list = IDR_result[2]
        MCE_load_name_list = IDR_result[3]
        story_name_window_reordered = IDR_result[4]
        
        # Plot
        count_x = 0
        count_y = 2
        count_avg = 0
        
        # DE Plot
        if len(DE_load_name_list) != 0:
            # H1_DE   
            fig9, ax9 = plt.subplots(1,1, figsize=(7*fig_scale, 10*fig_scale), dpi=200)
            fig9.tight_layout()
            
            # 지진파별 plot
            for load_name in DE_load_name_list:
                ax9.plot(IDR_result_each[count_x].iloc[:,-1], IDR_result_each[count_x].iloc[:,0]
                         , label='{}'.format(load_name), linewidth=0.7)
                ax9.plot(IDR_result_each[count_x+1].iloc[:,-1], IDR_result_each[count_x].iloc[:,0]
                         , linewidth=0.7)
                count_x += 4
                
            # 평균 plot
            ax9.plot(IDR_result_avg[count_avg].iloc[:,0], story_name_window_reordered, color='k', label='Average', linewidth=2)
            ax9.plot(IDR_result_avg[count_avg].iloc[:,1], story_name_window_reordered, color='k', linewidth=2)
            
            # reference line 그려서 허용치 나타내기
            ax9.axvline(x=-cri_DE, color='r', linestyle='--', label='LS')
            ax9.axvline(x=cri_DE, color='r', linestyle='--')
            
            ax9.set_xlim(-0.025, 0.025)
            ax9.set_yticks(story_name_window_reordered[::story_gap], story_name_window_reordered[::story_gap])
            
            # 기타
            ax9.grid(linestyle='-.')
            ax9.set_xlabel('Interstory Drift Ratios(m/m)')
            ax9.set_ylabel('Story')
            ax9.legend(loc=4, fontsize=8)
            # ax9.set_title('X 1.2$\star$DBE')
            
            plt.savefig('images/IDR_DE_X_fig.png', bbox_inches="tight")           
            plt.close()
            
            # H2_DE
            fig10, ax10 = plt.subplots(1,1, figsize=(7*fig_scale, 10*fig_scale), dpi=200)
            fig10.tight_layout()

            # 지진파별 plot
            for load_name in DE_load_name_list:
                ax10.plot(IDR_result_each[count_y].iloc[:,-1], IDR_result_each[count_y].iloc[:,0]
                         , label='{}'.format(load_name), linewidth=0.7)
                ax10.plot(IDR_result_each[count_y+1].iloc[:,-1], IDR_result_each[count_y].iloc[:,0]
                         , linewidth=0.7)
                count_y += 4
                
            # 평균 plot
            ax10.plot(IDR_result_avg[count_avg].iloc[:,0], story_name_window_reordered, color='k', label='Average', linewidth=2)
            ax10.plot(IDR_result_avg[count_avg].iloc[:,1], story_name_window_reordered, color='k', linewidth=2)
            count_avg += 1
            
            # reference line 그려서 허용치 나타내기
            ax10.axvline(x=-cri_DE, color='r', linestyle='--', label='LS')
            ax10.axvline(x=cri_DE, color='r', linestyle='--')
            
            ax10.set_xlim(-0.025, 0.025)
            ax10.set_yticks(story_name_window_reordered[::story_gap], story_name_window_reordered[::story_gap])
            
            # 기타
            ax10.grid(linestyle='-.')
            ax10.set_xlabel('Interstory Drift Ratios(m/m)')
            ax10.set_ylabel('Story')
            ax10.legend(loc=4, fontsize=8)
            # ax10.set_title('Y 1.2$\star$DBE')
            
            plt.savefig('images/IDR_DE_Y_fig.png', bbox_inches="tight")
            plt.close()

            # hwp 파일의 각각의 필드에 text or image 넣기
            hwp.MoveToField('IDR_DE_fig')
            hwp.HAction.Run('TableLeftCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/IDR_DE_X_fig.png'), sizeoption=1, Width=70, Height=100)
            hwp.FindCtrl()
            change_num_type()
            hwp.HAction.Run('TableRightCell')
            hwp.HAction.Run('TableRightCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/IDR_DE_Y_fig.png'), sizeoption=1, Width=70, Height=100)
            hwp.FindCtrl()
            change_num_type()
            
        # MCE Plot
        if len(MCE_load_name_list) != 0:
            # H1_MCE
            fig11, ax11 = plt.subplots(1,1, figsize=(7*fig_scale, 10*fig_scale), dpi=200)
            fig11.tight_layout()
            
            # 지진파별 plot
            for load_name in MCE_load_name_list:
                ax11.plot(IDR_result_each[count_x].iloc[:,-1], IDR_result_each[count_x].iloc[:,0]
                         , label='{}'.format(load_name), linewidth=0.7)
                ax11.plot(IDR_result_each[count_x+1].iloc[:,-1], IDR_result_each[count_x].iloc[:,0]
                         , linewidth=0.7)
                count_x += 4
                
            # 평균 plot
            ax11.plot(IDR_result_avg[count_avg].iloc[:,0], story_name_window_reordered, color='k', label='Average', linewidth=2)
            ax11.plot(IDR_result_avg[count_avg].iloc[:,1], story_name_window_reordered, color='k', linewidth=2)
            
            # reference line 그려서 허용치 나타내기
            ax11.axvline(x=-cri_MCE, color='r', linestyle='--', label='CP')
            ax11.axvline(x=cri_MCE, color='r', linestyle='--')
            
            ax11.set_xlim(-0.025, 0.025)
            ax11.set_yticks(story_name_window_reordered[::story_gap], story_name_window_reordered[::story_gap])
            
            # 기타
            ax11.grid(linestyle='-.')
            ax11.set_xlabel('Interstory Drift Ratios(m/m)')
            ax11.set_ylabel('Story')
            ax11.legend(loc=4, fontsize=8)
            # ax11.set_title('X MCE')
            
            plt.savefig('images/IDR_MCE_X_fig.png', bbox_inches="tight")
            plt.close()

            # H2_MCE
            fig12, ax12 = plt.subplots(1,1, figsize=(7*fig_scale, 10*fig_scale), dpi=200)
            fig12.tight_layout()

            # 지진파별 plot
            for load_name in MCE_load_name_list:
                ax12.plot(IDR_result_each[count_y].iloc[:,-1], IDR_result_each[count_y].iloc[:,0]
                         , label='{}'.format(load_name), linewidth=0.7)
                ax12.plot(IDR_result_each[count_y+1].iloc[:,-1], IDR_result_each[count_y].iloc[:,0]
                         , linewidth=0.7)
                count_y += 4
                
            # 평균 plot
            ax12.plot(IDR_result_avg[count_avg].iloc[:,0], story_name_window_reordered, color='k', label='Average', linewidth=2)
            ax12.plot(IDR_result_avg[count_avg].iloc[:,1], story_name_window_reordered, color='k', linewidth=2)
            count_avg += 1
            
            # reference line 그려서 허용치 나타내기
            ax12.axvline(x=-cri_MCE, color='r', linestyle='--', label='CP')
            ax12.axvline(x=cri_MCE, color='r', linestyle='--')
            
            ax12.set_xlim(-0.025, 0.025)
            ax12.set_yticks(story_name_window_reordered[::story_gap], story_name_window_reordered[::story_gap])
            
            # 기타
            ax12.grid(linestyle='-.')
            ax12.set_xlabel('Interstory Drift Ratios(m/m)')
            ax12.set_ylabel('Story')
            ax12.legend(loc=4, fontsize=8)
            # ax12.set_title('Y MCE')

            plt.savefig('images/IDR_MCE_Y_fig.png', bbox_inches="tight")
            plt.close()            
            
            # hwp 파일의 각각의 필드에 text or image 넣기
            hwp.MoveToField('IDR_MCE_fig')
            hwp.HAction.Run('TableLeftCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/IDR_MCE_X_fig.png'), sizeoption=1, Width=70, Height=100)
            hwp.FindCtrl()
            change_num_type()
            hwp.HAction.Run('TableRightCell')
            hwp.HAction.Run('TableRightCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/IDR_MCE_Y_fig.png'), sizeoption=1, Width=70, Height=100)       
            hwp.FindCtrl()
            change_num_type()

    #%% Beam Rotation
    if get_BR == True:
        # Load Pickle Files
        with open('pkl/BR.pkl', 'rb') as f:
            BR_result = pickle.load(f)
        
        # 결과값 classify & assign
        BR_plot = BR_result[0]
        story_info = BR_result[1]
        DE_load_name_list = BR_result[2]
        MCE_load_name_list = BR_result[3]
        
        # Plot
        # DE Plot
        if len(DE_load_name_list) != 0:

            fig13, ax13 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig13.tight_layout()
            
            # DCR plot                
            ax13.scatter(BR_plot['DCR(DE_pos)'], BR_plot['Height(mm)'], color='k', s=1)
            ax13.scatter(BR_plot['DCR(DE_neg)'], BR_plot['Height(mm)'], color='k', s=1)

            # 허용치(DCR) 기준선
            ax13.axvline(x = DCR_criteria, color='r', linestyle='--')
            ax13.axvline(x = -DCR_criteria, color='r', linestyle='--')

            ax13.set_xlim(-xlim, xlim)
            ax13.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax13.grid(linestyle='-.')
            ax13.set_xlabel('DCR')
            ax13.set_ylabel('Story')
            # ax13.set_title('Beam Rotation (1.2$\star$DBE)')
            
            plt.savefig('images/BR_DE_fig.png', bbox_inches="tight")
            plt.close()
            
            # hwp 파일의 각각의 필드에 text or image 넣기
            hwp.MoveToField('BR_fig')
            hwp.HAction.Run('TableLeftCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/BR_DE_fig.png'), sizeoption=1, Width=70, Height=95)
            hwp.FindCtrl()
            change_num_type()
            
        if len(MCE_load_name_list) != 0:
            
            fig14, ax14 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig14.tight_layout()

            # DCR plot                
            ax14.scatter(BR_plot['DCR(MCE_pos)'], BR_plot['Height(mm)'], color='k', s=1)
            ax14.scatter(BR_plot['DCR(MCE_neg)'], BR_plot['Height(mm)'], color='k', s=1)

            # 허용치(DCR) 기준선
            ax14.axvline(x = DCR_criteria, color='r', linestyle='--')
            ax14.axvline(x = -DCR_criteria, color='r', linestyle='--')

            ax14.set_xlim(-xlim, xlim)
            ax14.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax14.grid(linestyle='-.')
            ax14.set_xlabel('DCR')
            ax14.set_ylabel('Story')
            # ax14.set_title('Beam Rotation (MCE)')
            
            plt.savefig('images/BR_MCE_fig.png', bbox_inches="tight")
            plt.close()

            # hwp 파일의 각각의 필드에 text or image 넣기
            hwp.MoveToField('BR_fig')
            hwp.HAction.Run('TableRightCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/BR_MCE_fig.png'), sizeoption=1, Width=70, Height=95)
            hwp.FindCtrl()
            change_num_type()
            
    #%% Beam Shear Force
    if get_BSF == True:
        # Load Pickle Files
        with open('pkl/BSF.pkl', 'rb') as f:
            BSF_result = pickle.load(f)
        
        # 결과값 classify & assign
        BSF_plot = BSF_result[0]
        story_info = BSF_result[1]
        DE_load_name_list = BSF_result[2]
        MCE_load_name_list = BSF_result[3]
        
        # Plot
        # DE Plot
        if len(DE_load_name_list) != 0:
            fig15, ax15 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig15.tight_layout()
            
            # DCR plot                
            ax15.axes.scatter(BSF_plot['DE'], BSF_plot['Height(mm)'], color='k', s=1)
    
            # 허용치(DCR) 기준선
            ax15.axes.axvline(x = DCR_criteria, color='r', linestyle='--')
    
            ax15.axes.set_xlim(0, xlim)
            ax15.axes.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])
    
            # 기타
            ax15.axes.grid(linestyle='-.')
            ax15.axes.set_xlabel('DCR')
            ax15.axes.set_ylabel('Story')
            # ax15.axes.set_title('Shear Strength (1.2$\star$DBE)')
            
            plt.savefig('images/BSF_DE_fig.png', bbox_inches="tight")
            plt.close()
            
            # hwp 파일의 각각의 필드에 text or image 넣기
            hwp.MoveToField('BSF_fig')
            hwp.HAction.Run('TableLeftCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/BSF_DE_fig.png'), sizeoption=1, Width=70, Height=95)
            hwp.FindCtrl()
            change_num_type()
            
        # Plot
        # MCE Plot
        if len(MCE_load_name_list) != 0:
            fig16, ax16 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig16.tight_layout()
            
            # DCR plot                
            ax16.axes.scatter(BSF_plot['MCE'], BSF_plot['Height(mm)'], color='k', s=1)
    
            # 허용치(DCR) 기준선
            ax16.axes.axvline(x = DCR_criteria, color='r', linestyle='--')
    
            ax16.axes.set_xlim(0, xlim)
            ax16.axes.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])
    
            # 기타
            ax16.axes.grid(linestyle='-.')
            ax16.axes.set_xlabel('DCR')
            ax16.axes.set_ylabel('Story')
            # ax16.axes.set_title('Shear Strength (MCE)')
            
            plt.savefig('images/BSF_MCE_fig.png', bbox_inches="tight")
            plt.close()

            # hwp 파일의 각각의 필드에 text or image 넣기
            hwp.MoveToField('BSF_fig')
            hwp.HAction.Run('TableRightCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/BSF_MCE_fig.png'), sizeoption=1, Width=70, Height=95)
            hwp.FindCtrl()
            change_num_type()
            
    #%% Wall Axial Strain
    if get_WAS == True:
        with open('pkl/WAS.pkl', 'rb') as f:
            WAS_result = pickle.load(f)
        
        # 결과값 classify & assign
        WAS_plot = WAS_result[0]
        story_info = WAS_result[1]
        DE_load_name_list = WAS_result[2]
        MCE_load_name_list = WAS_result[3]           
        
        # DE Plot
        if len(DE_load_name_list) != 0:
            # DE_Neg
            fig17, ax17 = plt.subplots(1,1, figsize=(7*fig_scale, 10*fig_scale), dpi=200)
            fig17.tight_layout()
            
            # WAS plot
            ax17.scatter(WAS_plot['DE(Compressive)'], WAS_plot['Height(mm)'], color='k', s=5)
            ax17.scatter(WAS_plot['DE(Tensile)'], WAS_plot['Height(mm)'], color='k', s=5)

            # 허용치 기준선
            ax17.axvline(x=min_criteria, color='r', linestyle='--')
            ax17.axvline(x=max_criteria, color='r', linestyle='--')

            ax17.set_xlim(-0.003, 0)
            ax17.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax17.grid(linestyle='-.')
            ax17.set_xlabel('Axial Strain (m/m)')
            ax17.set_ylabel('Story')
            # ax17.set_title('1.2$\star$DBE (Compressive)')
            
            plt.savefig('images/WAS_DE_C_fig.png', bbox_inches="tight")
            plt.close()
            
            # DE_Pos
            fig18, ax18 = plt.subplots(1,1, figsize=(7*fig_scale, 10*fig_scale), dpi=200)
            fig18.tight_layout()
            
            # WAS plot
            ax18.scatter(WAS_plot['DE(Compressive)'], WAS_plot['Height(mm)'], color='k', s=5)
            ax18.scatter(WAS_plot['DE(Tensile)'], WAS_plot['Height(mm)'], color='k', s=5)

            # 허용치 기준선
            ax18.axvline(x=min_criteria, color='r', linestyle='--')
            ax18.axvline(x=max_criteria, color='r', linestyle='--')

            ax18.set_xlim(0, 0.013)
            ax18.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax18.grid(linestyle='-.')
            ax18.set_xlabel('Axial Strain (m/m)')
            ax18.set_ylabel('Story')
            # ax18.set_title('1.2$\star$DBE (Tensile)')
            
            plt.savefig('images/WAS_DE_T_fig.png', bbox_inches="tight")
            plt.close()          
            
            # hwp 파일의 각각의 필드에 text or image 넣기
            hwp.MoveToField('WAS_DE_fig')
            hwp.HAction.Run('TableLeftCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/WAS_DE_C_fig.png'), sizeoption=1, Width=70, Height=100)
            hwp.FindCtrl()
            change_num_type()
            hwp.HAction.Run('TableRightCell')
            hwp.HAction.Run('TableRightCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/WAS_DE_T_fig.png'), sizeoption=1, Width=70, Height=100)
            hwp.FindCtrl()
            change_num_type()
            
        # MCE Plot
        if len(MCE_load_name_list) != 0:
            # MCE_Neg
            fig19, ax19 = plt.subplots(1,1, figsize=(7*fig_scale, 10*fig_scale), dpi=200)
            fig19.tight_layout()
            
            # WAS plot
            ax19.scatter(WAS_plot['MCE(Compressive)'], WAS_plot['Height(mm)'], color='k', s=5)
            ax19.scatter(WAS_plot['MCE(Tensile)'], WAS_plot['Height(mm)'], color='k', s=5)

            # 허용치 기준선
            ax19.axvline(x=min_criteria, color='r', linestyle='--')
            ax19.axvline(x=max_criteria, color='r', linestyle='--')

            ax19.set_xlim(-0.003, 0)
            ax19.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax19.grid(linestyle='-.')
            ax19.set_xlabel('Axial Strain (m/m)')
            ax19.set_ylabel('Story')
            # ax19.set_title('MCE (Compressive)')
            
            plt.savefig('images/WAS_MCE_C_fig.png', bbox_inches="tight")
            plt.close()
            
            # MCE_Pos
            fig20, ax20 = plt.subplots(1,1, figsize=(7*fig_scale, 10*fig_scale), dpi=200)
            fig20.tight_layout()
            
            # WAS plot
            ax20.scatter(WAS_plot['MCE(Compressive)'], WAS_plot['Height(mm)'], color='k', s=5)
            ax20.scatter(WAS_plot['MCE(Tensile)'], WAS_plot['Height(mm)'], color='k', s=5)

            # 허용치 기준선
            ax20.axvline(x=min_criteria, color='r', linestyle='--')
            ax20.axvline(x=max_criteria, color='r', linestyle='--')

            ax20.set_xlim(0, 0.013)
            ax20.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax20.grid(linestyle='-.')
            ax20.set_xlabel('Axial Strain (m/m)')
            ax20.set_ylabel('Story')
            # ax20.set_title('MCE (Tensile)')
            
            plt.savefig('images/WAS_MCE_T_fig.png', bbox_inches="tight")
            plt.close()      
            
            # hwp 파일의 각각의 필드에 text or image 넣기
            hwp.MoveToField('WAS_MCE_fig')
            hwp.HAction.Run('TableLeftCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/WAS_MCE_C_fig.png'), sizeoption=1, Width=70, Height=100)
            hwp.FindCtrl()
            change_num_type()
            hwp.HAction.Run('TableRightCell')
            hwp.HAction.Run('TableRightCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/WAS_MCE_T_fig.png'), sizeoption=1, Width=70, Height=100)
            hwp.FindCtrl()
            change_num_type()
            
    #%% Wall Rotation
    if get_WR == True:
        with open('pkl/WR.pkl', 'rb') as f:
            WR_result = pickle.load(f)
        
        # 결과값 classify & assign
        WR_plot = WR_result[0]
        story_info = WR_result[1]
        DE_load_name_list = WR_result[2]
        MCE_load_name_list = WR_result[3]           
        
        # DE Plot
        if len(DE_load_name_list) != 0:
            
            fig21, ax21 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig21.tight_layout()
            
            # WR plot
            ax21.scatter(WR_plot['DCR(DE_pos)'], WR_plot['Height(mm)'], color='k', s=1)
            ax21.scatter(WR_plot['DCR(DE_neg)'], WR_plot['Height(mm)'], color='k', s=1)

            # 허용치 기준선
            ax21.axvline(x = DCR_criteria, color='r', linestyle='--')
            ax21.axvline(x = -DCR_criteria, color='r', linestyle='--')

            ax21.set_xlim(-xlim, xlim)
            ax21.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])
            
            # 기타
            ax21.grid(linestyle='-.')
            ax21.set_xlabel('DCR')
            ax21.set_ylabel('Story')
            # ax21.set_title('Wall Rotation (1.2$\star$DBE)')
            
            plt.savefig('images/WR_DE_fig.png', bbox_inches="tight")
            plt.close()
            
            # hwp 파일의 각각의 필드에 text or image 넣기
            hwp.MoveToField('WR_fig')
            hwp.HAction.Run('TableLeftCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/WR_DE_fig.png'), sizeoption=1, Width=70, Height=95)
            hwp.FindCtrl()
            change_num_type()

        # MCE Plot
        if len(MCE_load_name_list) != 0:
            
            fig22, ax22 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig22.tight_layout()
            
            # WR plot
            ax22.scatter(WR_plot['DCR(MCE_pos)'], WR_plot['Height(mm)'], color='k', s=1)
            ax22.scatter(WR_plot['DCR(MCE_neg)'], WR_plot['Height(mm)'], color='k', s=1)

            # 허용치 기준선
            ax22.axvline(x = DCR_criteria, color='r', linestyle='--')
            ax22.axvline(x = -DCR_criteria, color='r', linestyle='--')

            ax22.set_xlim(-xlim, xlim)
            ax22.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax22.grid(linestyle='-.')
            ax22.set_xlabel('DCR')
            ax22.set_ylabel('Story')
            # ax22.set_title('Wall Rotation (MCE)')
            
            plt.savefig('images/WR_MCE_fig.png', bbox_inches="tight")
            plt.close()
            
            # hwp 파일의 각각의 필드에 text or image 넣기
            hwp.MoveToField('WR_fig')
            hwp.HAction.Run('TableRightCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/WR_MCE_fig.png'), sizeoption=1, Width=70, Height=95)
            hwp.FindCtrl()
            change_num_type()

    #%% Wall Shear Force
    if get_WSF == True:
        with open('pkl/WSF.pkl', 'rb') as f:
            WSF_result = pickle.load(f)
        
        # 결과값 classify & assign
        wall_result = WSF_result[0]
        story_info = WSF_result[1]
        DE_load_name_list = WSF_result[2]
        MCE_load_name_list = WSF_result[3]
                 
        
        # DE Plot
        if len(DE_load_name_list) != 0:
            
            fig23, ax23 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig23.tight_layout()
            
            # WSF plot
            ax23.scatter(wall_result['DE'], wall_result['Height(mm)'], color = 'k', s=1)

            # 허용치 기준선
            ax23.axvline(x = DCR_criteria, color='r', linestyle='--')

            ax23.set_xlim(0, xlim)
            ax23.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax23.grid(linestyle='-.')
            ax23.set_xlabel('DCR')
            ax23.set_ylabel('Story')
            # ax23.set_title('Shear Strength (1.2$\star$DBE)')
            
            plt.savefig('images/WSF_DE_fig.png', bbox_inches="tight")
            plt.close()
            
            # hwp 파일의 각각의 필드에 text or image 넣기
            hwp.MoveToField('WSF_fig')
            hwp.HAction.Run('TableLeftCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/WSF_DE_fig.png'), sizeoption=1, Width=70, Height=95)
            hwp.FindCtrl()
            change_num_type()
            
        # MCE Plot
        if len(MCE_load_name_list) != 0:
            
            fig24, ax24 = plt.subplots(1,1, figsize=(7*fig_scale, 9.5*fig_scale), dpi=200)
            fig24.tight_layout()
            
            # WSF plot
            ax24.scatter(wall_result['MCE'], wall_result['Height(mm)'], color = 'k', s=1)

            # 허용치 기준선
            ax24.axvline(x = DCR_criteria, color='r', linestyle='--')

            ax24.set_xlim(0, xlim)
            ax24.set_yticks(story_info['Height(mm)'][::-story_gap], story_info['Story Name'][::-story_gap])

            # 기타
            ax24.grid(linestyle='-.')
            ax24.set_xlabel('DCR')
            ax24.set_ylabel('Story')
            # ax24.set_title('Shear Strength (MCE)')
            
            plt.savefig('images/WSF_MCE_fig.png', bbox_inches="tight")
            plt.close()  
            
            # hwp 파일의 각각의 필드에 text or image 넣기
            hwp.MoveToField('WSF_fig')
            hwp.HAction.Run('TableRightCell')
            hwp.InsertPicture(os.path.join(os.getcwd(), 'images/WSF_MCE_fig.png'), sizeoption=1, Width=70, Height=95)
            hwp.FindCtrl()
            change_num_type()
            
    # 결과 저장할 경로
    # Path 지정
    result_path = os.path.dirname(result_xlsx_path[0])
    hwp_file_path = os.path.join(result_path,'Result Plots.hwp')
    
    count = 1
    while True:
        if os.path.exists(hwp_file_path):            
            hwp_file_path = os.path.join(result_path,'Result Plots(%s).hwp' %count)
            count += 1            
        else:
            # 결과 저장
            hwp.SaveAs(hwp_file_path)
            break

    hwp.Quit()
    
#%% Test Functions
# Result DataFrame Checking Function
def main_df() -> pd.DataFrame:
    # File Paths
    input_xlsx_path = r'D:/이형우/5_PBSD/용현학익7단지/708D/test/YH-708_Data Conversion_Ver.3.5_구조심의_240216.xlsx'
    result_xlsx_path1 = r"D:/이형우/5_PBSD/용현학익7단지/708D/test/YH-708_Analysis Result_DE.xlsx"
    result_xlsx_path2 = r"D:/이형우/5_PBSD/용현학익7단지/708D/test/YH-708_Analysis Result_MCE.xlsx"
    result_xlsx_path = [result_xlsx_path1, result_xlsx_path2]
    beam_design_xlsx_path = r'D:/이형우/5_PBSD/용현학익7단지/708D/test/YH-708_Seismic Design_Coupling Beam_Ver.2.2_After_re.xlsx'
    wall_design_xlsx_path = r'D:/이형우/5_PBSD/용현학익7단지/708D/test/YH-708_Seismic Design_Shear Wall_Ver.2.2_After.xlsx'
    
    # PostProc class 생성
    result = PostProc(input_xlsx_path, result_xlsx_path
                      , get_base_SF=False, get_story_SF=False
                      , get_IDR=False, get_BR=True, get_BSF=False
                      , get_E_BSF=False, get_CR=False, get_CSF=False
                      , get_E_CSF=False, get_WAS=False, get_WR=False
                      , get_WSF=False, BR_scale_factor=1.0)
    
    # pkl 파일 읽기
    # result.IDR(cri_DE=0.015, cri_MCE=0.02, yticks=2)
    result.BR(input_xlsx_path, beam_design_xlsx_path, graph=False)
    with open('pkl/BR.pkl', 'rb') as f:
        result_df = pickle.load(f)
    # pkl 폴더 삭제
    os.path.exists('pkl')
    shutil.rmtree('pkl')
    
    return result_df

# HWP Print Checking Function
def main_hwp() -> None:
    # File Paths
    input_xlsx_path = r'D:/이형우/5_PBSD/용현학익7단지/708D/test/YH-708_Data Conversion_Ver.3.5_구조심의_240216.xlsx'
    result_xlsx_path1 = r"D:/이형우/5_PBSD/용현학익7단지/708D/test/YH-708_Analysis Result_DE.xlsx"
    result_xlsx_path2 = r"D:/이형우/5_PBSD/용현학익7단지/708D/test/YH-708_Analysis Result_MCE.xlsx"
    result_xlsx_path = [result_xlsx_path1, result_xlsx_path2]
    beam_design_xlsx_path = r'D:/이형우/5_PBSD/용현학익7단지/708D/test/YH-708_Seismic Design_Coupling Beam_Ver.2.2_After_re.xlsx'
    wall_design_xlsx_path = r'D:/이형우/5_PBSD/용현학익7단지/708D/test/YH-708_Seismic Design_Shear Wall_Ver.2.2_After.xlsx'
    
    get_base_SF = True
    get_story_SF = False
    get_IDR = True
    get_BR = True
    get_BSF = True
    get_WAS = True
    get_WR = True
    get_WSF = True
    
    ylim = 70000
    
    # PostProc class 생성
    result = PostProc(input_xlsx_path, result_xlsx_path
                      , get_base_SF=get_base_SF, get_story_SF=get_story_SF
                      , get_IDR=get_IDR, get_BR=get_BR, get_BSF=get_BSF
                      , get_E_BSF=False, get_CR=False, get_CSF=False
                      , get_E_CSF=False, get_WAS=get_WAS, get_WR=get_WR
                      , get_WSF=get_WSF, BR_scale_factor=1.0)
    
    # pkl 파일 읽기
    result.base_SF(ylim=ylim)
    result.story_SF(xlim=70000)
    result.IDR()
    result.BR_plot(beam_design_xlsx_path)
    result.BSF_plot(beam_design_xlsx_path)
    result.WAS_plot(wall_design_xlsx_path)
    result.WR_plot(wall_design_xlsx_path)
    result.WSF_plot(wall_design_xlsx_path)
    
    # hwp 출력
    print_hwp(result_xlsx_path, get_base_SF=get_base_SF, get_story_SF=get_story_SF
                  , get_IDR=get_IDR, get_BR=get_BR, get_BSF=get_BSF, get_E_BSF=False
                  , get_CR=False, get_CSF=False, get_E_CSF=False, get_WAS=get_WAS
                  , get_WR=get_WR, get_WSF=get_WSF, project_name='성능기반 내진설계'
                  , bldg_name='101동', story_gap=2, max_shear=ylim)
    
    # pkl 폴더 삭제
    os.path.exists('pkl')
    shutil.rmtree('pkl')

# Execute Testing
if __name__ == '__main__':
    # main_hwp()
    # File Paths
    input_xlsx_path = r'D:/이형우/3_PBSD/용현학익7단지/706D/test/YH-706_Data Conversion_Ver.3.5_구조심의_240207.xlsx'
    result_xlsx_path1 = r"D:/이형우/3_PBSD/용현학익7단지/706D/test/YH_706_DE_Result.xlsx"
    result_xlsx_path2 = r"D:/이형우/3_PBSD/용현학익7단지/706D/test/YH_706_MCE_Result.xlsx"
    result_xlsx_path = [result_xlsx_path1, result_xlsx_path2]
    beam_design_xlsx_path = r'D:/이형우/3_PBSD/용현학익7단지/706D/test/Seismic Design_Coupling Beam_Ver.2.0_240123.xlsx'
    dbeam_design_xlsx_path = r'D:/이형우/3_PBSD/용현학익7단지/706D/test/Seismic Design_Divided Beam_Ver.3.0_240308.xlsx'
    wall_design_xlsx_path = r'D:/이형우/3_PBSD/용현학익7단지/706D/test/YH-708_Seismic Design_Shear Wall_Ver.2.2_After.xlsx'
    
    get_base_SF = False
    get_story_SF = False
    get_IDR = False
    get_BR = True
    get_BSF = True
    get_WAS = False
    get_WR = False
    get_WSF = False
    
    ylim = 70000
    
    # PostProc class 생성
    result = PostProc(input_xlsx_path, result_xlsx_path
                      , get_base_SF=get_base_SF, get_story_SF=get_story_SF
                      , get_IDR=get_IDR, get_BR=get_BR, get_BSF=get_BSF
                      , get_E_BSF=False, get_CR=False, get_CSF=False
                      , get_E_CSF=False, get_WAS=get_WAS, get_WR=get_WR
                      , get_WSF=get_WSF, BR_scale_factor=1.0)
    
    story_info = result.story_info
    cbeam_info = result.beam_info.copy()
    dbeam_info = result.dbeam_info.copy()
    rebar_info = result.rebar_info

    # Analysis Result Sheets
    node_data = result.node_data
    element_data = result.frame_data
    beam_rot_data = result.beam_rot_data

    # Seismic Loads List
    load_name_list = result.load_name_list
    gravity_load_name = result.gravity_load_name
    seismic_load_name_list = result.seismic_load_name_list
    DE_load_name_list = result.DE_load_name_list
    MCE_load_name_list = result.MCE_load_name_list
    
    scale_factor = 1.0
