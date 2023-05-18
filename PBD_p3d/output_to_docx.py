#%% Import
import os
import pandas as pd
import time
from io import BytesIO # 파일처럼 취급되는 문자열 객체 생성(메모리 낭비 down)
import multiprocessing as mp
from collections import deque
import pickle

import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docx.oxml.ns import qn

import PBD_p3d as pbd

# 자료형 비교 위해 불러올 것
import matplotlib.pyplot as plt
from decimal import Decimal

#%% test2

#%% Word로 결과 정리

class OutputDocx():

    def __init__(self, bldg_name):
        
        # template 불러와서 Document 생성
        # template = 성능기반 내진설계 보고서
        self.document = docx.Document("template/report_template.docx")
        
        # 동 이름 replace(paragraph level)
        for paragraph in self.document.paragraphs:
            for run in paragraph.runs:        
                if 'building_name' in run.text:
                    run.text = bldg_name
                    
        # 동 이름 replac(table level)
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if 'building_name' in run.text:
                                run.text = bldg_name
                                
    def save_docx(self, result_xlsx_path, output_docx):
        # 결과 저장할 경로
        output_path = os.path.dirname(result_xlsx_path[0])
        # 결과 저장
        self.document.save(os.path.join(output_path, output_docx))
        
#%% Base_SF

    def base_SF_docx(self, base_SF):
        # generator -> list       
        base_SF_list = list(base_SF)

        # 결과를 값과 그래프로 나누기(by data type)
        base_SF_markers = []
        base_SF_values = deque()
        base_SF_plots = deque()
        for i in base_SF_list:
            if isinstance(i, Decimal):
                base_SF_values.append(i)
            elif isinstance(i, plt.Figure):
                base_SF_plots.append(i)
            elif isinstance(i, str):
                base_SF_markers.append(i)

        # Avg. Base Shear 표 작성
        # template의 1,2번 표 불러오기
        base_SF_values_table = self.document.tables[0]
        base_SF_plots_table = self.document.tables[1]
        
        # DE가 있는 경우, DE 값,그래프 채우기
        if 'DE' in base_SF_markers:            
            # 첫번째 표에 avg 값 넣기
            values_row = base_SF_values_table.rows[4]
            values_cell_x = values_row.cells[2]
            values_cell_y = values_row.cells[3]
            values_para_x = values_cell_x.paragraphs[0]
            values_para_y = values_cell_y.paragraphs[0]
            values_para_x.text = f'{base_SF_values.popleft():,} kN' # 1000 자리마다 , 찍기
            values_para_y.text = f'{base_SF_values.popleft():,} kN'
            values_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
            values_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 두번째 표에 그래프 넣기
            memfile = BytesIO()
            memfile2 = BytesIO()
            base_SF_plots.popleft().savefig(memfile)
            base_SF_plots.popleft().savefig(memfile2)
            
            plots_row = base_SF_plots_table.rows[0]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(9))
            plots_run_y.add_picture(memfile2, width=Cm(9))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        if 'MCE' in base_SF_markers:            
            # 첫번째 표에 avg 값 넣기
            values_row = base_SF_values_table.rows[5]
            values_cell_x = values_row.cells[2]
            values_cell_y = values_row.cells[3]
            values_para_x = values_cell_x.paragraphs[0]
            values_para_y = values_cell_y.paragraphs[0]
            values_para_x.text = f'{base_SF_values.popleft():,} kN' # 1000 자리마다 , 찍기
            values_para_y.text = f'{base_SF_values.popleft():,} kN'
            values_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
            values_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 두번째 표에 그래프 넣기
            memfile = BytesIO()
            memfile2 = BytesIO()
            base_SF_plots.popleft().savefig(memfile)
            base_SF_plots.popleft().savefig(memfile2)
            
            plots_row = base_SF_plots_table.rows[3]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(9))
            plots_run_y.add_picture(memfile2, width=Cm(9))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER

#%% IDR

    def IDR_docx(self, IDR):
        # generator -> list       
        IDR_list = list(IDR)

        # 결과를 값과 그래프로 나누기(by data type)
        IDR_markers = []
        IDR_plots = deque()
        for i in IDR_list:
            if isinstance(i, plt.Figure):
                IDR_plots.append(i)
            elif isinstance(i, str):
                IDR_markers.append(i)

        # IDR 표 작성
        # template의 표 불러오기
        IDR_plots_table = self.document.tables[2]

        # DE가 있는 경우, DE 값,그래프 채우기
        if 'DE' in IDR_markers:            
            
            # 표에 그래프 넣기
            memfile = BytesIO()
            memfile2 = BytesIO()
            IDR_plots.popleft().savefig(memfile)
            IDR_plots.popleft().savefig(memfile2)
            
            plots_row = IDR_plots_table.rows[0]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(7))
            plots_run_y.add_picture(memfile2, width=Cm(7))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        if 'MCE' in IDR_markers:            
            
            # 표에 그래프 넣기
            memfile = BytesIO()
            memfile2 = BytesIO()
            IDR_plots.popleft().savefig(memfile)
            IDR_plots.popleft().savefig(memfile2)
            
            plots_row = IDR_plots_table.rows[3]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(7))
            plots_run_y.add_picture(memfile2, width=Cm(7))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER

#%% Wall Axial Strain
            
    def WAS_docx(self, AS):
        
        document = docx.Document("template/report_template.docx")
        # generator -> list       
        AS_list = list(AS)

        # 결과를 값과 그래프로 나누기(by data type)
        AS_markers = []
        AS_values = deque()
        AS_plots = deque()
        for i in AS_list:
            if isinstance(i, pd.DataFrame):
                AS_values.append(i)
            elif isinstance(i, plt.Figure):
                AS_plots.append(i)
            elif isinstance(i, str):
                AS_markers.append(i)

        # Avg. AS 표 작성
        # template의 3,4,5번 표 불러오기
        AS_plots_table = self.document.tables[3]
        AS_values_table = self.document.tables[4]
        AS_values_table_2 = self.document.tables[5]        

        # DE가 있는 경우, DE 값,그래프 채우기
        if 'DE' in AS_markers:            
            
            # 1번 표에 그래프 넣기
            memfile = BytesIO()
            memfile2 = BytesIO()
            AS_plots.popleft().savefig(memfile)
            AS_plots.popleft().savefig(memfile2)
            
            plots_row = AS_plots_table.rows[0]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(8))
            plots_run_y.add_picture(memfile2, width=Cm(8))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 2번 표에 error값 넣기
            # error 개소가 2개소 이상인 경우, table row 늘리기
            errors_df = AS_values.popleft()
            # Axial strain 값 round(소수점 5째자리)
            errors_df.iloc[:,[3,4,5,6]] = errors_df.iloc[:,[3,4,5,6]].round(5)
            errors_df = errors_df.iloc[:,[0,1,2,3,4]]
            if errors_df.shape[0] > 1:
                for i in range(int(errors_df.shape[0] - 1)):
                    AS_values_table.add_row().cells
            
            # row 마다 loop 돌리면서 좌표/error값 입력
            for i in range(errors_df.shape[0]):
                
                values_row = AS_values_table.rows[3 + i]
                
                cell_count = 0
                for values_cell in values_row.cells:
                    values_para = values_cell.paragraphs[0]
                    values_run = values_para.add_run()
                    values_run.text = str(list(errors_df.iloc[:,cell_count])[i])        
                    values_run.font.name = '맑은 고딕'
                    values_run.font.size = Pt(9)
                    values_para.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
                    cell_count += 1
                    
        # MCE가 있는 경우, DE 값,그래프 채우기
        if 'MCE' in AS_markers:            
            
            # 1번 표에 그래프 넣기
            memfile = BytesIO()
            memfile2 = BytesIO()
            AS_plots.popleft().savefig(memfile)
            AS_plots.popleft().savefig(memfile2)
            
            plots_row = AS_plots_table.rows[3]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(8))
            plots_run_y.add_picture(memfile2, width=Cm(8))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 3번 표에 error값 넣기
            # error 개소가 2개소 이상인 경우, table row 늘리기
            errors_df = AS_values.popleft()
            # Axial strain 값 round(소수점 5째자리)
            errors_df.iloc[:,[3,4,5,6]] = errors_df.iloc[:,[3,4,5,6]].round(5)
            errors_df = errors_df.iloc[:,[0,1,2,5,6]]
            if errors_df.shape[0] > 1:
                for i in range(int(errors_df.shape[0] - 1)):
                    AS_values_table_2.add_row().cells
            
            # row 마다 loop 돌리면서 좌표/error값 입력
            for i in range(errors_df.shape[0]):
                
                values_row = AS_values_table_2.rows[3 + i]
                
                cell_count = 0
                for values_cell in values_row.cells:
                    values_para = values_cell.paragraphs[0]
                    values_run = values_para.add_run()
                    values_run.text = str(list(errors_df.iloc[:,cell_count])[i])        
                    values_run.font.name = '맑은 고딕'
                    values_run.font.size = Pt(9)
                    values_para.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
                    cell_count += 1
        
#%% Beam Rotation (DCR)
        
    def BR_docx(self, BR_DCR):
        
        # generator -> list       
        BR_list = list(BR_DCR)

        # 결과를 값과 그래프로 나누기(by data type)
        BR_markers = []
        BR_values = deque()
        BR_plots = deque()
        for i in BR_list:
            if isinstance(i, pd.DataFrame):
                BR_values.append(i)
            elif isinstance(i, plt.Figure):
                BR_plots.append(i)
            elif isinstance(i, str):
                BR_markers.append(i)

        # Avg. AS 표 작성
        # template의 3,4,5번 표 불러오기
        BR_plots_table = self.document.tables[6]
        BR_values_table = self.document.tables[7]
        BR_values_table_2 = self.document.tables[8]

        # DE가 있는 경우, DE 값,그래프 채우기
        if 'DE' in BR_markers:            
            
            # 1번 표에 그래프 넣기
            memfile = BytesIO()
            BR_plots.popleft().savefig(memfile)
            
            plots_row = BR_plots_table.rows[0]
            plots_cell = plots_row.cells[0]
            plots_para = plots_cell.paragraphs[0]  
            plots_run = plots_para.add_run()
            plots_run.add_picture(memfile, width=Cm(8))  
            plots_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 2번 표에 error값 넣기
            # error 개소가 2개소 이상인 경우, table row 늘리기
            errors_df = BR_values.popleft()
            # Axial strain 값 round(소수점 5째자리)
            errors_df.iloc[:,[3,4]] = errors_df.iloc[:,[3,4]].round(5)
            errors_df = errors_df.iloc[:,[1,3,4]]
            if errors_df.shape[0] > 1:
                for i in range(int(errors_df.shape[0] - 1)):
                    BR_values_table.add_row().cells
            
            # row 마다 loop 돌리면서 좌표/error값 입력
            for i in range(errors_df.shape[0]):
                
                values_row = BR_values_table.rows[3 + i]
                
                cell_count = 0
                for values_cell in values_row.cells:
                    values_para = values_cell.paragraphs[0]
                    values_run = values_para.add_run()
                    values_run.text = str(list(errors_df.iloc[:,cell_count])[i])        
                    values_run.font.name = '맑은 고딕'
                    values_run.font.size = Pt(9)
                    values_para.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
                    cell_count += 1
                    
        # MCE가 있는 경우, DE 값,그래프 채우기
        if 'MCE' in BR_markers:            
            
            # 1번 표에 그래프 넣기
            memfile = BytesIO()
            BR_plots.popleft().savefig(memfile)
            
            plots_row = BR_plots_table.rows[0]
            plots_cell = plots_row.cells[1]
            plots_para = plots_cell.paragraphs[0]  
            plots_run = plots_para.add_run()
            plots_run.add_picture(memfile, width=Cm(8))  
            plots_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 2번 표에 error값 넣기
            # error 개소가 2개소 이상인 경우, table row 늘리기
            errors_df = BR_values.popleft()
            # Axial strain 값 round(소수점 5째자리)
            errors_df.iloc[:,[3,4]] = errors_df.iloc[:,[3,4]].round(5)
            errors_df = errors_df.iloc[:,[1,3,4]]
            if errors_df.shape[0] > 1:
                for i in range(int(errors_df.shape[0] - 1)):
                    BR_values_table_2.add_row().cells
            
            # row 마다 loop 돌리면서 좌표/error값 입력
            for i in range(errors_df.shape[0]):
                
                values_row = BR_values_table_2.rows[3 + i]
                
                cell_count = 0
                for values_cell in values_row.cells:
                    values_para = values_cell.paragraphs[0]
                    values_run = values_para.add_run()
                    values_run.text = str(list(errors_df.iloc[:,cell_count])[i])        
                    values_run.font.name = '맑은 고딕'
                    values_run.font.size = Pt(9)
                    values_para.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
                    cell_count += 1

#%% Wall Rotation (DCR)
        
    def WR_docx(self, SWR_DCR):
        # generator -> list       
        WR_list = list(SWR_DCR)

        # 결과를 값과 그래프로 나누기(by data type)
        WR_markers = []
        WR_values = deque()
        WR_plots = deque()
        for i in WR_list:
            if isinstance(i, pd.DataFrame):
                WR_values.append(i)
            elif isinstance(i, plt.Figure):
                WR_plots.append(i)
            elif isinstance(i, str):
                WR_markers.append(i)

        # Avg. AS 표 작성
        # template의 3,4,5번 표 불러오기
        WR_plots_table = self.document.tables[9]
        WR_values_table = self.document.tables[10]
        WR_values_table_2 = self.document.tables[11]

        # DE가 있는 경우, DE 값,그래프 채우기
        if 'DE' in WR_markers:            
            
            # 1번 표에 그래프 넣기
            memfile = BytesIO()
            WR_plots.popleft().savefig(memfile)
            
            plots_row = WR_plots_table.rows[0]
            plots_cell = plots_row.cells[0]
            plots_para = plots_cell.paragraphs[0]  
            plots_run = plots_para.add_run()
            plots_run.add_picture(memfile, width=Cm(8))  
            plots_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 2번 표에 error값 넣기
            # error 개소가 2개소 이상인 경우, table row 늘리기
            errors_df = WR_values.popleft()
            # Axial strain 값 round(소수점 5째자리)
            errors_df.iloc[:,[1,2]] = errors_df.iloc[:,[1,2]].round(5)
            if errors_df.shape[0] > 1:
                for i in range(int(errors_df.shape[0] - 1)):
                    WR_values_table.add_row().cells
            
            # row 마다 loop 돌리면서 좌표/error값 입력
            for i in range(errors_df.shape[0]):
                
                values_row = WR_values_table.rows[3 + i]
                
                cell_count = 0
                for values_cell in values_row.cells:
                    values_para = values_cell.paragraphs[0]
                    values_run = values_para.add_run()
                    values_run.text = str(list(errors_df.iloc[:,cell_count])[i])        
                    values_run.font.name = '맑은 고딕'
                    values_run.font.size = Pt(9)
                    values_para.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
                    cell_count += 1
                    
        # MCE가 있는 경우, DE 값,그래프 채우기
        if 'MCE' in WR_markers:            
            
            # 1번 표에 그래프 넣기
            memfile = BytesIO()
            WR_plots.popleft().savefig(memfile)
            
            plots_row = WR_plots_table.rows[0]
            plots_cell = plots_row.cells[1]
            plots_para = plots_cell.paragraphs[0]  
            plots_run = plots_para.add_run()
            plots_run.add_picture(memfile, width=Cm(8))  
            plots_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 2번 표에 error값 넣기
            # error 개소가 2개소 이상인 경우, table row 늘리기
            errors_df = WR_values.popleft()
            # Axial strain 값 round(소수점 5째자리)
            errors_df.iloc[:,[1,2]] = errors_df.iloc[:,[1,2]].round(5)
            if errors_df.shape[0] > 1:
                for i in range(int(errors_df.shape[0] - 1)):
                    WR_values_table_2.add_row().cells
            
            # row 마다 loop 돌리면서 좌표/error값 입력
            for i in range(errors_df.shape[0]):
                
                values_row = WR_values_table_2.rows[3 + i]
                
                cell_count = 0
                for values_cell in values_row.cells:
                    values_para = values_cell.paragraphs[0]
                    values_run = values_para.add_run()
                    values_run.text = str(list(errors_df.iloc[:,cell_count])[i])        
                    values_run.font.name = '맑은 고딕'
                    values_run.font.size = Pt(9)
                    values_para.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
                    cell_count += 1

#%% Column Rotation (DCR)
        
    def CR_docx(self, CR_DCR):
        
        # generator -> list       
        CR_list = list(CR_DCR)

        # 결과를 값과 그래프로 나누기(by data type)
        CR_markers = []
        CR_values = deque()
        CR_plots = deque()
        for i in CR_list:
            if isinstance(i, pd.DataFrame):
                CR_values.append(i)
            elif isinstance(i, plt.Figure):
                CR_plots.append(i)
            elif isinstance(i, str):
                CR_markers.append(i)

        # Avg. AS 표 작성
        # template의 3,4,5번 표 불러오기
        CR_plots_table = self.document.tables[12]
        CR_values_table = self.document.tables[13]
        CR_values_table_2 = self.document.tables[14]
        CR_values_table_3 = self.document.tables[15]
        CR_values_table_4 = self.document.tables[16]

        # DE가 있는 경우, DE 값,그래프 채우기
        if 'DE' in CR_markers:            
            
            # 1번 표에 그래프 넣기
            memfile = BytesIO()
            memfile2 = BytesIO()
            CR_plots.popleft().savefig(memfile)
            CR_plots.popleft().savefig(memfile2)
            
            plots_row_x = CR_plots_table.rows[0]
            plots_row_y = CR_plots_table.rows[3]
            plots_cell_x = plots_row_x.cells[0]
            plots_cell_y = plots_row_y.cells[0]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(8))
            plots_run_y.add_picture(memfile2, width=Cm(8))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 1,2번 표에 error값 넣기
            errors_df = CR_values.popleft()
            errors_df2 = CR_values.popleft()
            # Column Rotation 값 round(소수점 5째자리)
            errors_df.iloc[:,[3,4]] = errors_df.iloc[:,[3,4]].round(5)
            errors_df2.iloc[:,[3,4]] = errors_df2.iloc[:,[3,4]].round(5)
            errors_df = errors_df.iloc[:,[1,3,4]]
            errors_df2 = errors_df2.iloc[:,[1,3,4]]
            # error 개소가 2개소 이상인 경우, table row 늘리기
            if errors_df.shape[0] > 1:
                for i in range(int(errors_df.shape[0] - 1)):
                    CR_values_table.add_row().cells
            if errors_df2.shape[0] > 1:
                for i in range(int(errors_df2.shape[0] - 1)):
                    CR_values_table_2.add_row().cells
            
            # row 마다 loop 돌리면서 좌표/error값 입력
            for i in range(errors_df.shape[0]):
                
                values_row = CR_values_table.rows[3 + i]                
                cell_count = 0
                for values_cell in values_row.cells:
                    values_para = values_cell.paragraphs[0]
                    values_run = values_para.add_run()
                    values_run.text = str(list(errors_df.iloc[:,cell_count])[i])        
                    values_run.font.name = '맑은 고딕'
                    values_run.font.size = Pt(9)
                    values_para.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
                    cell_count += 1

            for i in range(errors_df2.shape[0]):
                
                values_row = CR_values_table_2.rows[3 + i]                
                cell_count = 0
                for values_cell in values_row.cells:
                    values_para = values_cell.paragraphs[0]
                    values_run = values_para.add_run()
                    values_run.text = str(list(errors_df2.iloc[:,cell_count])[i])        
                    values_run.font.name = '맑은 고딕'
                    values_run.font.size = Pt(9)
                    values_para.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
                    cell_count += 1
                    
        # MCE가 있는 경우, DE 값,그래프 채우기
        if 'MCE' in CR_markers:            
            
            # 1번 표에 그래프 넣기
            memfile = BytesIO()
            memfile2 = BytesIO()
            CR_plots.popleft().savefig(memfile)
            CR_plots.popleft().savefig(memfile2)
            
            plots_row_x = CR_plots_table.rows[0]
            plots_row_y = CR_plots_table.rows[3]
            plots_cell_x = plots_row_x.cells[1]
            plots_cell_y = plots_row_y.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(8))
            plots_run_y.add_picture(memfile2, width=Cm(8))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 3,4번 표에 error값 넣기
            errors_df = CR_values.popleft()
            errors_df2 = CR_values.popleft()
            # Column Rotation 값 round(소수점 5째자리)
            errors_df.iloc[:,[3,4]] = errors_df.iloc[:,[3,4]].round(5)
            errors_df2.iloc[:,[3,4]] = errors_df2.iloc[:,[3,4]].round(5)
            errors_df = errors_df.iloc[:,[1,3,4]]
            errors_df2 = errors_df2.iloc[:,[1,3,4]]
            # error 개소가 2개소 이상인 경우, table row 늘리기
            if errors_df.shape[0] > 1:
                for i in range(int(errors_df.shape[0] - 1)):
                    CR_values_table_3.add_row().cells
            if errors_df2.shape[0] > 1:
                for i in range(int(errors_df2.shape[0] - 1)):
                    CR_values_table_4.add_row().cells
            
            # row 마다 loop 돌리면서 좌표/error값 입력
            for i in range(errors_df.shape[0]):
                
                values_row = CR_values_table_3.rows[3 + i]                
                cell_count = 0
                for values_cell in values_row.cells:
                    values_para = values_cell.paragraphs[0]
                    values_run = values_para.add_run()
                    values_run.text = str(list(errors_df.iloc[:,cell_count])[i])        
                    values_run.font.name = '맑은 고딕'
                    values_run.font.size = Pt(9)
                    values_para.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
                    cell_count += 1

            for i in range(errors_df2.shape[0]):
                
                values_row = CR_values_table_4.rows[3 + i]                
                cell_count = 0
                for values_cell in values_row.cells:
                    values_para = values_cell.paragraphs[0]
                    values_run = values_para.add_run()
                    values_run.text = str(list(errors_df2.iloc[:,cell_count])[i])        
                    values_run.font.name = '맑은 고딕'
                    values_run.font.size = Pt(9)
                    values_para.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
                    cell_count += 1

#%% Wall Shear Force (DCR)
        
    def WSF_docx(self, wall_SF):
        
        # generator -> list       
        SF_list = list(wall_SF)

        # 결과를 값과 그래프로 나누기(by data type)
        SF_markers = []
        SF_plots = deque()
        for i in SF_list:
            if isinstance(i, plt.Figure):
                SF_plots.append(i)
            elif isinstance(i, str):
                SF_markers.append(i)

        # Avg. SF 표 작성
        # template의 3,4,5번 표 불러오기
        SF_plots_table = self.document.tables[17]

        # DE가 있는 경우, DE 값,그래프 채우기
        if 'DE' in SF_markers:            
            
            # 1번 표에 그래프 넣기
            memfile = BytesIO()
            memfile2 = BytesIO()
            SF_plots.popleft().savefig(memfile)
            SF_plots.popleft().savefig(memfile2)
            
            plots_row_x = SF_plots_table.rows[0]
            plots_row_y = SF_plots_table.rows[2]
            plots_cell_x = plots_row_x.cells[0]
            plots_cell_y = plots_row_x.cells[0]
            plots_para_x = plots_cell_x.paragraphs[0]  
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(8))
            plots_run_y.add_picture(memfile2, width=Cm(8))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
        # MCE가 있는 경우, DE 값,그래프 채우기
        if 'MCE' in SF_markers:            
            
            # 1번 표에 그래프 넣기
            memfile = BytesIO()
            memfile2 = BytesIO()
            SF_plots.popleft().savefig(memfile)
            SF_plots.popleft().savefig(memfile2)
            
            plots_row_x = SF_plots_table.rows[0]
            plots_row_y = SF_plots_table.rows[2]
            plots_cell_x = plots_row_x.cells[1]
            plots_cell_y = plots_row_x.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]  
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(8))
            plots_run_y.add_picture(memfile2, width=Cm(8))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER

#%% Column Shear Force (DCR)
        
    def CSF_docx(self, CSF):
        
        # generator -> list       
        SF_list = list(CSF)

        # 결과를 값과 그래프로 나누기(by data type)
        SF_markers = []
        SF_plots = deque()
        for i in SF_list:
            if isinstance(i, plt.Figure):
                SF_plots.append(i)
            elif isinstance(i, str):
                SF_markers.append(i)

        # Avg. SF 표 작성
        # template의 3,4,5번 표 불러오기
        SF_plots_table = self.document.tables[19]

        # DE가 있는 경우, DE 값,그래프 채우기
        if 'DE' in SF_markers:            
            
            # 1번 표에 그래프 넣기
            memfile = BytesIO()
            memfile2 = BytesIO()
            SF_plots.popleft().savefig(memfile)
            SF_plots.popleft().savefig(memfile2)
            
            plots_row_x = SF_plots_table.rows[0]
            plots_row_y = SF_plots_table.rows[3]
            plots_cell_x = plots_row_x.cells[0]
            plots_cell_y = plots_row_x.cells[0]
            plots_para_x = plots_cell_x.paragraphs[0]  
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(8))
            plots_run_y.add_picture(memfile2, width=Cm(8))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
        # MCE가 있는 경우, DE 값,그래프 채우기
        if 'MCE' in SF_markers:            
            
            # 1번 표에 그래프 넣기
            memfile = BytesIO()
            memfile2 = BytesIO()
            SF_plots.popleft().savefig(memfile)
            SF_plots.popleft().savefig(memfile2)
            
            plots_row_x = SF_plots_table.rows[0]
            plots_row_y = SF_plots_table.rows[3]
            plots_cell_x = plots_row_x.cells[1]
            plots_cell_y = plots_row_x.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]  
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(8))
            plots_run_y.add_picture(memfile2, width=Cm(8))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
#%% E.Beam Shear Force (DCR)
        
    def E_BSF_docx(self, E_BSF):
        
        # generator -> list       
        SF_list = list(E_BSF)

        # 결과를 값과 그래프로 나누기(by data type)
        # SF_values = deque()
        SF_plots = deque()
        for i in SF_list:
            if isinstance(i, plt.Figure):
                SF_plots.append(i)

        # E.Beam DCR 도면에 표시한 표 작성
        # template의 3,4,5번 표 불러오기
        SF_plots_table = self.document.tables[12]
        # SF_values_table = self.document.tables[13]

        # 그래프 넣기    
            
        # 1번 표에 그래프 넣기
        memfile = BytesIO()
        memfile2 = BytesIO()
        SF_plots.popleft().savefig(memfile)
        SF_plots.popleft().savefig(memfile2)
        
        plots_row_x = SF_plots_table.rows[0]
        plots_row_y = SF_plots_table.rows[2]
        plots_cell_x = plots_row_x.cells[0]
        plots_cell_y = plots_row_x.cells[0]
        plots_para_x = plots_cell_x.paragraphs[0]  
        plots_para_y = plots_cell_y.paragraphs[0]
        plots_run_x = plots_para_x.add_run()
        plots_run_y = plots_para_y.add_run()
        plots_run_x.add_picture(memfile, width=Cm(8))
        plots_run_y.add_picture(memfile2, width=Cm(8))
        plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
        plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER

#%% Story_SF

    def story_SF_docx(self, story_SF):
        # generator -> list       
        story_SF_list = list(story_SF)
        
        # 결과를 값과 그래프로 나누기(by data type)
        story_SF_markers = []
        story_SF_plots = deque()
        for i in story_SF_list:
            if isinstance(i, plt.Figure):
                story_SF_plots.append(i)
            elif isinstance(i, str):
                story_SF_markers.append(i)
        
        # Avg. Story Shear 표 작성
        # template의 14번표 불러오기
        story_SF_plots_table = self.document.tables[20]
        
        # DE가 있는 경우, DE 값,그래프 채우기
        if 'DE' in story_SF_markers:                        
            # 표에 그래프 넣기
            memfile = BytesIO()
            memfile2 = BytesIO()
            story_SF_plots.popleft().savefig(memfile)
            story_SF_plots.popleft().savefig(memfile2)
            
            plots_row_x = story_SF_plots_table.rows[0]
            plots_row_y = story_SF_plots_table.rows[3]
            plots_cell_x = plots_row_x.cells[0]
            plots_cell_y = plots_row_y.cells[0]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(9))
            plots_run_y.add_picture(memfile2, width=Cm(9))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        if 'MCE' in story_SF_markers:                        
            # 두번째 표에 그래프 넣기
            memfile = BytesIO()
            memfile2 = BytesIO()
            story_SF_plots.popleft().savefig(memfile)
            story_SF_plots.popleft().savefig(memfile2)
            
            plots_row_x = story_SF_plots_table.rows[0]
            plots_row_y = story_SF_plots_table.rows[3]
            plots_cell_x = plots_row_x.cells[1]
            plots_cell_y = plots_row_y.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(9))
            plots_run_y.add_picture(memfile2, width=Cm(9))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER

#%% Base_SF

    def base_SF_docx_test(self, base_SF):
        # memfile.dat 불러오기
        with open('memfile_plot.dat', 'rb') as f1, open('memfile_value.dat', 'rb') as f2, open('memfile_marker.dat', 'rb') as f3:
            base_SF_plots = pickle.load(f1)
            base_SF_values = pickle.load(f2)
            base_SF_markers = pickle.load(f3)
        
        # Avg. Base Shear 표 작성
        # template의 1,2번 표 불러오기
        base_SF_values_table = self.document.tables[0]
        base_SF_plots_table = self.document.tables[1]
        
        # DE가 있는 경우, DE 값,그래프 채우기
        if 'DE' in base_SF_markers:            
            # 첫번째 표에 avg 값 넣기
            values_row = base_SF_values_table.rows[4]
            values_cell_x = values_row.cells[2]
            values_cell_y = values_row.cells[3]
            values_para_x = values_cell_x.paragraphs[0]
            values_para_y = values_cell_y.paragraphs[0]
            values_para_x.text = f'{base_SF_values.popleft():,} kN' # 1000 자리마다 , 찍기
            values_para_y.text = f'{base_SF_values.popleft():,} kN'
            values_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
            values_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 두번째 표에 그래프 넣기
            # ax -> fig 형식으로 만들기
            fig1 = plt.figure(dpi=150)
            with open('memfile_ax1.dat', 'rb') as f:
                ax = pickle.load(f)       
            memfile = BytesIO()
            plt.savefig(memfile)
            memfile.close()
            plt.close()
            
            fig2 = plt.figure(dpi=150)
            with open('memfile_ax2.dat', 'rb') as f:
                ax2 = pickle.load(f)
            memfile = BytesIO()
            plt.savefig(memfile)
            memfile.close()
            plt.close()
            
            # 표에 그래프 넣기
            plots_row = base_SF_plots_table.rows[0]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(9))
            plots_run_y.add_picture(memfile, width=Cm(9))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        if 'MCE' in base_SF_markers:            
            # 첫번째 표에 avg 값 넣기
            values_row = base_SF_values_table.rows[5]
            values_cell_x = values_row.cells[2]
            values_cell_y = values_row.cells[3]
            values_para_x = values_cell_x.paragraphs[0]
            values_para_y = values_cell_y.paragraphs[0]
            values_para_x.text = f'{base_SF_values.popleft():,} kN' # 1000 자리마다 , 찍기
            values_para_y.text = f'{base_SF_values.popleft():,} kN'
            values_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
            values_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 두번째 표에 그래프 넣기
            # ax -> fig 형식으로 만들기
            fig3 = plt.figure(dpi=150)
            with open('memfile_ax3.dat', 'rb') as f:
                ax3 = pickle.load(f)
            memfile = BytesIO()
            plt.savefig(memfile)
            memfile.close()
            plt.close()
            
            fig4 = plt.figure(dpi=150)
            with open('memfile_ax4.dat', 'rb') as f:
                ax4 = pickle.load(f)         
            memfile = BytesIO()
            plt.savefig(memfile)
            memfile.close()
            plt.close()
            
            plots_row = base_SF_plots_table.rows[3]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(9))
            plots_run_y.add_picture(memfile, width=Cm(9))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER

#%% C.Beam SF (엑셀에 자동입력)
        
    def BSF_docx(self, BSF):
        
        pass


#%% C.Beam, G.Column의 Plastic Hinge 찾기 (Rotation값 엑셀에 자동입력)
        
    def plastic_hinge_docx(self, plastic_hinge):
        
        pass

#%% 배근 Wall Shear Force (DCR)
                    
#%% 전체 결과 그래프 그리기
'''
count = 0
for i in plot_list:
    
    memfile = BytesIO()
    i.savefig(memfile)
    
    document_table_faster_para = document_table_faster[count]\
                                    .paragraphs[0]
    document_table_faster_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document_table_faster_run = document_table_faster_para.add_run()
    # document_table_faster_run.add_picture(memfile, height = figsize_y\
    #                                          , width = figsize_x)
    # document_table_faster_run.add_picture(memfile, width=Cm(figsize_x)
    document_table_faster_run.add_picture(memfile, width=Cm(figsize_x))
    document_table_faster_run.add_break(WD_BREAK.PAGE)
    
    memfile.close()    
    count += 1

# Table 스타일  
document_table.style = 'no_borderlines'
document_table.autofit = False
document_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
# 스타일 지정(global)
document_style = document.styles['Normal']
document_style.font.name = '맑은 고딕'
document_style._element.rPr.rFonts\
    .set(qn('w:eastAsia'), '맑은 고딕') # 한글 폰트를 따로 설정해 준다
document_style.font.size = Pt(8) 

# 저장~
document.save(output_path + '\\' + output_docx)
'''
#%% Appendix
'''
# Document 생성
appendix_word = docx.Document()

# Changing the page margins
appendix_word_sections = appendix_word.sections
for section in appendix_word_sections:
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(0.44)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

# 제목
appendix_word_title_para = appendix_word.add_paragraph()
appendix_word_title_run = appendix_word_title_para.add_run(bldg_name+'(Appendix)')
appendix_word_title_run.font.size = Pt(12)
appendix_word_title_run.bold = True

# 표 삽입  # int(-(-x//1)) = math.ceil()
appendix_word_table = appendix_word.add_table(int(-(-len(plot_list)//2)), 2)
appendix_word_table_faster = appendix_word_table._cells


#%% 전체 결과 그래프 그리기

# 지진파별 그래프
count = 1

for i in input_wall_name:  
    
    shear_force_major_max_temp = shear_force_major_max[(shear_force_major_max.index.str.contains(i + '_'))\
                                                       & (shear_force_major_max.index.str[0] == i[0])]

    story_temp = shear_force_major_max_temp.index.tolist()
    story_temp = pd.Series(list(map(lambda x: x.rsplit('_', 1)[1], story_temp))).tolist()
    
    ### DE
    memfile = BytesIO()
    plt.figure(count, dpi=150, figsize=(5, 4.8))
    # plt.xlim(0, story_shear_xlim)
    
    # 지진파별 plot
    for j in range(14):
        plt.plot(shear_force_major_max_temp.iloc[:,j], range(shear_force_major_max_temp.shape[0]), label=seismic_load_name_list[j], linewidth=0.7)
        
    # 평균 plot
    plt.plot(shear_force_major_max_temp.iloc[:,0:14].mean(axis=1), range(shear_force_major_max_temp.shape[0]), color='k', label='Average', linewidth=2)
    
    plt.yticks(range(shear_force_major_max_temp.shape[0])[::story_shear_yticks], story_temp[::story_shear_yticks])
    # plt.xticks(range(14), range(1,15))
    
    # 기타
    plt.grid(linestyle='-.')
    plt.xlabel('Story Shear(kN)')
    plt.ylabel('Story')
    plt.legend(loc=1, fontsize=8)
    plt.title('{}'.format(i.split('_')[0]) + ' (DE)')
    
    plt.tight_layout()
    plt.savefig(memfile)
    plt.close()
    
    SF_word_table_faster_para = SF_word_table_faster[count-1].paragraphs[0]
    SF_word_table_faster_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    SF_word_table_faster_run = SF_word_table_faster_para.add_run()
    # SF_word_table_faster_run.add_picture(memfile, height = figsize_y, width = figsize_x)
    SF_word_table_faster_run.add_picture(memfile, width=Cm(figsize_x))

    memfile.close()
    count += 1
    
    ### MCE
    memfile2 = BytesIO()
    plt.figure(count, dpi=150, figsize=(5, 4.8))
    # plt.xlim(0, story_shear_xlim)
    
    # 지진파별 plot
    for j in range(14):
        plt.plot(shear_force_major_max_temp.iloc[:,j+14], range(shear_force_major_max_temp.shape[0]), label=seismic_load_name_list[j+14], linewidth=0.7)
        
    # 평균 plot
    plt.plot(shear_force_major_max_temp.iloc[:,14:28].mean(axis=1), range(shear_force_major_max_temp.shape[0]), color='k', label='Average', linewidth=2)
    
    plt.yticks(range(shear_force_major_max_temp.shape[0])[::story_shear_yticks], story_temp[::story_shear_yticks])
    # plt.xticks(range(14), range(1,15))
    
    # 기타
    plt.grid(linestyle='-.')
    plt.xlabel('Story Shear(kN)')
    plt.ylabel('Story')
    plt.legend(loc=1, fontsize=8)
    plt.title('{}'.format(i.split('_')[0]) + ' (MCE)')
    
    plt.tight_layout()
    plt.savefig(memfile2)
    plt.close()
    
    SF_word_table_faster_para2 = SF_word_table_faster[count-1].paragraphs[0]
    SF_word_table_faster_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    SF_word_table_faster_run2 = SF_word_table_faster_para2.add_run()
    # SF_word_table_faster_run2.add_picture(memfile2, height = figsize_y, width = figsize_x)
    SF_word_table_faster_run2.add_picture(memfile2, width=Cm(figsize_x))
        
    memfile2.close()
    count += 1

# Table 스타일  
appendix_word_table.style = 'Table Grid'
appendix_word_table.autofit = False
appendix_word_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 스타일 지정(global)
appendix_word_style = appendix_word.styles['Normal']
appendix_word_style.font.name = '맑은 고딕'
appendix_word_style._element.rPr.rFonts\
    .set(qn('w:eastAsia'), '맑은 고딕') # 한글 폰트를 따로 설정해 준다
appendix_word_style.font.size = Pt(8) 
        
# 저장~
appendix_word.save(output_path + '\\' + appendix_docx)

'''

