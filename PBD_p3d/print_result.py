#%% Import
import os
import pandas as pd
import numpy as np
import time
from io import BytesIO # 파일처럼 취급되는 문자열 객체 생성(메모리 낭비 down)
# import multiprocessing as mp
# from collections import deque
# import pickle
import win32com.client
import pythoncom
from PyPDF2 import PdfMerger, PdfFileReader


import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docx.oxml.ns import qn

import PBD_p3d as pbd

# 자료형 비교 위해 불러올 것
import matplotlib.pyplot as plt
from decimal import Decimal, ROUND_UP

from PyQt5.QtCore import QObject

import matplotlib.pyplot as plt
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg
from matplotlib.backends.backend_qt5agg import FigureCanvas as FigureCanvas
from matplotlib.backends.backend_qt5agg import NavigationToolbar2QT as NavigationToolbar
from matplotlib.figure import Figure

#%% Matplotlib으로 결과 보여주기
class ShowResult(FigureCanvasQTAgg, QObject):
    
    def __init__(self, parent=None, width=5, height=4):
        self.fig = Figure(figsize=(width, height), layout='tight')
        self.axes = self.fig.add_subplot(111)
        FigureCanvasQTAgg.__init__(self, self.fig)
        FigureCanvasQTAgg.setMinimumSize(self, self.size())
        # super(ShowResult, self).__init__(self.fig)

#%% Word로 결과 정리

class PrintResult(QObject):

    def __init__(self, input_xlsx_path, result_xlsx_path, bldg_name, story_gap, max_shear):
        super().__init__()
        
        ### 변수 정리(str -> int)
        self.story_gap = int(story_gap)
        self.max_shear = int(max_shear)
        self.bldg_name = bldg_name

        ### Data Conversion Sheets 불러오기
        self.input_xlsx_path = input_xlsx_path

        story_info = pd.DataFrame()
        rebar_info = pd.DataFrame()
        c_beam_deform_cap = pd.DataFrame()
        wall_deform_cap = pd.DataFrame()
        e_column_info = pd.DataFrame()
        e_beam_info = pd.DataFrame()
        wall_info = pd.DataFrame()

        # 시트 불러오기
        input_data_raw = pd.ExcelFile(input_xlsx_path)
        input_data_sheets = pd.read_excel(input_data_raw, ['ETC', 'Story Data', 'Output_C.Beam Properties'
                                                           , 'Output_E.Column Properties', 'Output_E.Beam Properties'
                                                           , 'Output_Wall Properties', 'Results_Wall'], skiprows=3)
        input_data_raw.close()

        # Story Data 추출
        story_info = input_data_sheets['Story Data'].iloc[:,[0,1,2]]
        # Rebar Data 추출
        rebar_info = input_data_sheets['ETC'].iloc[:,[0,3,4]]
        # Deformation Capacity 추출
        c_beam_deform_cap = input_data_sheets['Output_C.Beam Properties'].iloc[:,[0,80,81]]
        wall_deform_cap = input_data_sheets['Results_Wall'].iloc[:,[0,11,12,13,14,48,49,54,55]]
        # E.Column Data 추출
        e_column_info = input_data_sheets['Output_E.Column Properties'].iloc[:,0:17]
        # E.Beam Data 추출
        e_beam_info = input_data_sheets['Output_E.Beam Properties'].iloc[:,0]
        # Wall Data 추출
        wall_info = input_data_sheets['Output_Wall Properties'].iloc[:,0:10]

        # 변수 정리(다른 메소드에서도 사용하기 위함)
        self.story_info = story_info
        self.story_info.columns = ['Index', 'Story Name', 'Height(mm)']
        self.rebar_info = rebar_info
        self.rebar_info.columns = ['Type', '일반용', '내진용']
        self.story_name = self.story_info.loc[:, 'Story Name']
        self.c_beam_deform_cap = c_beam_deform_cap
        self.c_beam_deform_cap.columns = ['Name', 'LS', 'CP']
        self.wall_deform_cap = wall_deform_cap
        self.wall_deform_cap.columns = ['Name', 'Vu_DE_H1', 'Vu_DE_H2', 'Vu_MCE_H1', 'Vu_MCE_H2'
                                        , 'LS(H1)', 'LS(H2)', 'CP(H1)', 'CP(H2)']
        self.e_column_info = e_column_info
        self.e_beam_info = e_beam_info
        self.e_beam_info.name = 'Name'
        self.wall_info = wall_info
        self.wall_info.columns = ['Name', 'Length(mm)', 'Thickness(mm)', 'Concrete Grade', 'Rebar Type', 'V.Rebar Type'
                                  , 'V.Rebar Spacing(mm)', 'V.Rebar EA', 'H.Rebar Type', 'H.Rebar Spacing(mm)']

        ### Analysis Result Sheets 불러오기
        self.result_xlsx_path = result_xlsx_path
        to_load_list = result_xlsx_path
        
        section_data = pd.DataFrame()
        drift_data = pd.DataFrame()
        beam_rot_data = pd.DataFrame()
        wall_rot_data = pd.DataFrame()
        wall_AS_data = pd.DataFrame()
        frame_SF_data = pd.DataFrame()

        # 시트 불러오기
        for i in to_load_list:
            result_data_raw = pd.ExcelFile(i)
            result_data_sheets = pd.read_excel(result_data_raw, ['Structure Section Forces', 'Drift Output'
                                                                 , 'Frame Results - Bending Deform', 'Node Coordinate Data'
                                                                 , 'Element Data - Frame Types', 'Element Data - Shear Wall'
                                                                 , 'Gage Data - Bar Type', 'Gage Data - Wall Type'
                                                                 , 'Gage Results - Bar Type', 'Gage Results - Wall Type'
                                                                 , 'Frame Results - End Forces'], skiprows=[0,2])
            result_data_raw.close()

            # Shear Force 데이터 추출
            column_to_slice_section = ['StrucSec Name', 'Load Case', 'Step Type', 'FH1', 'FH2', 'FV']
            section_data_temp = result_data_sheets['Structure Section Forces'].loc[:,column_to_slice_section]
            section_data = pd.concat([section_data, section_data_temp])
            # Drift 데이터 추출
            drift_data_temp = result_data_sheets['Drift Output'].iloc[:,[0,1,3,5,6]]
            drift_data = pd.concat([drift_data, drift_data_temp])
            # Beam Rotation 데이터 추출
            beam_rot_data_temp = result_data_sheets['Frame Results - Bending Deform'].iloc[:,[0,2,5,7,10,13,14]]
            beam_rot_data = pd.concat([beam_rot_data, beam_rot_data_temp])
            # Wall Rotation 데이터 추출
            column_to_slice_wall_rot = ['Group Name', 'Element Name', 'Load Case', 'Step Type', 'Rotation', 'Performance Level']
            wall_rot_data_temp = result_data_sheets['Gage Results - Wall Type'].loc[:,column_to_slice_wall_rot]
            wall_rot_data = pd.concat([wall_rot_data, wall_rot_data_temp])
            # Wall Axial Strain 데이터 추출
            wall_AS_data_temp = result_data_sheets['Gage Results - Bar Type'].iloc[:,[0,2,5,7,8,9]]
            wall_AS_data = pd.concat([wall_AS_data, wall_AS_data_temp])
            # Beam Shear Force 데이터 추출
            frame_SF_data_temp = result_data_sheets['Frame Results - End Forces'].iloc[:,[0,2,5,7,8,10,11,12,15,16,17,18]]
            frame_SF_data = pd.concat([frame_SF_data, frame_SF_data_temp])

        # Node 데이터 추출
        node_data = result_data_sheets['Node Coordinate Data'].iloc[:,[1,2,3,4]]
        # Element 데이터 추출
        element_data = result_data_sheets['Element Data - Frame Types'].iloc[:,[0,2,5,7,9]]
        # Wall Element 데이터 추출
        wall_element_data = result_data_sheets['Element Data - Shear Wall'].iloc[:,[2,5,7,9,11,13]]
        # Axial Strain Gage 정보 추출
        wall_AS_gage_data = result_data_sheets['Gage Data - Bar Type'].iloc[:,[0,2,7,9]]
        # Rotation Gage 정보 추출
        wall_rot_gage_data = result_data_sheets['Gage Data - Wall Type'].iloc[:,[2,7,9,11,13]]

        # 변수 정리(다른 메소드에서도 사용하기 위함)
        self.section_data = section_data
        self.section_data.columns = ['Name', 'Load Case', 'Step Type', 'H1(kN)', 'H2(kN)', 'V(kN)']        
        self.drift_data = drift_data
        self.beam_rot_data = beam_rot_data
        self.beam_rot_data.columns = ['Group Name', 'Element Name', 'Load Case', 'Step Type', 'Distance from I-End', 'H2 Rotation(rad)', 'H3 Rotation(rad)']
        self.wall_rot_data = wall_rot_data
        self.node_data = node_data
        self.element_data = element_data
        self.wall_element_data = wall_element_data
        self.wall_AS_data = wall_AS_data
        self.wall_AS_gage_data = wall_AS_gage_data
        self.frame_SF_data = frame_SF_data
        self.wall_rot_gage_data = wall_rot_gage_data

        ### 지진파 이름 list 만들기
        load_name_list = []
        for i in self.section_data['Load Case'].drop_duplicates():
            new_i = i.split('+')[1]
            new_i = new_i.strip()
            load_name_list.append(new_i)
        
        self.gravity_load_name = [x for x in load_name_list if ('DE' not in x) and ('MCE' not in x)]
        self.seismic_load_name_list = [x for x in load_name_list if ('DE' in x) or ('MCE' in x)]
        
        self.DE_load_name_list = [x for x in load_name_list if 'DE' in x] # base shear로 사용할 지진파 개수 산정을 위함
        self.MCE_load_name_list = [x for x in load_name_list if 'MCE' in x]

        self.seismic_load_name_list.sort()
        self.DE_load_name_list.sort()
        self.MCE_load_name_list.sort()

        # Marker 생성
        self.markers = []
        if len(self.DE_load_name_list) != 0:
            self.markers.append('DE')
        if len(self.MCE_load_name_list) != 0:
            self.markers.append('MCE')

    #%% Base Shear Force
    def base_SF(self):
        ''' 

        Perform-3D 해석 결과에서 각 지진파에 대한 Base층의 전단력을 막대그래프 형식으로 출력. (kN)
        
        Parameters
        ----------
        result_path : str
                    Perform-3D에서 나온 해석 파일의 경로.
                    
        result_xlsx : str, optional, default='Analysis Result'
                    Perform-3D에서 나온 해석 파일의 이름. 해당 파일 이름이 포함된 파일들을 모두 불러온다.
                    
        ylim : int, optional, default=70000
            그래프의 y축 limit 값. y축 limit 안의 값만 표기되므로, limit를 넘어가는 값을 확인하고 싶을 시에는 ylim 값을 더 크게 설정하면 된다.
        
        Returns
        -------
        '''
    #%% Base 전단력 추출
        base_shear_data = self.section_data[self.section_data['Name'].str.contains('base', case=False)]        
        base_shear_data.reset_index(inplace=True, drop=True)

    #%% 데이터 Grouping
        shear_force_H1_data_grouped = pd.DataFrame()
        shear_force_H2_data_grouped = pd.DataFrame()
        
        for load_name in self.seismic_load_name_list:
            shear_force_H1_data_grouped['{}_H1_max'.format(load_name)] = base_shear_data[(base_shear_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                        (base_shear_data['Step Type'] == 'Max')]['H1(kN)'].values
                
            shear_force_H1_data_grouped['{}_H1_min'.format(load_name)] = base_shear_data[(base_shear_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                        (base_shear_data['Step Type'] == 'Min')]['H1(kN)'].values
        
        for load_name in self.seismic_load_name_list:
            shear_force_H2_data_grouped['{}_H2_max'.format(load_name)] = base_shear_data[(base_shear_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                        (base_shear_data['Step Type'] == 'Max')]['H2(kN)'].values
                
            shear_force_H2_data_grouped['{}_H2_min'.format(load_name)] = base_shear_data[(base_shear_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                        (base_shear_data['Step Type'] == 'Min')]['H2(kN)'].values   
        
        # all 절대값
        shear_force_H1_abs = shear_force_H1_data_grouped.abs()
        shear_force_H2_abs = shear_force_H2_data_grouped.abs()
        
        # Min, Max 중 최대값
        shear_force_H1_max = shear_force_H1_abs.groupby([[i//2 for i in range(0,len(self.seismic_load_name_list)*2)]], axis=1).max()
        shear_force_H2_max = shear_force_H2_abs.groupby([[i//2 for i in range(0,len(self.seismic_load_name_list)*2)]], axis=1).max()
        
        shear_force_H1_max.columns = self.seismic_load_name_list
        shear_force_H2_max.columns = self.seismic_load_name_list
        
        shear_force_H1_max.index = base_shear_data['Name'].drop_duplicates()
        shear_force_H2_max.index = base_shear_data['Name'].drop_duplicates()
        
        #%% Base Shear 결과 출력
        base_shear_H1 = shear_force_H1_max.copy()
        base_shear_H2 = shear_force_H2_max.copy()
        
        result = [base_shear_H1, base_shear_H2, self.DE_load_name_list, self.MCE_load_name_list, self.markers]

        return result

    #%% Story Shear Force
    def story_SF(self):
        ''' 

        Perform-3D 해석 결과에서 각 지진파에 대한 각 층의 전단력을 그래프로 출력(kN).
        
        Parameters
        ----------
        input_path : str
                    Data Conversion 엑셀 파일의 경로.
                    
        input_xlsx : str
                    Data Conversion 엑셀 파일의 이름. result_xlsx와는 달리 확장자명(.xlsx)까지 기입해줘야한다. 하나의 파일만 불러온다.
                    
        result_path : str
                    Perform-3D에서 나온 해석 파일의 경로.
                    
        result_xlsx : str, optional, default='Analysis Result'
                    Perform-3D에서 나온 해석 파일의 이름. 해당 파일 이름이 포함된 파일들을 모두 불러온다.
                    
        yticks : int, optional, default=2
                그래프의 y축 눈금 간격(층간격). 층이 너무 높으면 y축에 너무 많은 층이 표기되기 때문에, 층간격을 조절해서 정돈된 그래프를 표기할 수 있다.
        
        xlim : int, optional, default=70000
            그래프의 x축 limit 값. x축 limit 안의 값만 표기되므로, limit를 넘어가는 값을 확인하고 싶을 시에는 더 큰 xlim 값을 사용하면 된다.
        
        Returns
        -------
        '''    
    #%% 필요없는 전단력 제거(층전단력)
        story_shear_data = self.section_data[self.section_data['Name'].str.count('_') != 2] # underbar가 두개 들어간 행들은 제거        
        story_shear_data.reset_index(inplace=True, drop=True)
        
        # 이름에서 _shear 제거
        story_shear_data['Name'] = story_shear_data['Name'].str.rstrip('_Shear')
                
    #%% 데이터 Grouping
        shear_force_H1_data_grouped = pd.DataFrame()
        shear_force_H2_data_grouped = pd.DataFrame()
        
        for load_name in self.seismic_load_name_list:
            shear_force_H1_data_grouped['{}_H1_max'.format(load_name)] = story_shear_data[(story_shear_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                        (story_shear_data['Step Type'] == 'Max')]['H1(kN)'].values
                
            shear_force_H1_data_grouped['{}_H1_min'.format(load_name)] = story_shear_data[(story_shear_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                        (story_shear_data['Step Type'] == 'Min')]['H1(kN)'].values
        
        for load_name in self.seismic_load_name_list:
            shear_force_H2_data_grouped['{}_H2_max'.format(load_name)] = story_shear_data[(story_shear_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                        (story_shear_data['Step Type'] == 'Max')]['H2(kN)'].values
                
            shear_force_H2_data_grouped['{}_H2_min'.format(load_name)] = story_shear_data[(story_shear_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                        (story_shear_data['Step Type'] == 'Min')]['H2(kN)'].values   
        
        # all 절대값
        shear_force_H1_abs = shear_force_H1_data_grouped.abs()
        shear_force_H2_abs = shear_force_H2_data_grouped.abs()
        
        # Min, Max 중 최대값
        shear_force_H1_max = shear_force_H1_abs.groupby([[i//2 for i in range(0,len(self.seismic_load_name_list)*2)]], axis=1).max()
        shear_force_H2_max = shear_force_H2_abs.groupby([[i//2 for i in range(0,len(self.seismic_load_name_list)*2)]], axis=1).max()
        
        shear_force_H1_max.columns = self.seismic_load_name_list
        shear_force_H2_max.columns = self.seismic_load_name_list
        
        shear_force_H1_max.index = story_shear_data['Name'].drop_duplicates()
        shear_force_H2_max.index = story_shear_data['Name'].drop_duplicates()

    #%% Story Shear 결과 출력

        result = [shear_force_H1_max, shear_force_H2_max, self.DE_load_name_list, self.MCE_load_name_list, self.markers]

        return result

    #%% Interstory Drift Ratio
    def IDR(self):   
        ''' 

        Perform-3D 해석 결과에서 각 지진파에 대한 층간변위비를 그래프로 출력.  
        
        Parameters
        ----------                  
        input_path : str
                    Data Conversion 엑셀 파일의 경로.
                    
        input_xlsx : str
                    Data Conversion 엑셀 파일의 이름. result_xlsx와는 달리 확장자명(.xlsx)까지 기입해줘야한다. 하나의 파일만 불러온다.
        
        result_path : str
                    Perform-3D에서 나온 해석 파일의 경로.
                    
        result_xlsx : str, optional, default='Analysis Result'
                    Perform-3D에서 나온 해석 파일의 이름. 해당 파일 이름이 포함된 파일들을 모두 불러온다.
                    
        cri_DE : float, optional, default=0.015
                LS(인명보호)를 만족하는 층간변위비 허용기준.
                
        cri_MCE : float, optional, default=0.02
                CP(붕괴방지)를 만족하는 층간변위비 허용기준.
                
        yticks : int, optional, default=2
                그래프의 y축 눈금 간격(층간격). 층이 너무 높으면 y축에 너무 많은 층이 표기되기 때문에, 층간격을 조절해서 정돈된 그래프를 표기할 수 있다.
        
        Yields
        -------
        fig1 : matplotlib.pyplot.figure or None
            DE(설계지진) 발생 시 x방향 층간변위비 그래프
        
        fig2 : matplotlib.pyplot.figure or None
            DE(설계지진) 발생 시 y방향 층간변위비 그래프
        
        fig3 : matplotlib.pyplot.figure or None
            MCE(최대고려지진) 발생 시 x방향 층간변위비 그래프
        
        fig4 : matplotlib.pyplot.figure or None
            MCE(최대고려지진) 발생 시 y방향 층간변위비 그래프
        
        Raises
        -------
        
        References
        -------
        [1] "철근콘크리트 건축구조물의 성능기반 내진설계 지침", 대한건축학회, p.103, 2021
        
        '''    
    #%% IDR 결과 파일 불러오기     
        IDR_result_data = self.drift_data.sort_values(by=['Load Case', 'Drift ID', 'Step Type']) # 지진파 순서가 섞여있을 때 sort
        
    #%% Drift Name에서 story, direction 뽑아내기
        drift_name = IDR_result_data['Drift Name']
        
        story = []
        direction = []
        position = []
        for i in drift_name:
            i = i.strip()  # drift_name 앞뒤에 있는 blank 제거
        
            if i.count('_') == 2:
                story.append(i.split('_')[0])
                direction.append(i.split('_')[-1])
                position.append(i.split('_')[1].split('_')[0])
            else:
                story.append(None)
                direction.append(None)
        
        # Load Case에서 지진파 이름만 뽑아서 다시 naming
        load_striped = []        
        for i in IDR_result_data['Load Case']:
            load_striped.append(i.strip().split(' ')[-1])
            
        IDR_result_data['Load Case'] = load_striped
            
        
        IDR_result_data.reset_index(inplace=True, drop=True)
        IDR_result_data = pd.concat([pd.Series(story, name='Name'),\
                                    pd.Series(direction, name='Direction'),\
                                    pd.Series(position, name='Position'), IDR_result_data], axis=1)

    #%% IDR값(방향에 따른)
        ### 지진파별 평균
        # 각 지진파들로 변수 생성 후, 값 대입
        for load_name in self.seismic_load_name_list:
            globals()['IDR_x_max_{}_avg'.format(load_name)] = IDR_result_data[(IDR_result_data['Load Case'] == '{}'.format(load_name)) &\
                                                                        (IDR_result_data['Direction'] == 'X') &\
                                                                        (IDR_result_data['Step Type'] == 'Max')].groupby(['Name', 'Position'])['Drift']\
                                                                        .agg(**{'X Max avg':'mean'}).groupby('Name').max()
            
            globals()['IDR_x_min_{}_avg'.format(load_name)] = IDR_result_data[(IDR_result_data['Load Case'] == '{}'.format(load_name)) &\
                                                                        (IDR_result_data['Direction'] == 'X') &\
                                                                        (IDR_result_data['Step Type'] == 'Min')].groupby(['Name'])['Drift']\
                                                                        .agg(**{'X Min avg':'mean'}).groupby('Name').min()
                
            globals()['IDR_y_max_{}_avg'.format(load_name)] = IDR_result_data[(IDR_result_data['Load Case'] == '{}'.format(load_name)) &\
                                                                        (IDR_result_data['Direction'] == 'Y') &\
                                                                        (IDR_result_data['Step Type'] == 'Max')].groupby(['Name'])['Drift']\
                                                                        .agg(**{'Y Max avg':'mean'}).groupby('Name').max()
            
            globals()['IDR_y_min_{}_avg'.format(load_name)] = IDR_result_data[(IDR_result_data['Load Case'] == '{}'.format(load_name)) &\
                                                                        (IDR_result_data['Direction'] == 'Y') &\
                                                                        (IDR_result_data['Step Type'] == 'Min')].groupby(['Name'])['Drift']\
                                                                        .agg(**{'Y Min avg':'mean'}).groupby('Name').min()
                
            globals()['IDR_x_max_{}_avg'.format(load_name)].reset_index(inplace=True)
            globals()['IDR_x_min_{}_avg'.format(load_name)].reset_index(inplace=True)
            globals()['IDR_y_max_{}_avg'.format(load_name)].reset_index(inplace=True)
            globals()['IDR_y_min_{}_avg'.format(load_name)].reset_index(inplace=True)
            
        # Story 정렬하기
        story_name_window = IDR_result_data['Name'].drop_duplicates()
        story_name_window_reordered = [x for x in self.story_name[::-1].tolist() \
                                        if x in story_name_window.tolist()]  # story name를 reference로 해서 정렬
        
        # 정렬된 Story에 따라 IDR값도 정렬 + result에 append
        result = []
        for load_name in self.seismic_load_name_list:   
            globals()['IDR_x_max_{}_avg'.format(load_name)]['Name'] = pd.Categorical(globals()['IDR_x_max_{}_avg'.format(load_name)]['Name'], self.story_name[::-1])
            globals()['IDR_x_max_{}_avg'.format(load_name)].sort_values('Name', inplace=True)
            globals()['IDR_x_max_{}_avg'.format(load_name)].reset_index(inplace=True, drop=True)
            
            globals()['IDR_x_min_{}_avg'.format(load_name)]['Name'] = pd.Categorical(globals()['IDR_x_min_{}_avg'.format(load_name)]['Name'], self.story_name[::-1])
            globals()['IDR_x_min_{}_avg'.format(load_name)].sort_values('Name', inplace=True)
            globals()['IDR_x_min_{}_avg'.format(load_name)].reset_index(inplace=True, drop=True)
            
            globals()['IDR_y_max_{}_avg'.format(load_name)]['Name'] = pd.Categorical(globals()['IDR_y_max_{}_avg'.format(load_name)]['Name'], self.story_name[::-1])
            globals()['IDR_y_max_{}_avg'.format(load_name)].sort_values('Name', inplace=True)
            globals()['IDR_y_max_{}_avg'.format(load_name)].reset_index(inplace=True, drop=True)
            
            globals()['IDR_y_min_{}_avg'.format(load_name)]['Name'] = pd.Categorical(globals()['IDR_y_min_{}_avg'.format(load_name)]['Name'], self.story_name[::-1])
            globals()['IDR_y_min_{}_avg'.format(load_name)].sort_values('Name', inplace=True)
            globals()['IDR_y_min_{}_avg'.format(load_name)].reset_index(inplace=True, drop=True)

            result.append(globals()['IDR_x_max_{}_avg'.format(load_name)])
            result.append(globals()['IDR_x_min_{}_avg'.format(load_name)])
            result.append(globals()['IDR_y_max_{}_avg'.format(load_name)])
            result.append(globals()['IDR_y_min_{}_avg'.format(load_name)])
            
    #%% IDR값(방향에 따른) 전체 평균 (여기부터 2023.03.20 수정)
        result_avg = []
        if len(self.DE_load_name_list) != 0:
                
            IDR_x_max_DE_total = pd.concat([globals()['IDR_x_max_{}_avg'.format(x)].iloc[:,-1] for x in self.DE_load_name_list], axis=1)
            IDR_x_min_DE_total = pd.concat([globals()['IDR_x_min_{}_avg'.format(x)].iloc[:,-1] for x in self.DE_load_name_list], axis=1)
            IDR_y_max_DE_total = pd.concat([globals()['IDR_y_max_{}_avg'.format(x)].iloc[:,-1] for x in self.DE_load_name_list], axis=1)
            IDR_y_min_DE_total = pd.concat([globals()['IDR_y_min_{}_avg'.format(x)].iloc[:,-1] for x in self.DE_load_name_list], axis=1)
            
            IDR_x_max_DE_avg = IDR_x_max_DE_total.mean(axis=1)
            IDR_x_min_DE_avg = IDR_x_min_DE_total.mean(axis=1)
            IDR_y_max_DE_avg = IDR_y_max_DE_total.mean(axis=1)
            IDR_y_min_DE_avg = IDR_y_min_DE_total.mean(axis=1)

            # x,y 방향 min,max 값들을 하나의 dataframe으로 합치기
            IDR_DE_avg = pd.concat([IDR_x_max_DE_avg, IDR_x_min_DE_avg, IDR_y_max_DE_avg, IDR_y_min_DE_avg], axis=1)
            result_avg.append(IDR_DE_avg)
        
        if len(self.MCE_load_name_list) != 0:
            
            IDR_x_max_MCE_total = pd.concat([globals()['IDR_x_max_{}_avg'.format(x)].iloc[:,-1] for x in self.MCE_load_name_list], axis=1)
            IDR_x_min_MCE_total = pd.concat([globals()['IDR_x_min_{}_avg'.format(x)].iloc[:,-1] for x in self.MCE_load_name_list], axis=1)
            IDR_y_max_MCE_total = pd.concat([globals()['IDR_y_max_{}_avg'.format(x)].iloc[:,-1] for x in self.MCE_load_name_list], axis=1)
            IDR_y_min_MCE_total = pd.concat([globals()['IDR_y_min_{}_avg'.format(x)].iloc[:,-1] for x in self.MCE_load_name_list], axis=1)
            
            IDR_x_max_MCE_avg = IDR_x_max_MCE_total.mean(axis=1)
            IDR_x_min_MCE_avg = IDR_x_min_MCE_total.mean(axis=1)
            IDR_y_max_MCE_avg = IDR_y_max_MCE_total.mean(axis=1)
            IDR_y_min_MCE_avg = IDR_y_min_MCE_total.mean(axis=1)

            IDR_MCE_avg = pd.concat([IDR_x_max_MCE_avg, IDR_x_min_MCE_avg, IDR_y_max_MCE_avg, IDR_y_min_MCE_avg], axis=1)
            result_avg.append(IDR_MCE_avg)
        
    #%% IDR 결과 출력
        return result, result_avg, story_name_window_reordered
    
    #%% Beam Rotation
    def BR(self, c_beam_group='C.Beam'):      

        # Beam Rotation, Node, Element 데이터 불러오기
        beam_rot_data = self.beam_rot_data.copy()
        node_data = self.node_data.iloc[:,[0,3]]
        element_data = self.element_data.iloc[:,[1,2,3]]

        # temporary ((L), (R) 등 지우기)
        element_data.loc[:, 'Property Name'] = element_data.loc[:, 'Property Name'].str.split('(').str[0]
        
        # 필요없는 부재 빼기, 필요한 부재만 추출
        beam_rot_data = beam_rot_data[beam_rot_data['Group Name'] == c_beam_group]
        beam_rot_data = beam_rot_data[beam_rot_data['Distance from I-End'] == 0]
        
    #%% Analysis Result에 Element, Node 정보 매칭
        
        element_data = element_data.drop_duplicates()
        node_data = node_data.drop_duplicates()
        
        beam_rot_data = pd.merge(beam_rot_data, element_data, how='left')
        beam_rot_data = pd.merge(beam_rot_data, node_data, how='left', left_on='I-Node ID', right_on='Node ID')
        
        beam_rot_data = beam_rot_data[beam_rot_data['Property Name'].notna()]
        
        beam_rot_data.reset_index(inplace=True, drop=True)
        
    #%% beam_rot_data의 값 수정(H1, H2 방향 중 major한 방향의 rotation값만 추출, 그리고 2배)
        major_rot = []
        for i, j in zip(beam_rot_data['H2 Rotation(rad)'], beam_rot_data['H3 Rotation(rad)']):
            if abs(i) >= abs(j):
                major_rot.append(i)
            else: major_rot.append(j)
        
        beam_rot_data['Major Rotation(rad)'] = major_rot
        
        # 필요한 정보들만 다시 모아서 new dataframe
        beam_rot_data = beam_rot_data.iloc[:, [0,1,7,10,2,3,11]]
        
    #%% 성능기준(LS, CP) 정리해서 merge
        
        beam_rot_data = pd.merge(beam_rot_data, self.c_beam_deform_cap, how='left', left_on='Property Name', right_on='Name')
        
        beam_rot_data['DE Rotation(rad)'] = beam_rot_data['Major Rotation(rad)'].abs() / beam_rot_data['LS']
        beam_rot_data['MCE Rotation(rad)'] = beam_rot_data['Major Rotation(rad)'].abs() / beam_rot_data['CP']
        
        beam_rot_data = beam_rot_data[beam_rot_data['Name'].notna()]
        
    #%% 조작용 코드
        # 없애고 싶은 부재의 이름 입력(error_beam 확인 후!, DE, MCE에서 다 없어짐)
        beam_rot_data = beam_rot_data.drop(beam_rot_data[(beam_rot_data['Property Name'].str.contains('LB22_1'))].index)
        # beam_rot_data = beam_rot_data.drop(beam_rot_data[(beam_rot_data['Property Name'].str.contains('PB1-8_1'))].index)
        # beam_rot_data = beam_rot_data.drop(beam_rot_data[(beam_rot_data['Property Name'].str.contains('LB1A_2'))].index)
        # beam_rot_data = beam_rot_data.drop(beam_rot_data[(beam_rot_data['Property Name'].str.contains('LB1A_4'))].index)
        # beam_rot_data = beam_rot_data.drop(beam_rot_data[(beam_rot_data['Property Name'].str.contains('LB2_1'))].index)
        # beam_rot_data = beam_rot_data.drop(beam_rot_data[(beam_rot_data['Property Name'].str.contains('WB4B_'))].index)
        # beam_rot_data = beam_rot_data.drop(beam_rot_data[(beam_rot_data['Property Name'].str.contains('WB3D_'))].index)

    #%% DE 결과
        result = []
        
        if len(self.DE_load_name_list) != 0:            
            beam_rot_data_total_DE = pd.DataFrame()
            
            for load_name in self.DE_load_name_list:       
                temp_df_max = beam_rot_data[(beam_rot_data['Load Case'].str.contains('{}'.format(load_name)))\
                                            & (beam_rot_data['Step Type'] == 'Max')]\
                            .groupby(['Element Name'])['DE Rotation(rad)']\
                            .agg(**{'Rotation avg':'mean'})['Rotation avg']
                            
                beam_rot_data_total_DE['{}_max'.format(load_name)] = temp_df_max.tolist()
                
                temp_df_min = beam_rot_data[(beam_rot_data['Load Case'].str.contains('{}'.format(load_name)))\
                                            & (beam_rot_data['Step Type'] == 'Min')]\
                            .groupby(['Element Name'])['DE Rotation(rad)']\
                            .agg(**{'Rotation avg':'mean'})['Rotation avg']
                            
                beam_rot_data_total_DE['{}_min'.format(load_name)] = temp_df_min.tolist()
                
            beam_rot_data_total_DE['Element Name'] = temp_df_max.index            
            beam_rot_data_total_DE.reset_index(inplace=True, drop=True)
            
            beam_rot_data_total_DE = pd.merge(beam_rot_data_total_DE, element_data, how='left')
            beam_rot_data_total_DE = pd.merge(beam_rot_data_total_DE, node_data, how='left', left_on='I-Node ID', right_on='Node ID')
            beam_rot_data_total_DE = pd.merge(beam_rot_data_total_DE, self.story_info, how='left', left_on='V', right_on='Height(mm)')
            beam_rot_data_total_DE.sort_values('Height(mm)', inplace=True)
            # beam_rot_data_total_DE.reset_index(inplace=True, drop=True)
            
        # 평균 열 생성            
            beam_rot_data_total_DE['DE Max avg'] = beam_rot_data_total_DE.iloc[:,list(range(0,len(self.DE_load_name_list)*2,2))].mean(axis=1)
            beam_rot_data_total_DE['DE Min avg'] = beam_rot_data_total_DE.iloc[:,list(range(1,len(self.DE_load_name_list)*2,2))].mean(axis=1)

            result.append(beam_rot_data_total_DE)
            
    #%% MCE 결과 Plot
        
        if len(self.MCE_load_name_list) != 0:            
            beam_rot_data_total_MCE = pd.DataFrame()    
            
            for load_name in self.MCE_load_name_list:            
                temp_df_max = beam_rot_data[(beam_rot_data['Load Case'].str.contains('{}'.format(load_name)))\
                                            & (beam_rot_data['Step Type'] == 'Max')]\
                            .groupby(['Element Name'])['MCE Rotation(rad)']\
                            .agg(**{'Rotation avg':'mean'})['Rotation avg']
                            
                beam_rot_data_total_MCE['{}_max'.format(load_name)] = temp_df_max.tolist()
                
                temp_df_min = beam_rot_data[(beam_rot_data['Load Case'].str.contains('{}'.format(load_name)))\
                                            & (beam_rot_data['Step Type'] == 'Min')]\
                            .groupby(['Element Name'])['MCE Rotation(rad)']\
                            .agg(**{'Rotation avg':'mean'})['Rotation avg']
                            
                beam_rot_data_total_MCE['{}_min'.format(load_name)] = temp_df_min.tolist()                
            beam_rot_data_total_MCE['Element Name'] = temp_df_max.index
            
            beam_rot_data_total_MCE.reset_index(inplace=True, drop=True)
            
            beam_rot_data_total_MCE = pd.merge(beam_rot_data_total_MCE, element_data, how='left')
            beam_rot_data_total_MCE = pd.merge(beam_rot_data_total_MCE, node_data, how='left', left_on='I-Node ID', right_on='Node ID')
            beam_rot_data_total_MCE = pd.merge(beam_rot_data_total_MCE, self.story_info, how='left', left_on='V', right_on='Height(mm)')
            beam_rot_data_total_MCE.sort_values('Height(mm)', inplace=True)
            # beam_rot_data_total_MCE.reset_index(inplace=True, drop=True)
            
        # 평균 열 생성            
            beam_rot_data_total_MCE['MCE Max avg'] = beam_rot_data_total_MCE.iloc[:,list(range(0,len(self.MCE_load_name_list)*2,2))].mean(axis=1)
            beam_rot_data_total_MCE['MCE Min avg'] = beam_rot_data_total_MCE.iloc[:,list(range(1,len(self.MCE_load_name_list)*2,2))].mean(axis=1)

            result.append(beam_rot_data_total_MCE)

        return result, self.story_info
    
    # E.Beam Shear Force
    def E_BSF(self, contour=True):
        
    #%% Input Sheets 불러오기
        # E.Beam 데이터 불러오기
        transfer_element_info = self.e_beam_info.copy()
        SF_info_data = self.frame_SF_data.iloc[:,[0,1,2,3,5,6,10,11]]
        SF_info_data = SF_info_data.sort_values(by=['Load Case', 'Element Name', 'Step Type'])
        
        # Element 데이터 정리
        element_data = self.element_data.copy()
        # 필요한 부재만 선별
        element_data = element_data[element_data['Property Name'].isin(transfer_element_info)]

        # 기둥과 겹치는 등 평가에 반영하지 않을 부재 제거
        element_to_remove = ['E1702','E163', 'E120', 'E1542', 'E168', 'E129', 'E154', 'E1528', 'E1527', 'E184', 'E293']
        element_data = element_data[~element_data['Element Name'].isin(element_to_remove)]

        # Node 데이터 정리
        node_data = self.node_data.copy()
        # 나중에 element_info_data 열이름 깔끔하게 하기 위해 미리 깔끔하게
        i_node_info_data, j_node_info_data = node_data.copy(), node_data.copy()
        i_node_info_data.columns = ['I-Node ID', 'i-H1', 'i-H2', 'i-V']
        j_node_info_data.columns = ['J-Node ID', 'j-H1', 'j-H2', 'j-V']
        
        element_data = pd.merge(element_data, i_node_info_data, how='left')
        element_data = pd.merge(element_data, j_node_info_data, how='left')
        
        element_data = element_data.drop_duplicates()
        
        # 전단력, 부재 이름 Matching (by Element Name)
        SF_ongoing = pd.merge(element_data.iloc[:, [1,2,7]], SF_info_data.iloc[:, 1:], how='left')
        
        SF_ongoing = SF_ongoing.sort_values(by=['Element Name', 'Load Case', 'Step Type'])
        
        SF_ongoing.reset_index(inplace=True, drop=True)

        #%% (중력하중에 대한) V, M값에 절대값, 최대값 뽑기
    
        # 절대값 0.2배
        SF_ongoing_G = SF_ongoing.copy()
        SF_ongoing_G.iloc[:,[5,6,7,8]] = SF_ongoing_G.iloc[:,[5,6,7,8]].abs() * 0.2
        
        # i, j 노드 중 최대값 뽑기
        SF_ongoing_G['V2 max(G)'] = SF_ongoing_G[['V2 I-End', 'V2 J-End']].max(axis = 1)
        SF_ongoing_G['M3 max(G)'] = SF_ongoing_G[['M3 I-End', 'M3 J-End']].max(axis = 1)

        # max, min 중 최대값 뽑기
        SF_ongoing_G_max = SF_ongoing_G.loc[SF_ongoing_G.groupby(SF_ongoing_G.index // 2)['V2 max(G)'].idxmax()]
        SF_ongoing_G_max['M3 max(G)'] = SF_ongoing_G.groupby(SF_ongoing_G.index // 2)['M3 max(G)'].max().tolist()

        # 필요한 하중만 포함된 데이터 slice (MCE, G)
        SF_ongoing_G_max = SF_ongoing_G_max[SF_ongoing_G_max['Load Case']\
                                        .str.contains('|'.join(self.gravity_load_name))] 
                                        # function equivalent of a combination of df.isin() and df.str.contains()

        # 부재별 최대값 뽑기
        SF_ongoing_elements = SF_ongoing_G_max.iloc[:,[0,1,2]]
        SF_ongoing_elements= SF_ongoing_elements.drop_duplicates()
        SF_ongoing_elements.set_index('Element Name', inplace=True)
        
        SF_ongoing_G_max_avg = SF_ongoing_elements.copy() # 평균값을 뽑진 않지만, 아래의 SF_ongoing_max_avg와 형태 맞춰주기위해 이렇게 naming됨 
        SF_ongoing_G_max_avg['V2 max(G)'] = SF_ongoing_G_max.groupby(['Element Name'])['V2 max(G)'].max()
        SF_ongoing_G_max_avg['M3 max(G)'] = SF_ongoing_G_max.groupby(['Element Name'])['M3 max(G)'].max()

        # 같은 부재(그러나 잘려있는) 경우 최대값 뽑기
        SF_ongoing_G_max_max = SF_ongoing_G_max_avg.loc[SF_ongoing_G_max_avg.groupby(['Property Name'])['V2 max(G)'].idxmax()]
        SF_ongoing_G_max_max['M3 max(G)'] = SF_ongoing_G_max_avg.groupby(['Property Name'])['M3 max(G)'].max().tolist()
        
        SF_ongoing_G_max_max.reset_index(inplace=True, drop=True) 

        #%% V, M값에 절대값, 최대값, 평균값 뽑기

        # 절대값, 1.2배
        SF_ongoing.iloc[:,[5,6,7,8]] = SF_ongoing.iloc[:,[5,6,7,8]].abs() * 1.2

        # i, j 노드 중 최대값 뽑기
        SF_ongoing['V2 max'] = SF_ongoing[['V2 I-End', 'V2 J-End']].max(axis = 1)
        SF_ongoing['M3 max'] = SF_ongoing[['M3 I-End', 'M3 J-End']].max(axis = 1)

        # max, min 중 최대값 뽑기
        SF_ongoing_max = SF_ongoing.loc[SF_ongoing.groupby(SF_ongoing.index // 2)['V2 max'].idxmax()]
        SF_ongoing_max['M3 max'] = SF_ongoing.groupby(SF_ongoing.index // 2)['M3 max'].max().tolist()

        # 필요한 하중만 포함된 데이터 slice (MCE, G)
        SF_ongoing_max = SF_ongoing_max[SF_ongoing_max['Load Case']\
                                        .str.contains('|'.join(self.MCE_load_name_list))]

        # 부재별 평균값 뽑기
        SF_ongoing_max_avg = SF_ongoing_elements.copy()    
        SF_ongoing_max_avg['V2 max'] = SF_ongoing_max.groupby(['Element Name'])['V2 max'].mean()
        SF_ongoing_max_avg['M3 max'] = SF_ongoing_max.groupby(['Element Name'])['M3 max'].mean()
    
        # 같은 부재(그러나 잘려있는) 경우 최대값 뽑기
        SF_ongoing_max_avg_max = SF_ongoing_max_avg.loc[SF_ongoing_max_avg.groupby(['Property Name'])['V2 max'].idxmax()]
        SF_ongoing_max_avg_max['M3 max'] = SF_ongoing_max_avg.groupby(['Property Name'])['M3 max'].max().tolist()
        
        SF_ongoing_max_avg_max.reset_index(inplace=True, drop=True)

    #%% 결과값 정리 후 input sheets에 넣기
        
        SF_ongoing_max_avg_max = pd.merge(transfer_element_info.rename('Property Name'),\
                                        SF_ongoing_max_avg_max, how='left')
            
        SF_ongoing_max_avg_max = SF_ongoing_max_avg_max.dropna()
        SF_ongoing_max_avg_max.reset_index(inplace=True, drop=True)
        
        # 중력하중, 지진하중에 대한 V,M값 합치기
        SF_output = pd.merge(SF_ongoing_max_avg_max, SF_ongoing_G_max_max, how='left')

        # SF_ongoing_max_avg 재정렬
        SF_output = SF_output.iloc[:,[0,2,4,3,5]]  
        
        # nan인 칸을 ''로 바꿔주기 (win32com으로 nan입력시 임의의 숫자가 입력되기때문 ㅠ)
        SF_output = SF_output.replace(np.nan, '', regex=True)
        
        # Using win32com...
        # Call CoInitialize function before using any COM object
        excel = win32com.client.gencache.EnsureDispatch('Excel.Application', pythoncom.CoInitialize()) # 엑셀 실행
        excel.Visible = True # 엑셀창 안보이게

        wb = excel.Workbooks.Open(self.input_xlsx_path)
        ws = wb.Sheets('Results_E.Beam')
        
        startrow, startcol = 5, 1
        
        # 이름 열 입력
        ws.Range(ws.Cells(startrow, startcol),\
                ws.Cells(startrow + SF_output.shape[0]-1,\
                        startcol)).Value\
        = [[i] for i in SF_output.iloc[:,0]] # series -> list 형식만 입력가능
        
        # V, M열 입력
        ws.Range(ws.Cells(startrow, startcol+13),\
                ws.Cells(startrow + SF_output.shape[0]-1,\
                        startcol + 13 + 4 - 1)).Value\
        = list(SF_output.iloc[:,[1,2,3,4]].itertuples(index=False, name=None)) # dataframe -> tuple list 형식만 입력가능
        
        wb.Save()
        # wb.Close(SaveChanges=1) # Closing the workbook
        # excel.Quit() # Closing the application
        
    #%% 부재의 위치별  V, M 값 확인을 위한 도면 작성
        if contour == True:
        
            # DCR 구하기    
            # 1.2Vus, 1.2Mus, 0.2Vuns, 0.2Muns 불러오기
            SF_ongoing_combined = pd.concat([SF_ongoing_G_max_avg, SF_ongoing_max_avg.iloc[:,[2,3]]]
                                        , axis=1)
            SF_ongoing_combined.reset_index(inplace=True, drop=False)
            
            # Vu, Mu 구하기
            SF_ongoing_combined['Vu'] = SF_ongoing_combined['V2 max'] - SF_ongoing_combined['V2 max(G)']
            SF_ongoing_combined['Mu'] = SF_ongoing_combined['M3 max'] - SF_ongoing_combined['M3 max(G)']

            # Mu unit 변경 (mm -> m)
            SF_ongoing_combined['Mu'] = SF_ongoing_combined['Mu'] / 1000
            
            # Vn, Mn값 불러오기
            e_beam_result = pd.read_excel(self.input_xlsx_path, sheet_name='Results_E.Beam'
                                        , skiprows=3, header=0)
            e_beam_result = e_beam_result.iloc[:,[0,29,32]]
            e_beam_result.columns = ['Property Name', 'Vn', 'Mn']
            e_beam_result.reset_index(inplace=True, drop=True)
            
            # DCR 구하기
            SF_ongoing_combined = pd.merge(SF_ongoing_combined, e_beam_result, how='left')
            SF_ongoing_combined['DCR(V)'] = SF_ongoing_combined['Vu'] / SF_ongoing_combined['Vn']
            SF_ongoing_combined['DCR(M)'] = SF_ongoing_combined['Mu'] / SF_ongoing_combined['Mn']
            

            # 도면을 그리기 위한 Node List 만들기    
            node_map_z = SF_ongoing_max_avg['i-V'].drop_duplicates()
            node_map_z.sort_values(ascending=False, inplace=True)
            node_map_list = node_data[node_data['V'].isin(node_map_z)]
            
            # 도면을 그리기 위한 Element List 만들기
            element_map_list = pd.merge(SF_ongoing_combined.iloc[:,[0,1,2,11,12]]
                                        , element_data.iloc[:,[1,5,6,8,9]]
                                        , how='left', on='Element Name')
            
            return node_map_z, node_map_list, element_map_list        

    # E.Column Shear Force
    def E_CSF(self, export_to_pdf=True, pdf_name='E.Column Results', bldg_name='#'):
        ''' 
        Perform-3D 해석 결과에서 기둥의 축력, 전단력, 모멘트를 불러와 Results_E.Column 엑셀파일을 작성. \n
        result_path : Perform-3D에서 나온 해석 파일의 경로. \n
        result_xlsx : Perform-3D에서 나온 해석 파일의 이름. 해당 파일 이름이 포함된 파일들을 모두 불러온다. \n
        input_path : Data Conversion 엑셀 파일의 경로 \n
        input_xlsx : Data Conversion 엑셀 파일의 이름. result_xlsx와는 달리 확장자명(.xlsx)까지 기입해줘야한다. 하나의 파일만 불러온다. \n
        column_xlsx : Results_E.Column 엑셀 파일의 이름.확장자명(.xlsx)까지 기입해줘야한다. \n
        export_to_pdf : 입력된 값에 따른 각 부재들의 결과 시트를 pdf로 출력. True = pdf 출력, False = pdf 미출력(Results_E.Column 엑셀파일만 작성됨).
        pdf_name = 출력할 pdf 파일 이름.
        
        '''
    #%% Input sheets 불러오기
        # E.Column, Rebar 데이터 불러오기
        transfer_element_info = self.e_column_info.copy()
        rebar_info = self.rebar_info
        SF_info_data = self.frame_SF_data.iloc[:,[0,1,2,3,4,5,7,8,9,10,11]]
        SF_info_data = SF_info_data.sort_values(by=['Load Case', 'Element Name', 'Step Type'])

        # E.Column 데이터 정리
        transfer_element_info.iloc[:,3] = transfer_element_info.iloc[:,3].fillna(method='ffill')

        transfer_element_info.columns = ['Name', 'b(mm)', 'h(mm)', 'c(mm)', 'Concrete Grade', 'Main Bar Type', 'Main Bar Diameter',\
                                        'Hoop Bar Type', 'Hoop Bar Diameter', 'Layer 1 EA', 'Layer 1 Row', 'Layer 2 EA',\
                                        'Layer 2 Row', 'Hoop X', 'Hoop Y', 'Hoop Spacing(mm)', 'Direction']
        
        # Rebar 데이터 정리
        main_bar_info = transfer_element_info.iloc[:,[5,6]]
        hoop_bar_info = transfer_element_info.iloc[:,[7,8]]
        
        main_bar_info = pd.merge(main_bar_info, rebar_info,\
                                            how='left', left_on='Main Bar Diameter', right_on='Type')
        hoop_bar_info = pd.merge(hoop_bar_info, rebar_info,\
                                            how='left', left_on='Hoop Bar Diameter', right_on='Type')
        
        # Main Bar 강도 리스트 만들기
        main_bar_strength = []
        for idx, row in main_bar_info.iterrows():
            if row[0] == '일반용':
                main_bar_strength.append(row[3])
            elif row[0] == '내진용':
                main_bar_strength.append(row[4])
        
        # Hoop Bar 강도 리스트 만들기
        hoop_bar_strength = []
        for idx, row in hoop_bar_info.iterrows():
            if row[0] == '일반용':
                hoop_bar_strength.append(row[3])
            elif row[0] == '내진용':
                hoop_bar_strength.append(row[4])
        
        transfer_element_info['Main Bar Strength'] = main_bar_strength
        transfer_element_info['Hoop Bar Strength'] = hoop_bar_strength  

        # Element 데이터 정리
        element_data = self.element_data.iloc[:,[0,1,2,3]]
        # 필요한 부재만 선별
        element_data = element_data[element_data['Property Name'].isin(transfer_element_info['Name'])]
        
        # 층 정보 Matching을 위한 Node 정보
        node_data = self.node_data

        element_data = pd.merge(element_data, node_data.iloc[:,[0,3]], how='left', left_on='I-Node ID', right_on='Node ID')

        element_data = element_data.drop_duplicates()

        # 전단력, 부재 이름 Matching (by Element Name)
        SF_ongoing = pd.merge(element_data.iloc[:, [1,2,5]], SF_info_data.iloc[:, 1:], how='left')

        SF_ongoing = SF_ongoing.sort_values(by=['Element Name', 'Load Case', 'Step Type'])

        SF_ongoing.reset_index(inplace=True, drop=True)

    #%% V, M값에 절대값, 최대값, 평균값 뽑기
        # 절대값, 1.2배
        SF_ongoing.iloc[:,[5,6,7,8,9,10,11]] = SF_ongoing.iloc[:,[5,6,7,8,9,10,11]].abs() * 1.2

        # i, j 노드 중 최대값 뽑기
        SF_ongoing['M2 max'] = SF_ongoing[['M2 I-End', 'M2 J-End']].max(axis = 1)
        SF_ongoing['M3 max'] = SF_ongoing[['M3 I-End', 'M3 J-End']].max(axis = 1)

        # max, min 중 최대값 뽑기
        SF_ongoing_max = SF_ongoing.loc[SF_ongoing.groupby(SF_ongoing.index // 2)['P I-End'].idxmax()]
        SF_ongoing_max['V2 max'] = SF_ongoing.groupby(SF_ongoing.index // 2)['V2 I-End'].max().tolist()
        SF_ongoing_max['V3 max'] = SF_ongoing.groupby(SF_ongoing.index // 2)['V3 I-End'].max().tolist()
        SF_ongoing_max['M2 max'] = SF_ongoing.groupby(SF_ongoing.index // 2)['M2 max'].max().tolist()
        SF_ongoing_max['M3 max'] = SF_ongoing.groupby(SF_ongoing.index // 2)['M3 max'].max().tolist()

        # 필요한 하중만 포함된 데이터 slice (MCE)
        SF_ongoing_max = SF_ongoing_max[SF_ongoing_max['Load Case']\
                                        .str.contains('|'.join(self.MCE_load_name_list))] # function equivalent of a combination of df.isin() and df.str.contains()
        
        # 부재별 평균값 뽑기
        SF_ongoing_max_avg = SF_ongoing_max.iloc[:,[0,1,2]]
        SF_ongoing_max_avg = SF_ongoing_max_avg.drop_duplicates()
        SF_ongoing_max_avg.set_index('Element Name', inplace=True)
        
        SF_ongoing_max_avg['P'] = SF_ongoing_max.groupby(['Element Name'])['P I-End'].mean()
        SF_ongoing_max_avg['V2 max'] = SF_ongoing_max.groupby(['Element Name'])['V2 max'].mean()
        SF_ongoing_max_avg['V3 max'] = SF_ongoing_max.groupby(['Element Name'])['V3 max'].mean()
        SF_ongoing_max_avg['M2 max'] = SF_ongoing_max.groupby(['Element Name'])['M2 max'].mean()
        SF_ongoing_max_avg['M3 max'] = SF_ongoing_max.groupby(['Element Name'])['M3 max'].mean()
    
        # 같은 부재(그러나 잘려있는) 경우 최대값 뽑기
        SF_ongoing_max_avg_max = SF_ongoing_max_avg.loc[SF_ongoing_max_avg.groupby(['Property Name'])['P'].idxmax()]
        SF_ongoing_max_avg_max['V2 max'] = SF_ongoing_max_avg.groupby(['Property Name'])['V2 max'].max().tolist()
        SF_ongoing_max_avg_max['V3 max'] = SF_ongoing_max_avg.groupby(['Property Name'])['V3 max'].max().tolist()
        SF_ongoing_max_avg_max['M2 max'] = SF_ongoing_max_avg.groupby(['Property Name'])['M2 max'].max().tolist()
        SF_ongoing_max_avg_max['M3 max'] = SF_ongoing_max_avg.groupby(['Property Name'])['M3 max'].max().tolist()
        
        SF_ongoing_max_avg_max.reset_index(inplace=True, drop=True)

    #%% 결과값 정리
        SF_ongoing_max_avg_max = pd.merge(transfer_element_info['Name'].rename('Property Name'),\
                                        SF_ongoing_max_avg_max, how='left')
            
        SF_ongoing_max_avg_max = SF_ongoing_max_avg_max.dropna()
        SF_ongoing_max_avg_max.reset_index(inplace=True, drop=True)
        
        SF_output = pd.merge(SF_ongoing_max_avg_max, transfer_element_info,\
                            how='left', left_on='Property Name', right_on='Name')
        
        # 기존 시트에 V, M 값 넣기(alt2)
        SF_output = SF_output.iloc[:,[0,8,9,23,10,11,13,24,15,25,16,17,\
                                    18,19,20,21,22,2,5,6,4,3]] # SF_output 재정렬
        
    # nan인 칸을 ''로 바꿔주기 (win32com으로 nan입력시 임의의 숫자가 입력되기때문 ㅠ)

        SF_output = SF_output.replace(np.nan, '', regex=True)

    #%% 출력 (Using win32com...)    
        # Using win32com...
        # Call CoInitialize function before using any COM object
        excel = win32com.client.gencache.EnsureDispatch('Excel.Application', pythoncom.CoInitialize()) # 엑셀 실행
        excel.Visible = True # 엑셀창 안보이게

        # Column Sheets 템플릿 열기
        wb = excel.Workbooks.Open(os.getcwd() + '\\template\\Results_E.Column_Ver.1.3.xlsx')
        ws = wb.Sheets('Results')
        
        # 결과값 입력 시작 셀 지정
        startrow, startcol = 5, 2
        
        ws.Range(ws.Cells(startrow, startcol),\
                ws.Cells(startrow+SF_output.shape[0]-1,\
                        startcol+SF_output.shape[1]-1)).Value\
        = list(SF_output.itertuples(index=False, name=None)) # dataframe -> tuple list 형식만 입력가능

        # 동 이름 입력
        ws.Range(ws.Cells(1,20),\
                ws.Cells(1,20)).Value\
        = bldg_name
        
        # pdf로 출력
        if export_to_pdf == True:
            # pdf Merge를 위한 PdfMerger 클래스 생성
            merger = PdfMerger()

            # pdf 저장 위치
            result_path = os.path.dirname(self.result_xlsx_path[0])
            result_path = result_path.replace('/', '\\') # 경로 구분자 변경 (왜인지 모르겠는데 여기서만 인식이 안돼서...)

            for i in range(SF_output.shape[0]):

                pdf_file_path = os.path.join(result_path, pdf_name+'({}).pdf'.format(i+1))
                
                wb.Worksheets(2).Select()            
                wb.Worksheets(2).Name = '({})'.format(i+1)
                
                xlTypePDF = 0
                xlQualityStandard = 0
                
                wb.ActiveSheet.ExportAsFixedFormat(xlTypePDF, pdf_file_path
                                                   , xlQualityStandard, True, False)    
                merger.append(pdf_file_path)
                
            merger.write(result_path+'\\'+'{}.pdf'.format(pdf_name))
            merger.close()

        # Merge한 후 개별 파일들 지우기    
        for i in range(SF_output.shape[0]):
            pdf_file_path = os.path.join(result_path, pdf_name+'({}).pdf'.format(i+1))
            os.remove(pdf_file_path)

        # 저장할 Column Sheets 이름 생성
        name_count = 1
        column_xlsx_name = self.bldg_name + '_E.Column_Ver.1.3.xlsx'
        while True:
            if os.path.isfile(column_xlsx_name):
                column_xlsx_name = self.bldg_name + '_E.Column_Ver.1.3(' + str(name_count) + ').xlsx'
                name_count += 1
            else:
                break

        column_xlsx_path = os.path.join(result_path, column_xlsx_name)

        wb.SaveAs(Filename = column_xlsx_path)
        wb.Close()
        # wb.Close(SaveChanges=1) # Closing the workbook
        # excel.Quit() # Closing the application



    # Wall Axial Strain
    def WAS(self):
        # Wall Axial Strain, Node 데이터 불러오기
        AS_result_data = self.wall_AS_data[self.wall_AS_data['Load Case']\
                                           .str.contains('|'.join(self.seismic_load_name_list))]
        AS_gage_data = self.wall_AS_gage_data
        node_data = self.node_data
    
        # gage 개수 얻기
        gage_num = len(AS_gage_data)
        
        ### Gage data에서 Element Name, I-Node ID 불러와서 v좌표 match하기
        AS_gage_data = AS_gage_data[['Element Name', 'I-Node ID']]; 
        
        # I-Node의 v좌표 match해서 추가
        AS_gage_data = AS_gage_data.join(node_data.set_index('Node ID')[['H1', 'H2', 'V']], on='I-Node ID')
        
        ### AS_gage_data와 AS_result_data의 순서 맞추기 (Element Name열 기준으로)
        # AS_gage_data의 Element Name 열을 Index로 설정
        AS_gage_data = AS_gage_data.set_index('Element Name')

        # AS_result_data의 Element Name 열 추출
        elem_name_list = AS_result_data['Element Name'].drop_duplicates()

        # AS_gage_data를 추출된 AS_result_data의 Element Name 열에 맞게 재구성
        AS_gage_data = AS_gage_data.loc[elem_name_list]

        AS_gage_data.reset_index(drop=False, inplace=True)
    
        ### AS_total data 만들기
        AS_max = AS_result_data[(AS_result_data['Step Type'] == 'Max') & (AS_result_data['Performance Level'] == 1)][['Axial Strain']].values # dataframe을 array로
        AS_max = AS_max.reshape(gage_num, len(self.seismic_load_name_list), order='F') # order = 'C' 인 경우 row 우선 변경, order = 'F'인 경우 column 우선 변경
        AS_max = pd.DataFrame(AS_max) # array를 다시 dataframe으로
        AS_min = AS_result_data[(AS_result_data['Step Type'] == 'Min') & (AS_result_data['Performance Level'] == 1)][['Axial Strain']].values
        AS_min = AS_min.reshape(gage_num, len(self.seismic_load_name_list), order='F')
        AS_min = pd.DataFrame(AS_min)
        AS_total = pd.concat([AS_max, AS_min], axis=1)

        ### AS_avg_data 만들기
        DE_max_avg = AS_total.iloc[:, 0:len(self.DE_load_name_list)].mean(axis=1)
        MCE_max_avg = AS_total.iloc[:, len(self.DE_load_name_list) : len(self.DE_load_name_list)+len(self.MCE_load_name_list)].mean(axis=1)
        DE_min_avg = AS_total.iloc[:, len(self.DE_load_name_list)+len(self.MCE_load_name_list) : 2*len(self.DE_load_name_list)+len(self.MCE_load_name_list)].mean(axis=1)
        MCE_min_avg = AS_total.iloc[:, 2*len(self.DE_load_name_list)+len(self.MCE_load_name_list) : 2*len(self.DE_load_name_list) + 2*len(self.MCE_load_name_list)].mean(axis=1)
        AS_avg_total = pd.concat([AS_gage_data.loc[:, ['H1', 'H2', 'V']], DE_max_avg, DE_min_avg, MCE_max_avg, MCE_min_avg], axis=1)
        AS_avg_total.columns = ['X(mm)', 'Y(mm)', 'Z(mm)', 'DE_max_avg', 'DE_min_avg', 'MCE_max_avg', 'MCE_min_avg']

        ### 층분할된 곳의 Axial strain gage는 max(abs(분할된 두 값))로 assign하기
        # 분할층 노드가 포함되지 않은 부재 slice
        AS_avg_total_no_divide = AS_avg_total[AS_avg_total['Z(mm)'].isin(self.story_info['Height(mm)'])] 
        
        # i-node가 분할층에 있는 부재 slice
        AS_avg_total_divide = AS_avg_total[~AS_avg_total['Z(mm)'].isin(self.story_info['Height(mm)'])]   
    
        # AS_avg_total_divide 노드들의 i-node의 z좌표를 아래 층으로 격하
        next_level_list = []
        for i in AS_avg_total_divide['Z(mm)']:
            level_smaller = self.story_info['Height(mm)'][i-self.story_info['Height(mm)'] >= 0]
            next_level = level_smaller.sort_values(ignore_index=True, ascending=False)[0]
    
            next_level_list.append(next_level)
        
        pd.options.mode.chained_assignment = None # SettingWithCopyWarning 안뜨게 하기
    
        AS_avg_total_divide.loc[:,'Z(mm)'] = next_level_list
            
        # divide, no_divide 정보 concat
        AS_avg_total_joined = pd.concat([AS_avg_total_divide, AS_avg_total_no_divide]\
                                        , ignore_index=True)

        AS_output = AS_avg_total_joined.groupby(['X(mm)', 'Y(mm)', 'Z(mm)'])\
                    .agg({'DE_max_avg':'max', 'DE_min_avg':'min', 'MCE_max_avg':'max', 'MCE_min_avg':'min'})\
                        [['DE_max_avg', 'DE_min_avg', 'MCE_max_avg', 'MCE_min_avg']]
        
        AS_output.reset_index(inplace=True)
        
    #%% ***조작용 코드
        # 데이터 없애기 위한 기준값 입력
        # AS_output = AS_output.drop(AS_output[(AS_output.loc[:,'DE_min_avg'] < -0.002)].index)
        # AS_output = AS_output.drop(AS_output[(AS_output.loc[:,'MCE_min_avg'] < -0.002)].index)
        # .....위와 같은 포맷으로 계속

        return AS_output, self.story_info
    
    def WR(self):
        # Wall Rotation, Node, Element 데이터 불러오기
        wall_rot_data = self.wall_rot_data.copy()
        node_data = self.node_data.copy()
        element_data = self.wall_element_data.copy()
        gage_data = self.wall_rot_gage_data.copy()
        deformation_cap = self.wall_deform_cap.copy()

    #%% Gage Data & Result에 Node 정보 매칭
    
        gage_data = gage_data.drop_duplicates()
        node_data = node_data.drop_duplicates()

    #%% 데이터 매칭 후 결과뽑기

        ### Gage data에서 Element Name, I-Node ID 불러와서 v좌표 match하기
        gage_num = len(gage_data) # gage 개수 얻기
        
        # Gage data의 i, j-node 좌표
        gage_data = gage_data.join(node_data.set_index('Node ID')[['H1', 'H2', 'V']], on='I-Node ID')
        gage_data.rename({'H1' : 'I_H1', 'H2' : 'I_H2', 'V' : 'I_V'}, axis=1, inplace = True) # I node 의 H1, H2, V 좌표 가져오기
        gage_data = gage_data.join(node_data.set_index('Node ID')[['H1', 'H2']], on='J-Node ID')
        gage_data.rename({'H1' : 'J_H1', 'H2' : 'J_H2'}, axis=1, inplace=True) # J node 의 H1 좌표 가져오기
        
        # vector ij의 x, y방향 성분 구하기
        gage_data['J_H1-I_H1'] = gage_data.apply(lambda x: x['J_H1']- x['I_H1'], axis=1)
        gage_data['J_H2-I_H2'] = gage_data.apply(lambda x: x['J_H2']- x['I_H2'], axis=1)
        gage_data['I_H1-J_H1'] = gage_data.apply(lambda x: x['I_H1']- x['J_H1'], axis=1)
        gage_data['I_H2-J_H2'] = gage_data.apply(lambda x: x['I_H2']- x['J_H2'], axis=1)

        # gage 벡터, (1,0)벡터 만들기(array)
        gage_vector_ij = gage_data.iloc[:,[10,11]].values
        gage_vector_ji = gage_data.iloc[:,[12,13]].values
        e1_vector = np.array([1,0])
        e2_vector = np.array([0,1])

        # Vector ij와 (1,0)의 Cosine Similarity 구하기
        def cos_sim(arr, unit_arr):
            result = np.dot(arr, unit_arr) / (np.linalg.norm(arr, axis=1)*np.linalg.norm(unit_arr))
            return result
            
        gage_data['Similarity ij-e1'] = cos_sim(gage_vector_ij, e1_vector)
        gage_data['Similarity ij-e2'] = cos_sim(gage_vector_ij, e2_vector)
        gage_data['Similarity ji-e1'] = cos_sim(gage_vector_ji, e1_vector)
        gage_data['Similarity ji-e2'] = cos_sim(gage_vector_ji, e2_vector)


        # Wall element data의 i, j-node 좌표
        element_data = element_data.join(node_data.set_index('Node ID')[['H1', 'H2', 'V']], on='I-Node ID')
        element_data.rename({'H1' : 'I_H1', 'H2' : 'I_H2', 'V' : 'I_V'}, axis=1, inplace=True)
        element_data = element_data.join(node_data.set_index('Node ID')[['H1', 'H2']], on='J-Node ID')
        element_data.rename({'H1' : 'J_H1', 'H2' : 'J_H2'}, axis=1, inplace=True)
        
        # vector ij의 x, y방향 성분 구하기
        element_data['J_H1-I_H1'] = element_data.apply(lambda x: x['J_H1']- x['I_H1'], axis=1)
        element_data['J_H2-I_H2'] = element_data.apply(lambda x: x['J_H2']- x['I_H2'], axis=1)
        element_data['I_H1-J_H1'] = element_data.apply(lambda x: x['I_H1']- x['J_H1'], axis=1)
        element_data['I_H2-J_H2'] = element_data.apply(lambda x: x['I_H2']- x['J_H2'], axis=1)

        # element 벡터, (1,0)벡터 만들기(array)
        element_vector_ij = element_data.iloc[:,[11,12]].values
        element_vector_ji = element_data.iloc[:,[13,14]].values

        # Vector ij와 (1,0)의 Cosine Similarity 구하기
        element_data['Similarity ij-e1'] = cos_sim(element_vector_ij, e1_vector)
        element_data['Similarity ij-e2'] = cos_sim(element_vector_ij, e2_vector)
        element_data['Similarity ji-e1'] = cos_sim(element_vector_ji, e1_vector)
        element_data['Similarity ji-e2'] = cos_sim(element_vector_ji, e2_vector)

        ### wall element data 와 SWR gage data 연결하기(wall 이름)
        gage_data = gage_data.join(element_data.set_index(['I-Node ID', 'Similarity ij-e1', 'Similarity ij-e2'])\
                                ['Property Name'], on=['I-Node ID', 'Similarity ij-e1', 'Similarity ij-e2'])
        gage_data.rename({'Property Name' : 'gage_name'}, axis=1, inplace=True)
        
        # i, j 노드가 반대로 설정된 경우
        gage_data = gage_data.join(element_data.set_index(['I-Node ID', 'Similarity ij-e1', 'Similarity ij-e2'])\
                                ['Property Name'], on=['J-Node ID', 'Similarity ji-e1', 'Similarity ji-e2'])
        gage_data.rename({'Property Name' : 'gage_name'}, axis=1, inplace=True)
        
        # 위에서 join한 두 가지 경우의 이름 열 합치기
        for i in range(len(gage_data)):
            if pd.isnull(gage_data.iloc[i, 18]):
                gage_data.iloc[i, 18] = gage_data.iloc[i, 19]
        
        gage_data = gage_data.iloc[:, 0:19]
        
        
        wall_rot_data = wall_rot_data[wall_rot_data['Load Case']\
                                    .str.contains('|'.join(self.seismic_load_name_list))]
        
        ### SWR gage data와 SWR result data 연결하기(Element Name 기준으로)
        wall_rot_data = wall_rot_data.join(gage_data.set_index('Element Name')['gage_name'], on='Element Name')    
            
        ### SWR_total data 만들기
        SWR_max = wall_rot_data[(wall_rot_data['Step Type'] == 'Max') & (wall_rot_data['Performance Level'] == 1)][['Rotation']].values # dataframe을 array로
        SWR_max_gagename = wall_rot_data[(wall_rot_data['Step Type'] == 'Max') & (wall_rot_data['Performance Level'] == 1)][['gage_name']].values # dataframe을 array로
        SWR_max = SWR_max.reshape(gage_num, len(self.DE_load_name_list)+len(self.MCE_load_name_list), order='F') # order = 'C' 인 경우 row 우선 변경, order = 'F'인 경우 column 우선 변경
        SWR_max_gagename = SWR_max_gagename.reshape(gage_num, len(self.DE_load_name_list)+len(self.MCE_load_name_list), order='F') # order = 'C' 인 경우 row 우선 변경, order = 'F'인 경우 column 우선 변경
        SWR_max = pd.DataFrame(SWR_max) # array를 다시 dataframe으로
        SWR_max_gagename = pd.DataFrame(SWR_max_gagename) # array를 다시 dataframe으로
        
        SWR_min = wall_rot_data[(wall_rot_data['Step Type'] == 'Min') & (wall_rot_data['Performance Level'] == 1)][['Rotation']].values
        SWR_min_gagename = wall_rot_data[(wall_rot_data['Step Type'] == 'Min') & (wall_rot_data['Performance Level'] == 1)][['gage_name']].values
        SWR_min = SWR_min.reshape(gage_num, len(self.DE_load_name_list)+len(self.MCE_load_name_list), order='F')
        SWR_min_gagename = SWR_min_gagename.reshape(gage_num, len(self.DE_load_name_list)+len(self.MCE_load_name_list), order='F')
        SWR_min = pd.DataFrame(SWR_min)
        SWR_min_gagename = pd.DataFrame(SWR_min_gagename)
        
        SWR_total = pd.concat([gage_data['I_V'], SWR_max_gagename.iloc[:,0], SWR_max, SWR_min], axis=1)
        
        #SWR_total 의 column 명 만들기
        SWR_total_column_max = []
        for load_name in self.seismic_load_name_list:
            SWR_total_column_max.extend([load_name + '_max'])
            
        SWR_total_column_min = []
        for load_name in self.seismic_load_name_list:
            SWR_total_column_min.extend([load_name + '_min'])
        
        SWR_total.columns = ['Height', 'gage_name'] + SWR_total_column_max + SWR_total_column_min
        
        ### SWR_avg_data 만들기
        DE_max_avg = SWR_total.iloc[:, 2:len(self.DE_load_name_list)+2].mean(axis=1) # 2를 더해준 건 앞에 Height와 gage_name이 추가되었기 때문
        MCE_max_avg = SWR_total.iloc[:, len(self.DE_load_name_list)+2 : len(self.DE_load_name_list) + len(self.MCE_load_name_list)+2].mean(axis=1)
        DE_min_avg = SWR_total.iloc[:, len(self.DE_load_name_list)+len(self.MCE_load_name_list)+2 : 2*len(self.DE_load_name_list)+len(self.MCE_load_name_list)+2].mean(axis=1)
        MCE_min_avg = SWR_total.iloc[:, 2*len(self.DE_load_name_list)+len(self.MCE_load_name_list)+2 : 2*len(self.DE_load_name_list) + 2*len(self.MCE_load_name_list)+2].mean(axis=1)
        SWR_avg_total = pd.concat([SWR_total[['Height', 'gage_name']], DE_max_avg, DE_min_avg, MCE_max_avg, MCE_min_avg], axis=1)
        SWR_avg_total.columns = ['Height', 'gage_name', 'DE_max_avg', 'DE_min_avg', 'MCE_max_avg', 'MCE_min_avg']   
        
        # LS 기준
        deformation_cap_DE = pd.DataFrame()
        for i in range(len(deformation_cap)):
            if deformation_cap.iloc[i, 1] > deformation_cap.iloc[i, 2]:
                deformation_cap_DE = pd.concat([deformation_cap_DE, pd.Series(deformation_cap.iloc[i, 5])], ignore_index=True)
            else:
                deformation_cap_DE = pd.concat([deformation_cap_DE, pd.Series(deformation_cap.iloc[i, 6])], ignore_index=True)
        
        # CP 기준
        deformation_cap_MCE = pd.DataFrame()
        for i in range(len(deformation_cap)):
            if deformation_cap.iloc[i, 3] > deformation_cap.iloc[i, 4]:
                deformation_cap_MCE = pd.concat([deformation_cap_MCE, pd.Series(deformation_cap.iloc[i, 7])], ignore_index=True)
            else:
                deformation_cap_MCE = pd.concat([deformation_cap_MCE, pd.Series(deformation_cap.iloc[i, 8])], ignore_index=True)
        
        SWR_criteria = pd.concat([deformation_cap['Name'], deformation_cap_DE, deformation_cap_MCE], axis = 1, ignore_index=True)
        SWR_criteria.columns = ['Name', 'DE criteria', 'MCE criteria']
        
        ### SWR avg total에 SWR criteria join(wall name 기준)
        SWR_avg_total = pd.merge(SWR_avg_total, SWR_criteria, how='left'\
                                , left_on='gage_name', right_on='Name')
        
        #SWR_avg_total.dropna(inplace=True)
        SWR_avg_total['DCR_DE_min'] = SWR_avg_total['DE_min_avg'].abs()/SWR_avg_total['DE criteria']
        SWR_avg_total['DCR_DE_max'] = SWR_avg_total['DE_max_avg']/SWR_avg_total['DE criteria']
        SWR_avg_total['DCR_MCE_min'] = SWR_avg_total['MCE_min_avg'].abs()/SWR_avg_total['MCE criteria']
        SWR_avg_total['DCR_MCE_max'] = SWR_avg_total['MCE_max_avg']/SWR_avg_total['MCE criteria']

        print(deformation_cap_DE)
        
        #%% ***조작용 코드
        SWR_avg_total = SWR_avg_total[SWR_avg_total['DCR_DE_min'] <= 1] # DE
        SWR_avg_total = SWR_avg_total[SWR_avg_total['DCR_DE_max'] <= 1]
        SWR_avg_total = SWR_avg_total[SWR_avg_total['DCR_MCE_min'] <= 1] # MCE
        SWR_avg_total = SWR_avg_total[SWR_avg_total['DCR_MCE_max'] <= 1]
        

        return SWR_avg_total
    
    def WSF(self, graph=True):
        ''' 

        Perform-3D 해석 결과에서 벽체의 축력, 전단력(DE, MCE)을 불러와 Data Conversion 엑셀파일의 Results_Wall 시트를 작성하고, 벽체 전단력 DCR 그래프를 출력(optional). \n
        
        
        벽체 회전각과 기준에서 계산한 허용기준을 각각의 벽체에 대해 비교하여 DCR 방식으로 산포도 그래프를 출력.
        
        Parameters
        ----------
        input_path : str
                    Data Conversion 엑셀 파일의 경로.
                    
        input_xlsx : str
                    Data Conversion 엑셀 파일의 이름. result_xlsx와는 달리 확장자명(.xlsx)까지 기입해줘야한다. 하나의 파일만 불러온다.
                    
        result_path : str
                    Perform-3D에서 나온 해석 파일의 경로.
                    
        result_xlsx : str, optional, default='Analysis Result'
                    Perform-3D에서 나온 해석 파일의 이름. 해당 파일 이름이 포함된 파일들을 모두 불러온다.                  

        graph : bool, optional, default=True
                True = Data Conversion 엑셀 파일에 입력된 값으로 DCR 그래프를 그릴지 설정.
                False = Data Conversion 엑셀파일만 작성. (그래프 X)    
        
        DCR_criteria : float, optional, default=1
                    DCR 기준값.
                    
        yticks : int, optional, default=2
                그래프의 y축 눈금 간격(층간격). 층이 너무 높으면 y축에 너무 많은 층이 표기되기 때문에, 층간격을 조절해서 정돈된 그래프를 표기할 수 있다.

        xlim : int, optional, default=3
            그래프의 x축 limit 값. x축 limit 안의 값만 표기되므로, limit를 넘어가는 값을 확인하고 싶을 시에는 더 큰 xlim 값을 사용하면 된다.

        Yields
        -------
        Min, Max값 모두 출력됨. 
        
        fig1 : matplotlib.pyplot.figure or None
            DE(설계지진) 발생 시 벽체 회전각 DCR 그래프
        
        fig2 : matplotlib.pyplot.figure or None
            MCE(최대고려지진) 발생 시 벽체 회전각 DCR 그래프
        
        error_wall_DE : pandas.core.frame.DataFrame or None
                        DE(설계지진) 발생 시 DCR 기준값을 초과하는 벽체의 정보
                        
        error_wall_MCE : pandas.core.frame.DataFrame or None
                        MCE(최대고려지진) 발생 시 DCR 기준값을 초과하는 벽체의 정보                                          
        
        Raises
        -------
        
        References
        -------
        .. [1] "철근콘크리트 건축구조물의 성능기반 내진설계 지침", 대한건축학회, p.79, 2021
        
        '''

    #%% Input Sheet 정보 load
        # Wall 데이터 불러오기
        transfer_element_info = self.wall_info.copy()

        # 필요없는 전단력 제거(층전단력)
        wall_SF_data = self.section_data[self.section_data['Name'].str.count('_') == 2] # underbar가 두개 들어간 행만 선택
        wall_SF_data.reset_index(inplace=True, drop=True)

        #%% 중력하중에 대한 전단력 데이터 grouping

        shear_force_H1_G_data_grouped = pd.DataFrame()
        shear_force_H2_G_data_grouped = pd.DataFrame()
        
        # G를 max, min으로 grouping
        for load_name in self.gravity_load_name:
            shear_force_H1_G_data_grouped['G_H1_max'] = wall_SF_data[(wall_SF_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                    (wall_SF_data['Step Type'] == 'Max')]['H1(kN)'].values
                
            shear_force_H1_G_data_grouped['G_H1_min'] = wall_SF_data[(wall_SF_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                    (wall_SF_data['Step Type'] == 'Min')]['H1(kN)'].values

            shear_force_H2_G_data_grouped['G_H2_max'] = wall_SF_data[(wall_SF_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                    (wall_SF_data['Step Type'] == 'Max')]['H2(kN)'].values
                
            shear_force_H2_G_data_grouped['G_H2_min'] = wall_SF_data[(wall_SF_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                    (wall_SF_data['Step Type'] == 'Min')]['H2(kN)'].values   

        # all 절대값
        shear_force_H1_G_abs = shear_force_H1_G_data_grouped.abs()
        shear_force_H2_G_abs = shear_force_H2_G_data_grouped.abs()
        
        # 최대값 뽑기 & 0.2배
        shear_force_H1_G_max = 0.2 * shear_force_H1_G_abs.max(axis=1)
        shear_force_H2_G_max = 0.2 * shear_force_H2_G_abs.max(axis=1)

    #%% DE, MCE에 대한 전단력 데이터 Grouping

        shear_force_H1_DE_data_grouped = pd.DataFrame()
        shear_force_H2_DE_data_grouped = pd.DataFrame()
        shear_force_H1_MCE_data_grouped = pd.DataFrame()
        shear_force_H2_MCE_data_grouped = pd.DataFrame()

        # DE를 max, min으로 grouping
        for load_name in self.DE_load_name_list:
            shear_force_H1_DE_data_grouped['{}_H1_max'.format(load_name)] = wall_SF_data[(wall_SF_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                        (wall_SF_data['Step Type'] == 'Max')]['H1(kN)'].values
                
            shear_force_H1_DE_data_grouped['{}_H1_min'.format(load_name)] = wall_SF_data[(wall_SF_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                        (wall_SF_data['Step Type'] == 'Min')]['H1(kN)'].values

            shear_force_H2_DE_data_grouped['{}_H2_max'.format(load_name)] = wall_SF_data[(wall_SF_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                        (wall_SF_data['Step Type'] == 'Max')]['H2(kN)'].values
                
            shear_force_H2_DE_data_grouped['{}_H2_min'.format(load_name)] = wall_SF_data[(wall_SF_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                        (wall_SF_data['Step Type'] == 'Min')]['H2(kN)'].values   

        # MCE를 max, min으로 grouping
        for load_name in self.MCE_load_name_list:
            shear_force_H1_MCE_data_grouped['{}_H1_max'.format(load_name)] = wall_SF_data[(wall_SF_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                        (wall_SF_data['Step Type'] == 'Max')]['H1(kN)'].values
                
            shear_force_H1_MCE_data_grouped['{}_H1_min'.format(load_name)] = wall_SF_data[(wall_SF_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                        (wall_SF_data['Step Type'] == 'Min')]['H1(kN)'].values

            shear_force_H2_MCE_data_grouped['{}_H2_max'.format(load_name)] = wall_SF_data[(wall_SF_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                        (wall_SF_data['Step Type'] == 'Max')]['H2(kN)'].values
                
            shear_force_H2_MCE_data_grouped['{}_H2_min'.format(load_name)] = wall_SF_data[(wall_SF_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                        (wall_SF_data['Step Type'] == 'Min')]['H2(kN)'].values   

        if len(self.DE_load_name_list) != 0:

            # all 절대값
            shear_force_H1_DE_abs = shear_force_H1_DE_data_grouped.abs()
            shear_force_H2_DE_abs = shear_force_H2_DE_data_grouped.abs()            
            # 최대값 every 4 columns
            shear_force_H1_DE_max = shear_force_H1_DE_abs.groupby([[i//4 for i in range(0,2*len(self.DE_load_name_list))]], axis=1).max()
            shear_force_H2_DE_max = shear_force_H2_DE_abs.groupby([[i//4 for i in range(0,2*len(self.DE_load_name_list))]], axis=1).max()
            # 1.2 * 평균값
            shear_force_H1_DE_avg = 1.2 * shear_force_H1_DE_max.mean(axis=1)
            shear_force_H2_DE_avg = 1.2 * shear_force_H2_DE_max.mean(axis=1)
            
        else : 
            shear_force_H1_DE_avg = ''
            shear_force_H2_DE_avg = ''

        if len(self.MCE_load_name_list) != 0:

            # all 절대값
            shear_force_H1_MCE_abs = shear_force_H1_MCE_data_grouped.abs()
            shear_force_H2_MCE_abs = shear_force_H2_MCE_data_grouped.abs()            
            # 최대값 every 4 columns
            shear_force_H1_MCE_max = shear_force_H1_MCE_abs.groupby([[i//4 for i in range(0,2*len(self.MCE_load_name_list))]], axis=1).max()
            shear_force_H2_MCE_max = shear_force_H2_MCE_abs.groupby([[i//4 for i in range(0,2*len(self.MCE_load_name_list))]], axis=1).max()
            # 1.2 * 평균값
            shear_force_H1_MCE_avg = 1.2 * shear_force_H1_MCE_max.mean(axis=1)
            shear_force_H2_MCE_avg = 1.2 * shear_force_H2_MCE_max.mean(axis=1)
            
        else : 
            shear_force_H1_MCE_avg = ''
            shear_force_H2_MCE_avg = ''

    #%% V(축력) 값 뽑기

        # 축력 불러와서 Grouping
        axial_force_data = wall_SF_data[wall_SF_data['Load Case'].str.contains(self.gravity_load_name[0])]['V(kN)']
        # 절대값
        axial_force_abs = axial_force_data.abs()
        # result
        axial_force_abs.reset_index(inplace=True, drop=True)
        axial_force = axial_force_abs.groupby([[i//2 for i in range(0, len(axial_force_abs))]], axis=0).max()
        
    #%% 결과 정리 후 Input Sheets에 넣기

    # 출력용 Dataframe 만들기
        SF_output = pd.DataFrame()
        SF_output['Name'] = wall_SF_data['Name'].drop_duplicates()
        SF_output.reset_index(inplace=True, drop=True)

        SF_output['Nu'] = axial_force
        SF_output['1.2_DE_H1'] = shear_force_H1_DE_avg
        SF_output['1.2_DE_H2'] = shear_force_H2_DE_avg
        SF_output['1.2_MCE_H1'] = shear_force_H1_MCE_avg
        SF_output['1.2_MCE_H2'] = shear_force_H2_MCE_avg
        SF_output['0.2_G_DE_H1'] = shear_force_H1_G_max
        SF_output['0.2_G_DE_H2'] = shear_force_H2_G_max
        SF_output['0.2_G_MCE_H1'] = shear_force_H1_G_max
        SF_output['0.2_G_MCE_H2'] = shear_force_H2_G_max
        
        SF_output = pd.merge(SF_output, transfer_element_info, how='left')

        SF_output = SF_output.iloc[:,[0,1,2,3,4,5,6,7,8,9]] # SF_output 재정렬
        
    # nan인 칸을 ''로 바꿔주기 (win32com으로 nan입력시 임의의 숫자가 입력되기때문 ㅠ)

        SF_output = SF_output.replace(np.nan, '', regex=True)
        
    # 엑셀로 출력(Using win32com)
        
        # Using win32com...
        # Call CoInitialize function before using any COM object
        excel = win32com.client.gencache.EnsureDispatch('Excel.Application', pythoncom.CoInitialize()) # 엑셀 실행
        excel.Visible = True # 엑셀창 안보이게
        
        wb = excel.Workbooks.Open(self.input_xlsx_path)
        ws = wb.Sheets('Results_Wall')
        
        startrow, startcol = 5, 1
        
        # 이름 열 입력
        ws.Range(ws.Cells(startrow, startcol),\
                 ws.Cells(startrow + SF_output.shape[0]-1,\
                        startcol)).Value\
        = [[i] for i in SF_output.iloc[:,0]] # series -> list 형식만 입력가능
        
        # 축력, 전단력 열 입력
        ws.Range(ws.Cells(startrow, startcol+10),\
                 ws.Cells(startrow + SF_output.shape[0]-1,\
                        startcol + 10 + 9 - 1)).Value\
        = list(SF_output.iloc[:,[1,2,3,4,5,6,7,8,9]].itertuples(index=False, name=None)) # dataframe -> tuple list 형식만 입력가능
        
        wb.Save()
        # wb.Close(SaveChanges=1) # Closing the workbook
        # excel.Quit() # Closing the application 

        if graph == True:
            # Wall 정보 load
            wall_result = pd.read_excel(self.input_xlsx_path,
                                sheet_name='Results_Wall', skiprows=3, header=0)
            
            wall_result = wall_result.iloc[:, [0,29,31,33,35]]
            wall_result.columns = ['Name', 'DE_H1', 'DE_H2', 'MCE_H1', 'MCE_H2']
            wall_result.reset_index(inplace=True, drop=True)

        #%% ***조작용 코드
            # wall_name_to_delete = ['84A-W1_1','84A-W3_1_40F'] 
            # # 지우고싶은 층들을 대괄호 안에 입력(벽 이름만 입력하면 벽 전체 다 없어짐, 벽+층 이름 입력하면 특정 층의 벽만 없어짐)
            
            # for i in wall_name_to_delete:
            #     wall_result = wall_result[wall_result['Name'].str.contains(i) == False]
            
        #%% 벽체 해당하는 층 높이 할당
            floor = []
            for i in wall_result['Name']:
                floor.append(i.split('_')[-1])
            
            wall_result['Story Name'] = floor
            
            wall_result_output = pd.merge(wall_result, self.story_info.iloc[:,[1,2]], how='left')

            return wall_result_output

#%% 

class PrintDocx(QObject):
    def __init__(self, bldg_name):
        super().__init__()
        
        # template 불러와서 Document 생성
        # template = 성능기반 내진설계 보고서
        self.document = docx.Document("template/report_template.docx")
        
        # 동 이름 replace(paragraph level)
        for paragraph in self.document.paragraphs:
            for run in paragraph.runs:        
                if 'building_name' in run.text:
                    run.text = bldg_name
                    
        # 동 이름 replac(table level)
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if 'building_name' in run.text:
                                run.text = bldg_name

    def save_docx(self, result_xlsx_path, output_docx):
        # 결과 저장할 경로
        output_path = os.path.dirname(result_xlsx_path[0])
        # 결과 저장
        self.document.save(os.path.join(output_path, output_docx))

    def print_docx(self, item, result):
        # template 불러와서 Document 생성
        # template = 성능기반 내진설계 보고서
        self.document = docx.Document("template/report_template.docx")
        
        # 동 이름 replace(paragraph level)
        for paragraph in self.document.paragraphs:
            for run in paragraph.runs:        
                if 'building_name' in run.text:
                    run.text = self.bldg_name
                    
        # 동 이름 replac(table level)
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if 'building_name' in run.text:
                                run.text = self.bldg_name

        # tuple -> list 
        base_SF_list = list(base_SF_result)

        # Base Shear SF Plot
        # 결과값 classify & assign
        base_shear_H1 = base_SF_list[0]
        base_shear_H2 = base_SF_list[1]
        DE_load_name_list = base_SF_list[2]
        MCE_load_name_list = base_SF_list[3]
        base_SF_markers = base_SF_list[4]

        # Avg. Base Shear 표 작성
        # template의 1,2번 표 불러오기
        base_SF_values_table = self.document.tables[0]
        base_SF_plots_table = self.document.tables[1]
        
        # DE가 있는 경우, DE 값,그래프 채우기
        if 'DE' in base_SF_markers:
            # plot 그리기
            # H1_DE
            fig1, ax1 = plt.subplots(1,1)
            ax1.set_ylim(0, self.max_shear)
        
            ax1.bar(range(len(DE_load_name_list)), base_shear_H1.iloc[0, 0:len(DE_load_name_list)]\
                    , color='darkblue', edgecolor='k', label = 'Max. Base Shear')
            ax1.axhline(y= base_shear_H1.iloc[0, 0:len(DE_load_name_list)].mean(), color='r', linestyle='-', label='Average')
            ax1.set_xticks(range(14), range(1,15))
            
            ax1.set_xlabel('Ground Motion No.')
            ax1.set_ylabel('Base Shear(kN)')
            ax1.legend(loc = 2)
            ax1.set_title('X DE')
            
            base_SF_avg_DE_x = Decimal(str(base_shear_H1.iloc[0, 0:len(DE_load_name_list)].mean()))\
                .quantize(Decimal('.01'), rounding=ROUND_UP)        
            
            memfile = BytesIO()
            plt.savefig(memfile)           
            plt.close()
            
            # H2_DE
            fig2, ax2 = plt.subplots(1,1)
            ax2.set_ylim(0, self.max_shear)
            
            ax2.bar(range(len(DE_load_name_list)), base_shear_H2.iloc[0, 0:len(DE_load_name_list)], color='darkblue', edgecolor='k', label = 'Max. Base Shear')
            ax2.axhline(y= base_shear_H2.iloc[0, 0:len(DE_load_name_list)].mean(), color='r', linestyle='-', label='Average')
            ax2.set_xticks(range(14), range(1,15))
            
            ax2.set_xlabel('Ground Motion No.')
            ax2.set_ylabel('Base Shear(kN)')
            ax2.legend(loc = 2)
            ax2.set_title('Y DE')
            
            base_SF_avg_DE_y = Decimal(str(base_shear_H2.iloc[0, 0:len(DE_load_name_list)].mean()))\
                .quantize(Decimal('.01'), rounding=ROUND_UP)
            
            memfile2 = BytesIO()
            plt.savefig(memfile2)
            plt.close()

            # 첫번째 표에 avg 값 넣기
            values_row = base_SF_values_table.rows[4]
            values_cell_x = values_row.cells[2]
            values_cell_y = values_row.cells[3]
            values_para_x = values_cell_x.paragraphs[0]
            values_para_y = values_cell_y.paragraphs[0]
            values_para_x.text = f'{base_SF_avg_DE_x:,} kN' # 1000 자리마다 , 찍기
            values_para_y.text = f'{base_SF_avg_DE_y:,} kN'
            values_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
            values_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 두번째 표에 그래프 넣기            
            plots_row = base_SF_plots_table.rows[0]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(9))
            plots_run_y.add_picture(memfile2, width=Cm(9))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        if 'MCE' in base_SF_markers:
            # plot 그리기
            # H1_MCE
            fig3, ax3 = plt.subplots(1,1)
            ax3.set_ylim(0, self.max_shear)
            
            ax3.bar(range(len(MCE_load_name_list)), base_shear_H1\
                    .iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)]\
                    , color='darkblue', edgecolor='k', label = 'Max. Base Shear')
            ax3.axhline(y= base_shear_H1.iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)]\
                        .mean(), color='r', linestyle='-', label='Average')
            ax3.set_xticks(range(14), range(1,15))
            
            ax3.set_xlabel('Ground Motion No.')
            ax3.set_ylabel('Base Shear(kN)')
            ax3.legend(loc = 2)
            ax3.set_title('X MCE')
            
            base_SF_avg_MCE_x = Decimal(str(base_shear_H1.iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)].mean()))\
                .quantize(Decimal('.01'), rounding=ROUND_UP)
            
            memfile = BytesIO()
            plt.savefig(memfile)           
            plt.close()

            # H2_MCE
            fig4, ax4 = plt.subplots(1,1)
            ax4.set_ylim(0, self.max_shear)
            
            plt.bar(range(len(MCE_load_name_list)), base_shear_H2\
                    .iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)]\
                    , color='darkblue', edgecolor='k', label = 'Max. Base Shear')
            plt.axhline(y= base_shear_H2.iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)]\
                        .mean(), color='r', linestyle='-', label='Average')
            ax4.set_xticks(range(14), range(1,15))
            
            ax4.set_xlabel('Ground Motion No.')
            ax4.set_ylabel('Base Shear(kN)')
            ax4.legend(loc = 2)
            ax4.set_title('Y MCE')
            
            base_SF_avg_MCE_y = Decimal(str(base_shear_H2.iloc[0, len(DE_load_name_list):len(DE_load_name_list)+len(MCE_load_name_list)].mean()))\
                .quantize(Decimal('.01'), rounding=ROUND_UP)

            memfile2 = BytesIO()
            plt.savefig(memfile2)
            plt.close()            

            # 첫번째 표에 avg 값 넣기
            values_row = base_SF_values_table.rows[5]
            values_cell_x = values_row.cells[2]
            values_cell_y = values_row.cells[3]
            values_para_x = values_cell_x.paragraphs[0]
            values_para_y = values_cell_y.paragraphs[0]
            values_para_x.text = f'{base_SF_avg_MCE_x:,} kN' # 1000 자리마다 , 찍기
            values_para_y.text = f'{base_SF_avg_MCE_y:,} kN'
            values_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER # 입력된 값 center alignment
            values_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 두번째 표에 그래프 넣기            
            plots_row = base_SF_plots_table.rows[3]
            plots_cell_x = plots_row.cells[0]
            plots_cell_y = plots_row.cells[1]
            plots_para_x = plots_cell_x.paragraphs[0]
            plots_para_y = plots_cell_y.paragraphs[0]
            plots_run_x = plots_para_x.add_run()
            plots_run_y = plots_para_y.add_run()
            plots_run_x.add_picture(memfile, width=Cm(9))
            plots_run_y.add_picture(memfile2, width=Cm(9))
            plots_para_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plots_para_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
